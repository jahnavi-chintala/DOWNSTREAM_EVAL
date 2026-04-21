#!/usr/bin/env python3
"""
reference_shape_verify.py
-------------------------
Check that **candidate** eval outputs match **reference** *structure*:

  • JSON / YAML — same nested keys and list-element shapes as the reference
    (values may differ unless --strict-values).
  • Word (.docx) — same number of tables, identical **header row** text per table,
    and the same sequence of heading *styles* + **normalized** heading text
    (study ids / ISO dates replaced so labels stay comparable).

Browsers do not render .docx reliably; this uses **python-docx** instead of Selenium.

Examples
--------
  python reference_shape_verify.py \\
    --ref-json reference_specs/cmp_eval_B7981027.json \\
    --cand-json Outputs/eval/eval_report_B7981027_2026-04-06.json

  python reference_shape_verify.py \\
    --ref-docx reference_specs/CMP_Eval_Report_B7981027.docx \\
    --cand-docx Outputs/eval/CMP_Eval_Report_B7981027.docx
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, List, Tuple

try:
    import yaml
except ImportError:  # pragma: no cover
    yaml = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore


_STUDY_TOKEN = re.compile(r"\bB\d{7}\b", re.I)
_DATE_TOKEN = re.compile(r"\b\d{4}-\d{2}-\d{2}\b")


def normalize_heading_text(s: str) -> str:
    t = re.sub(r"\s+", " ", (s or "").strip())
    t = _STUDY_TOKEN.sub("<STUDY>", t)
    t = _DATE_TOKEN.sub("<DATE>", t)
    return t


def _cell_header_text(cell) -> str:
    if hasattr(cell, "text"):
        return re.sub(r"\s+", " ", (cell.text or "").strip())
    return ""


def extract_docx_signature(path: Path) -> dict:
    if Document is None:
        raise RuntimeError("python-docx is required for --ref-docx/--cand-docx")
    doc = Document(str(path))
    headings: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        st = getattr(p.style, "name", "") or ""
        if st.startswith("Heading"):
            raw = (p.text or "").strip()
            if raw:
                headings.append((st, normalize_heading_text(raw)))
    table_headers: List[List[str]] = []
    for tbl in doc.tables:
        if not tbl.rows:
            table_headers.append([])
            continue
        row0 = tbl.rows[0]
        table_headers.append([_cell_header_text(c) for c in row0.cells])
    return {"headings": headings, "table_headers": table_headers}


def compare_docx(ref_path: Path, cand_path: Path) -> List[str]:
    r = extract_docx_signature(ref_path)
    c = extract_docx_signature(cand_path)
    errs: List[str] = []
    if r["headings"] != c["headings"]:
        errs.append(
            "Heading sequence or normalized text differs (reference vs candidate):\n"
            f"  ref ({len(r['headings'])}):  {r['headings'][:12]}{'...' if len(r['headings']) > 12 else ''}\n"
            f"  cand ({len(c['headings'])}): {c['headings'][:12]}{'...' if len(c['headings']) > 12 else ''}"
        )
    if len(r["table_headers"]) != len(c["table_headers"]):
        errs.append(
            f"Table count mismatch: reference={len(r['table_headers'])} candidate={len(c['table_headers'])}"
        )
    n = min(len(r["table_headers"]), len(c["table_headers"]))
    for i in range(n):
        rh, ch = r["table_headers"][i], c["table_headers"][i]
        if rh != ch:
            errs.append(f"Table {i} header row mismatch:\n  ref:  {rh}\n  cand: {ch}")
    return errs


def _type_name(x: Any) -> str:
    if x is None:
        return "null"
    if isinstance(x, bool):
        return "bool"
    if isinstance(x, int) and not isinstance(x, bool):
        return "int"
    if isinstance(x, float):
        return "float"
    if isinstance(x, str):
        return "str"
    if isinstance(x, list):
        return "list"
    if isinstance(x, dict):
        return "dict"
    return type(x).__name__


def structure_diff(
    ref: Any,
    cand: Any,
    path: str,
    *,
    strict_values: bool,
) -> Tuple[List[str], List[str]]:
    """
    Returns (errors, warnings). Candidate must include all keys from ref at each dict level.
    """
    errors: List[str] = []
    warnings: List[str] = []

    if type(ref) != type(cand):
        errors.append(f"{path}: type mismatch ref={_type_name(ref)} cand={_type_name(cand)}")
        return errors, warnings

    if isinstance(ref, dict):
        rkeys, ckeys = set(ref.keys()), set(cand.keys())
        for k in sorted(rkeys - ckeys):
            errors.append(f"{path}.{k}: missing key in candidate (present in reference)")
        for k in sorted(ckeys - rkeys):
            warnings.append(f"{path}.{k}: extra key in candidate (not_in_reference)")
        for k in sorted(rkeys & ckeys):
            e, w = structure_diff(
                ref[k],
                cand[k],
                f"{path}.{k}" if path else k,
                strict_values=strict_values,
            )
            errors.extend(e)
            warnings.extend(w)
        return errors, warnings

    if isinstance(ref, list):
        if not ref:
            return errors, warnings
        if not cand:
            errors.append(f"{path}: reference list non-empty but candidate list empty")
            return errors, warnings
        e0, w0 = structure_diff(
            ref[0],
            cand[0],
            f"{path}[0]",
            strict_values=strict_values,
        )
        errors.extend(e0)
        warnings.extend(w0)
        if strict_values and ref != cand:
            errors.append(f"{path}: strict-values enabled but list bodies differ")
        return errors, warnings

    if strict_values and ref != cand:
        errors.append(f"{path}: value mismatch ref={ref!r} cand={cand!r}")
    return errors, warnings


def load_json(path: Path) -> Any:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def load_yaml(path: Path) -> Any:
    if yaml is None:
        raise RuntimeError("PyYAML required for YAML comparison (pip install pyyaml)")
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass
    if hasattr(sys.stderr, "reconfigure"):
        try:
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass

    ap = argparse.ArgumentParser(description="Compare eval output shape to reference (JSON/YAML/DOCX).")
    ap.add_argument("--ref-json", type=Path, default=None)
    ap.add_argument("--cand-json", type=Path, default=None)
    ap.add_argument("--ref-yaml", type=Path, default=None)
    ap.add_argument("--cand-yaml", type=Path, default=None)
    ap.add_argument("--ref-docx", type=Path, default=None)
    ap.add_argument("--cand-docx", type=Path, default=None)
    ap.add_argument(
        "--strict-values",
        action="store_true",
        help="Require list/scalar values to match reference exactly (default: keys/shape only).",
    )
    args = ap.parse_args()

    if not any(
        [
            args.ref_json and args.cand_json,
            args.ref_yaml and args.cand_yaml,
            args.ref_docx and args.cand_docx,
        ]
    ):
        ap.print_help()
        print(
            "\nProvide at least one pair: --ref-json/--cand-json, --ref-yaml/--cand-yaml, "
            "or --ref-docx/--cand-docx.",
            file=sys.stderr,
        )
        return 2

    all_errors: List[str] = []
    all_warnings: List[str] = []

    if args.ref_json and args.cand_json:
        for label, p in ("reference", args.ref_json), ("candidate", args.cand_json):
            if not p.exists():
                all_errors.append(f"{label} JSON not found: {p}")
        if not all_errors:
            ref, cand = load_json(args.ref_json), load_json(args.cand_json)
            e, w = structure_diff(ref, cand, "", strict_values=args.strict_values)
            all_errors.extend(e)
            all_warnings.extend(w)

    if args.ref_yaml and args.cand_yaml:
        yaml_missing = False
        for label, p in ("reference", args.ref_yaml), ("candidate", args.cand_yaml):
            if not p.exists():
                all_errors.append(f"{label} YAML not found: {p}")
                yaml_missing = True
        if not yaml_missing:
            try:
                ref, cand = load_yaml(args.ref_yaml), load_yaml(args.cand_yaml)
                e, w = structure_diff(ref, cand, "", strict_values=args.strict_values)
                all_errors.extend(e)
                all_warnings.extend(w)
            except Exception as ex:  # pragma: no cover
                all_errors.append(f"YAML load/compare failed: {ex}")

    if args.ref_docx and args.cand_docx:
        docx_missing = False
        for label, p in ("reference", args.ref_docx), ("candidate", args.cand_docx):
            if not p.exists():
                all_errors.append(f"{label} DOCX not found: {p}")
                docx_missing = True
        if not docx_missing:
            try:
                all_errors.extend(compare_docx(args.ref_docx, args.cand_docx))
            except Exception as ex:
                all_errors.append(f"DOCX compare failed: {ex}")

    for w in all_warnings:
        print(f"[WARN] {w}")
    for e in all_errors:
        print(f"[ERROR] {e}")

    if all_errors:
        print(f"\nFailed with {len(all_errors)} error(s), {len(all_warnings)} warning(s).")
        return 1
    print(f"OK — structure matches ({len(all_warnings)} warning(s)).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
