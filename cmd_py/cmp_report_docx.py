"""
CMP eval report → Word (.docx) via python-docx.
Layout aligns with PFIZER / CMP Evaluation Report template (JSON from ``build_eval_report``).
Uses reusable helpers: ``add_heading``, ``add_table``, ``style_table``.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

import cmp_eval_docx_spec as _cmp_spec

_BODY_PT = 10.0
_SECTION_HEAD_PT = 11
_TITLE_PT = 15

_FONT = "Arial"


def _apply_cmp_doc_defaults(doc: Document) -> None:
    """Arial 10pt, narrow margins — call once after Document()."""
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
_HEADER_SHADE = "D9D9D9"
_STRUCTURE_SHADE = "FFF2CC"
_HIGH_PRIORITY_SHADE = "FFE6E6"
_MISS_COLOR = RGBColor(0xC0, 0x00, 0x00)

_METRIC_KEYS_ORDER = (
    "M1_kri_recall",
    "M2_threshold_accuracy",
    "M3_qtl_recall",
    "M4_hallucinations",
)

_SECTION_CARD_ORDER = (
    "global_kris",
    "study_specific_kris",
    "qtls",
    "metadata",
)

_SCORE_BLOCK_FOR_CARD = {"metadata": "section_metadata"}

_SECTION_FALLBACK_LABELS = (
    ("global_kris", "Global Standard KRIs"),
    ("study_specific_kris", "Study-Specific KRIs"),
    ("qtls", "Quality Tolerance Limits"),
    ("section_metadata", "Section 1 Metadata"),
)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _paragraph_style_run(
    p: Paragraph,
    text: str,
    *,
    bold: bool = False,
    size_pt: float | None = None,
    rgb: RGBColor | None = None,
) -> None:
    r = p.add_run(text)
    r.bold = bold
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    if rgb is not None:
        r.font.color.rgb = rgb


def _clear_cell(cell) -> Paragraph:
    for child in list(cell._tc):
        if child.tag.endswith("}tcPr"):
            continue
        cell._tc.remove(child)
    return cell.add_paragraph()


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size_pt: float | None = None,
    rgb: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    p = _clear_cell(cell)
    _paragraph_style_run(p, str(text), bold=bold, size_pt=size_pt or _BODY_PT, rgb=rgb)
    if align is not None:
        p.alignment = align


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    if level <= 1:
        p = doc.add_paragraph()
        _paragraph_style_run(p, text, bold=True, size_pt=_SECTION_HEAD_PT)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    else:
        p = doc.add_paragraph()
        _paragraph_style_run(p, text, bold=True, size_pt=_BODY_PT + 1)
        p.paragraph_format.space_after = Pt(4)


def style_table(table: Table, *, header_row: bool = True) -> None:
    table.style = "Table Grid"
    table.autofit = False
    # Light-gray thin borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblB.append(el)
    tblPr.append(tblB)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    *,
    col_widths: Sequence[Any] | None = None,
    max_width_inches: float = 6.5,
    header_shade: str = _HEADER_SHADE,
    row_fill: Callable[[int, List[str]], Optional[str]] | None = None,
    cell_style: Callable[[int, int, Any], dict] | None = None,
) -> Table:
    ncols = len(headers)
    nrows = len(rows)
    table = doc.add_table(rows=max(1, 1 + nrows), cols=ncols)
    style_table(table)

    widths: List = list(col_widths) if col_widths else []
    if not widths and ncols:
        w = max_width_inches / ncols
        widths = [Inches(w)] * ncols

    hdr_cells = table.rows[0].cells
    for ci, h in enumerate(headers):
        _set_cell_text(hdr_cells[ci], h, bold=True, size_pt=_BODY_PT)
        _set_cell_shading(hdr_cells[ci], header_shade)
        if ci < len(widths):
            hdr_cells[ci].width = widths[ci]

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = row_fill(ri, list(map(str, row_data))) if row_fill else None
        for ci in range(ncols):
            cell = row.cells[ci]
            val = row_data[ci] if ci < len(row_data) else ""
            kw: dict = {}
            if cell_style:
                kw.update(cell_style(ri, ci, val) or {})
            rgb = kw.pop("rgb", None)
            bold = kw.pop("bold", False)
            align = kw.pop("align", None)
            size = float(kw.pop("size_pt", _BODY_PT))
            _set_cell_text(cell, str(val), bold=bold, size_pt=size, rgb=rgb, align=align)
            if fill:
                _set_cell_shading(cell, fill)
            if ci < len(widths):
                cell.width = widths[ci]

    doc.add_paragraph()
    return table


def _metric_target_str(m: Dict[str, Any]) -> str:
    t = m.get("target")
    if t is None:
        return ""
    if isinstance(t, float):
        if str(m.get("metric", "")).startswith("M4"):
            return str(int(t)) if t == int(t) else str(t)
        return f"{t * 100:.0f}%" if 0 < t <= 1 else f"{t:.0f}%"
    return str(t)


def _m1_recall_from_summary(sm: Dict[str, Any]) -> float | None:
    # Prefer canonical stable key, fall back to any study-specific variant
    for key in ["m1_kri_recall", *[k for k in sm if str(k).startswith("m1_kri_recall_")]]:
        v = sm.get(key)
        if v is not None:
            try:
                return float(v)
            except (TypeError, ValueError):
                pass
    return None


def _build_metric_table_rows(report: Dict[str, Any]) -> List[List[str]]:
    metrics_block = report.get("metrics") or {}
    rows: List[List[str]] = []
    for key in _METRIC_KEYS_ORDER:
        m = metrics_block.get(key) or {}
        if not m:
            continue
        if str(m.get("metric", "")).startswith("M4"):
            score_str = str(m.get("score", ""))
            tv = m.get("target")
            t_disp = str(int(tv)) if isinstance(tv, float) and tv == int(tv) else str(tv)
        else:
            score_str = str(m.get("score_pct", m.get("score", "")))
            t_disp = _metric_target_str(m)
        pf = "PASS ✓" if m.get("passed") else "FAIL ✗"
        src = _metric_source(key, report)
        rows.append([str(m.get("metric", key)), score_str, t_disp, pf, src])
    if rows:
        return rows

    sm = report.get("summary_metrics") or {}
    em = report.get("eval_metadata") or {}
    gt = em.get("ground_truth_sources") or []
    kri_src = str(gt[0]) if len(gt) > 0 else "cmp_kri_ground_truth.csv"
    qtl_src = str(gt[1]) if len(gt) > 1 else "cmp_qtl_ground_truth.csv"

    m1v = _m1_recall_from_summary(sm)
    if m1v is not None:
        t = sm.get("m1_target", 0.8)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        pf = "PASS ✓" if sm.get("m1_pass") else "FAIL ✗"
        rows.append(["M1 KRI Recall", f"{m1v * 100:.0f}%", t_disp, pf, kri_src])

    if sm.get("m2_threshold_accuracy") is not None:
        m2 = float(sm["m2_threshold_accuracy"])
        t = sm.get("m2_target", 0.9)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        pf = "PASS ✓" if sm.get("m2_pass") else "FAIL ✗"
        rows.append(["M2 Threshold Accuracy", f"{m2 * 100:.0f}%", t_disp, pf, kri_src])

    if sm.get("m3_qtl_recall") is not None:
        m3 = float(sm["m3_qtl_recall"])
        if m3 > 1.0:
            m3 = min(1.0, m3)
        t = sm.get("m3_target", 0.85)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        pf = "PASS ✓" if sm.get("m3_pass") else "FAIL ✗"
        rows.append(["M3 QTL Recall", f"{m3 * 100:.0f}%", t_disp, pf, qtl_src])

    if sm.get("m4_hallucinations") is not None:
        m4 = sm["m4_hallucinations"]
        tv = sm.get("m4_target", 0)
        score_str = str(int(m4)) if isinstance(m4, float) and m4 == int(m4) else str(m4)
        t_disp = str(int(tv)) if isinstance(tv, float) and tv == int(tv) else str(tv)
        pf = "PASS ✓" if sm.get("m4_pass") else "FAIL ✗"
        rows.append(["M4 Hallucinations", score_str, t_disp, pf, "iqmp_risk_id validation"])

    return rows


def _metric_source(metric_key: str, report: Dict[str, Any]) -> str:
    em = report.get("eval_metadata") or {}
    gt = em.get("ground_truth_sources") or []
    if metric_key in ("M1_kri_recall", "M2_threshold_accuracy"):
        return str(gt[0]) if len(gt) > 0 else "cmp_kri_ground_truth.csv"
    if metric_key == "M3_qtl_recall":
        return str(gt[1]) if len(gt) > 1 else "cmp_qtl_ground_truth.csv"
    return "iqmp/id validity + extra-item checks"


def _phase_display(phase: Any) -> str:
    p = str(phase or "—").strip().upper().replace("_", " ")
    if p.startswith("PHASE"):
        rest = p.replace("PHASE", "").strip()
        return f"Phase {rest}" if rest else p
    return p.title() if p else "—"


def _ta_phase_bucket_subtitle(report: Dict[str, Any]) -> str:
    em = report.get("eval_metadata") or {}
    ta = str(report.get("therapeutic_area") or em.get("therapeutic_area") or "—")
    phase = str(report.get("phase") or em.get("phase") or "—")
    bucket = em.get("yaml_bucket") or report.get("yaml_bucket")
    if bucket:
        return f"{ta} / {phase} — {bucket}"
    return f"{_slug_line(ta, phase)}"


def _build_cmp_metrics_spec_rows(report: Dict[str, Any]) -> List[List[str]]:
    """Always emit M1–M6 with stakeholder labels and sources."""
    study = str(report.get("study_id") or (report.get("eval_metadata") or {}).get("study_id", "?"))
    metrics_block = report.get("metrics") or {}
    sm = report.get("summary_metrics") or {}
    em = report.get("eval_metadata") or {}
    gt = em.get("ground_truth_sources") or []
    kri_src = str(gt[0]) if len(gt) > 0 else "cmp_kri_ground_truth.csv"
    qtl_src = str(gt[1]) if len(gt) > 1 else "cmp_qtl_ground_truth.csv"

    def pf(ok: Any) -> str:
        return "PASS ✓" if ok else "FAIL ✗"

    rows: List[List[str]] = []

    m1 = metrics_block.get("M1_kri_recall") or {}
    if m1:
        score_s = str(m1.get("score_pct", m1.get("score", "")))
        t_disp = _metric_target_str(m1)
        rows.append(
            [
                f"M1  KRI Recall ({study})",
                score_s,
                t_disp,
                pf(m1.get("passed")),
                _metric_source("M1_kri_recall", report),
            ]
        )
    else:
        m1v = _m1_recall_from_summary(sm)
        if m1v is not None:
            t = sm.get("m1_target", 0.8)
            t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
            rows.append([f"M1  KRI Recall ({study})", f"{m1v * 100:.0f}%", t_disp, pf(sm.get("m1_pass")), kri_src])
        else:
            rows.append([f"M1  KRI Recall ({study})", "—", "80%", "—", kri_src])

    m2 = metrics_block.get("M2_threshold_accuracy") or {}
    if m2:
        rows.append(
            [
                "M2  Threshold Accuracy",
                str(m2.get("score_pct", m2.get("score", ""))),
                _metric_target_str(m2),
                pf(m2.get("passed")),
                _metric_source("M2_threshold_accuracy", report),
            ]
        )
    elif sm.get("m2_threshold_accuracy") is not None:
        m2v = float(sm["m2_threshold_accuracy"])
        t = sm.get("m2_target", 0.9)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        rows.append(["M2  Threshold Accuracy", f"{m2v * 100:.0f}%", t_disp, pf(sm.get("m2_pass")), kri_src])
    else:
        rows.append(["M2  Threshold Accuracy", "—", "90%", "—", kri_src])

    m3 = metrics_block.get("M3_qtl_recall") or {}
    if m3:
        sp = m3.get("score_pct")
        if sp is not None:
            m3_score_disp = str(sp)
        else:
            sv = float(m3.get("score", 0) or 0)
            if sv > 1.0:
                sv = min(1.0, sv)
            m3_score_disp = f"{sv * 100:.0f}%"
        rows.append(
            [
                "M3  QTL Recall",
                m3_score_disp,
                _metric_target_str(m3),
                pf(m3.get("passed")),
                _metric_source("M3_qtl_recall", report),
            ]
        )
    elif sm.get("m3_qtl_recall") is not None:
        m3v = float(sm["m3_qtl_recall"])
        # Legacy bug: summary stored matched_count/gt which could exceed 1.0; recall is a fraction 0–1
        if m3v > 1.0:
            m3v = min(1.0, m3v)
        t = sm.get("m3_target", 0.85)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        rows.append(["M3  QTL Recall", f"{m3v * 100:.0f}%", t_disp, pf(sm.get("m3_pass")), qtl_src])
    else:
        rows.append(["M3  QTL Recall", "—", "85%", "—", qtl_src])

    m4 = metrics_block.get("M4_hallucinations") or {}
    if m4:
        tv = m4.get("target")
        t_disp = str(int(tv)) if isinstance(tv, float) and tv == int(tv) else str(tv)
        score_str = str(m4.get("score", ""))
        rows.append(
            [
                "M4  Hallucinations",
                score_str,
                t_disp,
                pf(m4.get("passed")),
                "iqmp_risk_id validation",
            ]
        )
    elif sm.get("m4_hallucinations") is not None:
        m4v = sm["m4_hallucinations"]
        tv = sm.get("m4_target", 0)
        score_str = str(int(m4v)) if isinstance(m4v, float) and m4v == int(m4v) else str(m4v)
        t_disp = str(int(tv)) if isinstance(tv, float) and tv == int(tv) else str(tv)
        rows.append(["M4  Hallucinations", score_str, t_disp, pf(sm.get("m4_pass")), "iqmp_risk_id validation"])
    else:
        rows.append(["M4  Hallucinations", "—", "0", "—", "iqmp_risk_id validation"])

    m5 = metrics_block.get("M5_forms_variables_accuracy") or metrics_block.get("M5_forms_variables") or {}
    if m5:
        p5 = m5.get("passed")
        rows.append(
            [
                "M5  forms_variables Accuracy",
                str(m5.get("score_pct", m5.get("score", ""))),
                _metric_target_str(m5) if m5.get("target") is not None else "95%",
                pf(p5) if p5 is not None else "—",
                "ASB.pdf (supplemental)",
            ]
        )
    elif sm.get("m5_forms_variables") is not None:
        m5v = sm["m5_forms_variables"]
        t = sm.get("m5_target", 0.95)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        try:
            pct = f"{float(m5v) * 100:.0f}%"
        except (TypeError, ValueError):
            pct = str(m5v)
        mp = sm.get("m5_pass")
        rows.append(
            ["M5  forms_variables Accuracy", pct, t_disp, pf(mp) if mp is not None else "—", "ASB.pdf (supplemental)"]
        )
    else:
        rows.append(
            [
                "M5  forms_variables Accuracy",
                "N/A",
                "95%",
                "N/A",
                "ASB.pdf not provided",
            ]
        )

    m6 = metrics_block.get("M6_data_sources_completeness") or metrics_block.get("M6_data_sources") or {}
    if m6:
        p6 = m6.get("passed")
        rows.append(
            [
                "M6  Data Sources Completeness",
                str(m6.get("score_pct", m6.get("score", ""))),
                _metric_target_str(m6) if m6.get("target") is not None else "90%",
                pf(p6) if p6 is not None else "—",
                "SDS Non-CRF tab (supplemental)",
            ]
        )
    elif sm.get("m6_data_sources") is not None:
        m6v = sm["m6_data_sources"]
        t = sm.get("m6_target", 0.9)
        t_disp = f"{float(t) * 100:.0f}%" if isinstance(t, (int, float)) and 0 < float(t) <= 1 else str(t)
        try:
            pct = f"{float(m6v) * 100:.0f}%"
        except (TypeError, ValueError):
            pct = str(m6v)
        m6p = sm.get("m6_pass")
        rows.append(
            [
                "M6  Data Sources Completeness",
                pct,
                t_disp,
                pf(m6p) if m6p is not None else "—",
                "SDS Non-CRF tab (supplemental)",
            ]
        )
    else:
        rows.append(
            [
                "M6  Data Sources Completeness",
                "N/A",
                "90%",
                "N/A",
                "SDS Non-CRF tab not provided",
            ]
        )

    return rows


def _slug_line(ta: str, phase: str) -> str:
    t = (ta or "unknown").strip().lower().replace(" ", "_").replace("&", "and")
    ph = (phase or "unknown").strip().lower().replace(" ", "_")
    return f"{t} / {ph}"


def _section_risk_status(score: float) -> str:
    return "PASS ✓" if float(score) >= 85.0 else "FAIL ✗"


def _build_section_score_rows(report: Dict[str, Any]) -> List[List[str]]:
    sec_score = report.get("section_scores") or {}
    scorecard = report.get("section_scorecard") or {}
    score_key = _SCORE_BLOCK_FOR_CARD
    srows: List[List[str]] = []

    if scorecard:
        for sk in _SECTION_CARD_ORDER:
            card = scorecard.get(sk) or {}
            name = card.get("name") or sk
            w = card.get("weight", 0)
            wp = f"{float(w) * 100:.0f}%" if isinstance(w, (int, float)) else str(w)
            sc = float(card.get("score", 0))
            wc = float(card.get("weighted_contribution", 0))
            block = sec_score.get(score_key.get(sk, sk), {})
            matched = block.get("matched")
            gt_ct = block.get("ground_truth_count")
            gen_ct = block.get("generated_count")
            if matched is not None and gt_ct is not None and gen_ct is not None:
                kri_q = f"{matched} / {gt_ct} (gen {gen_ct})"
            elif matched is not None and gt_ct is not None:
                kri_q = f"{matched} / {gt_ct}"
            else:
                kri_q = str(card.get("count", ""))
            status = _section_risk_status(sc)
            srows.append([name, wp, f"{sc:.1f}", f"{wc:.2f}", kri_q, status])
        return srows

    for block_key, label in _SECTION_FALLBACK_LABELS:
        block = sec_score.get(block_key, {})
        if not block:
            continue
        w = float(block.get("weight", 0))
        wp = f"{w * 100:.0f}%"
        sc = float(block.get("score", 0))
        wc = float(block.get("weighted_contribution", 0))
        matched = block.get("matched")
        gt_ct = block.get("ground_truth_count")
        kri_q = f"{matched} / {gt_ct}" if matched is not None and gt_ct is not None else "—"
        status = _section_risk_status(sc)
        srows.append([label, wp, f"{sc:.1f}", f"{wc:.2f}", kri_q, status])
    return srows


def _structure_note(sv: dict) -> str:
    errs = int(sv.get("error_count", 0))
    warns = int(sv.get("warning_count", 0))
    if errs == 0 and warns == 0:
        return "No structure errors or warnings — CluePoints compatible."
    if errs == 0:
        return (
            f"{warns} structure warning(s) — review paths in the JSON report for full CluePoints compatibility."
        )
    return "Structure errors present — see JSON report for paths and messages."


def _fmt_thr_pair(attr: dict | None) -> str:
    if not attr:
        return "—"
    g, t = attr.get("generated"), attr.get("ground_truth")
    if g is not None and t is not None:
        return f"{g} / {t}"
    if g is not None:
        return str(g)
    if t is not None:
        return str(t)
    return "—"


def _kri_table_rows(items: List[dict]) -> List[List[str]]:
    rows: List[List[str]] = []
    for item in items:
        ms = (item.get("match_status") or "").lower()
        is_miss = ms == "miss"
        gen = item.get("generated_label")
        glab = item.get("ground_truth_label") or "—"
        if is_miss:
            gen_disp = "**MISS**"
        else:
            gen_disp = str(gen) if gen is not None else "—"
        attrs = item.get("attributes") or {}
        mod = attrs.get("moderate_threshold") if isinstance(attrs, dict) else None
        high = attrs.get("high_threshold") if isinstance(attrs, dict) else None
        iq = attrs.get("iqmp_risk_id") if isinstance(attrs, dict) else None
        iqmp_s = "—"
        if isinstance(iq, dict):
            iqmp_s = str(iq.get("generated", iq.get("ground_truth", "—")))

        score_s = f"{float(item.get('kri_score', 0)):.1f}"
        match_s = str(item.get("match_status") or ("MISS" if is_miss else "—")).upper()

        rows.append(
            [
                gen_disp,
                str(glab),
                score_s,
                match_s,
                _fmt_thr_pair(mod if isinstance(mod, dict) else None),
                _fmt_thr_pair(high if isinstance(high, dict) else None),
                iqmp_s,
            ]
        )
    return rows


def _qtl_table_rows(items: List[dict]) -> List[List[str]]:
    rows = []
    for q in items:
        rows.append(
            [
                str(q.get("generated_name") or "—"),
                str(q.get("ground_truth_name") or q.get("gt_name") or "—"),
                str(q.get("match_status") or "—").upper(),
                _pct_cell(q.get("expectation_score")),
                _pct_cell(q.get("tolerance_score")),
                f"{float(q.get('qtl_score', 0)):.1f}",
            ]
        )
    return rows


def _pct_cell(v: Any) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
        if f <= 1.0 and f >= 0:
            return f"{f * 100:.1f}%"
        return f"{f:.1f}%"
    except (TypeError, ValueError):
        return str(v)


def _add_structure_block(doc: Document, sv: dict) -> None:
    add_heading(doc, "Structure validation")
    passed = sv.get("passed", False)
    res = "PASS ✓" if passed else "FAIL ✗"
    score = float(sv.get("structure_score", 0))
    errs = int(sv.get("error_count", 0))
    warns = int(sv.get("warning_count", 0))

    t = doc.add_table(rows=1, cols=1)
    style_table(t)
    cell = t.rows[0].cells[0]
    _set_cell_shading(cell, _STRUCTURE_SHADE)
    p = _clear_cell(cell)
    _paragraph_style_run(p, f"Result: {res}", bold=True, size_pt=_BODY_PT)
    p = cell.add_paragraph()
    _paragraph_style_run(p, f"Score: {score:.0f} / 100", size_pt=_BODY_PT)
    p = cell.add_paragraph()
    _paragraph_style_run(p, f"Errors: {errs}", size_pt=_BODY_PT)
    p = cell.add_paragraph()
    _paragraph_style_run(p, f"Warnings: {warns}", size_pt=_BODY_PT)
    p = cell.add_paragraph()
    _paragraph_style_run(p, _structure_note(sv), size_pt=_BODY_PT)
    doc.add_paragraph()


def _add_simple_data_table(
    doc: Document,
    headers: Sequence[str],
    data_rows: Sequence[Sequence[Any]],
    *,
    col_widths: Sequence[Any] | None = None,
) -> None:
    """One header row + data (uses ``add_table``)."""
    add_table(doc, list(headers), [list(r) for r in data_rows], col_widths=list(col_widths) if col_widths else None)


def _add_single_row_table(doc: Document, values: Sequence[str]) -> None:
    """Single body row, no separate header (reference breakdown cards)."""
    n = len(values)
    if n == 0:
        return
    t = doc.add_table(rows=1, cols=n)
    style_table(t)
    for i, val in enumerate(values):
        _set_cell_text(t.rows[0].cells[i], str(val), size_pt=_BODY_PT)
    doc.add_paragraph()


def _add_banner_one_cell(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    style_table(t)
    cell = t.rows[0].cells[0]
    p = _clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_style_run(p, text.replace("\n", " "), bold=True, size_pt=_TITLE_PT - 2)
    doc.add_paragraph()


def _add_document_score_table_three(doc: Document, left: str, mid: str, right: str) -> None:
    t = doc.add_table(rows=1, cols=3)
    style_table(t)
    cells = t.rows[0].cells
    for ci, (raw, bold_hdr) in enumerate(
        (
            (left, True),
            (mid, True),
            (right, False),
        )
    ):
        cell = cells[ci]
        p = _clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
        lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
        for li, line in enumerate(lines):
            if li:
                p = cell.add_paragraph()
            _paragraph_style_run(
                p,
                line,
                bold=bold_hdr,
                size_pt=_TITLE_PT if ci == 1 and li == 0 else _SECTION_HEAD_PT if ci == 0 else _BODY_PT,
            )
    doc.add_paragraph()


def _emit_cmp_block_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    r0 = rows[0]
    if len(rows) == 1 and len(r0) == 1:
        txt = r0[0]
        if "out of scope" in txt.lower():
            p = doc.add_paragraph()
            _paragraph_style_run(p, txt, size_pt=9)
            return
        if "MISS:" in txt.upper() or "NEAR MISS" in txt.upper():
            t = doc.add_table(rows=1, cols=1)
            style_table(t)
            cell = t.rows[0].cells[0]
            _set_cell_shading(cell, _STRUCTURE_SHADE)
            p = _clear_cell(cell)
            _paragraph_style_run(p, txt, size_pt=_BODY_PT)
            doc.add_paragraph()
        else:
            _add_banner_one_cell(doc, txt)
        return
    if len(rows) == 1 and len(r0) == 3 and "DOCUMENT SCORE" in r0[0].upper():
        _add_document_score_table_three(doc, r0[0], r0[1], r0[2])
        return
    if len(rows) == 1:
        _add_single_row_table(doc, r0)
        return
    hdr = [str(x) for x in r0]
    data = [[str(c) for c in row] for row in rows[1:]]
    ncols = len(hdr)
    default_w = max(6.5 / max(ncols, 1), 0.65)
    widths = [Inches(default_w)] * ncols
    _add_simple_data_table(doc, hdr, data, col_widths=widths)


def _write_cmp_eval_docx_from_blocks(report: Dict[str, Any], output_path: str) -> None:
    """Replay ``cmp_docx_blocks`` from reference JSON (paragraph/table order from Word)."""
    blocks = report.get("cmp_docx_blocks") or []
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_cmp_doc_defaults(doc)

    for block in blocks:
        if not block or len(block) < 2:
            continue
        kind, payload = block[0], block[1]
        if kind == "p":
            text = str(payload)
            if re.match(r"^\d+\.\s", text) or text.startswith(("Global ", "Study-", "Quality ", "Actions")):
                add_heading(doc, text)
            else:
                p = doc.add_paragraph()
                _paragraph_style_run(p, text, size_pt=_BODY_PT)
        else:
            _emit_cmp_block_table(doc, list(payload))
    em = report.get("eval_metadata") or {}
    study = report.get("study_id") or em.get("study_id", "?")
    eval_dt = report.get("eval_date") or em.get("eval_date", "—")
    fp = doc.add_paragraph()
    _paragraph_style_run(
        fp,
        f"End of eval report. Source JSON: cmp_eval_{study}.json | Config: {em.get('config_file', 'cmp_eval_config.yaml')}",
        size_pt=9,
    )
    doc.save(str(path))


def write_cmp_eval_docx(report: Dict[str, Any], output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    if report.get("cmp_docx_blocks"):
        _write_cmp_eval_docx_from_blocks(report, output_path)
        return

    doc = Document()
    _apply_cmp_doc_defaults(doc)

    em = report.get("eval_metadata") or {}
    study = report.get("study_id") or em.get("study_id", "?")
    ta = report.get("therapeutic_area") or em.get("therapeutic_area", "—")
    phase = report.get("phase") or em.get("phase", "—")
    eval_dt = report.get("eval_date") or em.get("eval_date", "—")
    gen_ver = em.get("generator_version") or report.get("generator_version", "v1.0")
    if gen_ver and not str(gen_ver).lower().startswith("v"):
        gen_ver = f"v{gen_ver}"
    cfg_ver = em.get("config_version") or report.get("config_version", "1.0")

    drug = str(em.get("drug_name") or report.get("drug_name") or "—")
    indication = str(em.get("indication") or report.get("indication") or "—")
    line3 = f"{study} — {drug} {_phase_display(phase)} {indication}"
    line4 = f"Eval date: {eval_dt} | Generator: {gen_ver} | Config: {cfg_ver}"
    _cmp_spec.add_dark_header_table(
        doc,
        "PFIZER PROTOCOL INTELLIGENCE PLATFORM",
        "CMP Generator — Eval Report",
        line3,
        line4,
    )

    ds = float(report.get("document_score", 0))
    doc_pass = bool(report.get("document_passed", report.get("document_pass")))
    thresh = float(report.get("pass_threshold", report.get("document_pass_threshold", 75)))
    tgt = float(report.get("target", report.get("document_target", 80)))
    _cmp_spec.add_document_score_banner(
        doc,
        ta_phase_bucket=_ta_phase_bucket_subtitle(report),
        score=ds,
        passed=doc_pass,
        threshold=thresh,
        target=tgt,
    )
    dsb = report.get("document_score_breakdown") or {}
    if dsb.get("pre_structure_score") is not None and dsb.get("structure_factor") is not None:
        p = doc.add_paragraph()
        _paragraph_style_run(
            p,
            f"Score calc: weighted sum {dsb.get('pre_structure_score')} x structure factor {dsb.get('structure_factor')}",
            size_pt=9,
        )

    doc.add_heading("1. Summary Metrics", level=2)
    doc.add_paragraph()
    mrows = _build_cmp_metrics_spec_rows(report)
    _cmp_spec.add_metrics_table_formatted(
        doc,
        ["Metric", "Score", "Target", "Pass/Fail", "Source"],
        mrows,
    )

    doc.add_heading("2. Section Scorecard", level=2)
    doc.add_paragraph()
    srows = _build_section_score_rows(report)
    if srows:
        _cmp_spec.add_section_scorecard_table(doc, srows)
        qtl_m3 = (report.get("metrics", {}) or {}).get("M3_qtl_recall", {})
        if qtl_m3:
            note = doc.add_paragraph()
            _paragraph_style_run(
                note,
                "Note: M3 is GT recall only; QTL section score also penalizes extra generated QTLs.",
                size_pt=9,
            )
        doc.add_heading("Section scores (weighted contribution)", level=3)
        doc.add_paragraph()
        for row in srows:
            try:
                sc = float(row[2])
            except (TypeError, ValueError, IndexError):
                sc = 0.0
            _cmp_spec.add_two_col_bar_row(doc, row[0], sc)

    doc.add_heading("3. KRI Detail", level=2)
    doc.add_paragraph()

    g_items = report.get("global_kris") or []
    if g_items:
        doc.add_heading("Global Standard KRIs", level=3)
        doc.add_paragraph()
        _cmp_spec.add_kri_eight_table_with_callouts(doc, g_items, yellow_miss_notes=False)

    ss_items = report.get("study_specific_kris") or []
    if ss_items:
        doc.add_heading("Study-Specific KRIs", level=3)
        doc.add_paragraph()
        _cmp_spec.add_kri_eight_table_with_callouts(doc, ss_items, yellow_miss_notes=True)

    qtls = report.get("qtls") or []
    if qtls:
        doc.add_heading("Quality Tolerance Limits", level=3)
        doc.add_paragraph()
        _cmp_spec.add_qtl_table_with_near_miss_notes(doc, qtls)

    sv = dict(report.get("structure_validation") or {})
    if not sv:
        # Backfill counts from improvement actions when aligned reference shape omits structure block.
        imp = report.get("improvement_actions") or []
        se = [a for a in imp if str(a.get("type")) == "structure_error"]
        sw = [a for a in imp if str(a.get("type")) == "structure_warning"]
        sv = {
            "passed": len(se) == 0,
            "structure_score": 0.0 if se else 100.0,
            "error_count": len(se),
            "warning_count": len(sw),
            "errors": [],
            "warnings": [],
        }
    _add_structure_block(doc, sv)

    actions = report.get("improvement_actions") or []
    doc.add_heading("4. Improvement Actions", level=2)
    doc.add_paragraph()
    intro = doc.add_paragraph()
    ir = intro.add_run("Actions derived from eval scoring. Prioritised for generator developer.")
    ir.italic = True
    doc.add_paragraph()
    if actions:
        _cmp_spec.add_improvement_table(doc, _cmp_spec.improvement_rows_sorted(list(actions)))

    oos = em.get("out_of_scope_sections_excluded") or []
    oos_reason = em.get("out_of_scope_reason")
    oos_txt = "Out of scope — not penalised\n"
    if oos:
        oos_txt += ", ".join(str(x) for x in oos)
        oos_txt += " — excluded from scoring."
    else:
        oos_txt += (
            "DQA (Data Quality Assessment) and Duplicate Patient Analysis — excluded from scoring."
        )
    if oos_reason:
        oos_txt += f" {oos_reason}"
    _cmp_spec.add_shaded_callout(doc, oos_txt, _cmp_spec.GREY_OOS)

    cfg_name = em.get("config_file") or "cmp_eval_config.yaml"
    _cmp_spec.footer_cmp(doc, study, cfg_name, version=str(gen_ver))

    form_ref = em.get("cmp_form_reference")
    if form_ref:
        fp = doc.add_paragraph()
        present = em.get("cmp_form_present")
        suffix = "" if present else " (path from config — file not found at eval time)"
        _paragraph_style_run(fp, f"CMP form reference: {form_ref}{suffix}", size_pt=9)

    doc.save(str(path))


def write_cmp_eval_docx_from_json(json_path: str, output_path: str | None = None) -> str:
    p = Path(json_path)
    with open(p, encoding="utf-8") as f:
        report = json.load(f)
    study = report.get("study_id", p.stem)
    out = output_path or str(p.parent / f"CMP_Eval_Report_{study}.docx")
    write_cmp_eval_docx(report, out)
    return out
