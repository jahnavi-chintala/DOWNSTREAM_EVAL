"""
Report builder — renders D5_Eval_Report_{study}.docx from eval JSON results.
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


_GREEN = RGBColor(0x00, 0x80, 0x00)
_RED = RGBColor(0xCC, 0x00, 0x00)
_ORANGE = RGBColor(0xFF, 0x8C, 0x00)


def _status_color(status: str) -> RGBColor:
    if "PASS" in status:
        return _GREEN
    if "FAIL" in status or "STOP" in status:
        return _RED
    return _ORANGE


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_metric_table(doc: Document, metrics: dict):
    """Add the scorecard summary table."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metric", "Score", "Target", "Status", "Detail"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    metric_order = ["m1_global", "m1_country", "m2", "m3", "m4", "m5"]
    labels = {
        "m1_global": "M1 Global — Domain Recall",
        "m1_country": "M1 Country — Domain Recall",
        "m2": "M2 — Category Accuracy",
        "m3": "M3 — Hallucination Rate",
        "m4": "M4 — Lineage Completeness",
        "m5": "M5 — Confidence Calibration",
    }

    for key in metric_order:
        m = metrics.get(key, {})
        if m.get("score") is None and m.get("status") == "N/A":
            continue
        row = table.add_row().cells
        row[0].text = labels.get(key, key)
        score = m.get("score")
        row[1].text = f"{score:.3f}" if score is not None else "N/A"
        target = m.get("target")
        row[2].text = f"{target:.2f}" if target is not None else "—"
        status = m.get("status", "N/A")
        row[3].text = status
        for p in row[3].paragraphs:
            for run in p.runs:
                run.font.color.rgb = _status_color(status)
                run.bold = True

        detail_parts = []
        if "numerator" in m:
            detail_parts.append(f"{m['numerator']}/{m['denominator']}")
        if "hallucinations" in m:
            detail_parts.append(f"{len(m['hallucinations'])} hallucinations")
        if "compliant_changes" in m:
            detail_parts.append(f"{m['compliant_changes']} compliant")
        row[4].text = ", ".join(detail_parts) if detail_parts else "—"


def _add_m1_detail_table(doc: Document, title: str, detail: list[dict]):
    """Add M1 per-change detail table."""
    if not detail:
        return
    _add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Entity", "Scope", "Description", "Matched"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for row_data in detail:
        row = table.add_row().cells
        row[0].text = row_data.get("actual_entity", "")
        row[1].text = row_data.get("scope", "")
        row[2].text = row_data.get("description", "")
        matched = row_data.get("matched", False)
        row[3].text = "Yes" if matched else "No"
        for p in row[3].paragraphs:
            for run in p.runs:
                run.font.color.rgb = _GREEN if matched else _RED


def _add_m5_detail_table(doc: Document, m5_data: dict):
    """Add M5 confidence calibration detail."""
    detail = m5_data.get("high_confidence_detail", [])
    if not detail:
        return
    _add_heading(doc, "M5 — High Confidence Calibration Detail", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Change", "Entity", "Scope", "Confidence", "Validated"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for row_data in detail:
        row = table.add_row().cells
        row[0].text = row_data.get("change", "")
        row[1].text = row_data.get("entity", "")
        row[2].text = row_data.get("scope", "")
        row[3].text = row_data.get("confidence", "")
        matched = row_data.get("matched", False)
        row[4].text = "Yes" if matched else "No"
        for p in row[4].paragraphs:
            for run in p.runs:
                run.font.color.rgb = _GREEN if matched else _RED


def _add_improvement_actions(doc: Document, actions: list[dict]):
    if not actions:
        return
    _add_heading(doc, "Improvement Actions", level=2)
    for a in actions:
        priority = a.get("priority", "")
        action = a.get("action", "")
        impact = a.get("impact", "")
        doc.add_paragraph(
            f"[{priority}] {action}",
            style="List Bullet",
        )
        if impact:
            doc.add_paragraph(f"Impact: {impact}", style="List Bullet 2")


def build_report(
    eval_result: dict,
    output_path: str | Path,
) -> Path:
    """Build a Word report from the structured eval JSON."""
    output_path = Path(output_path)
    study_id = eval_result.get("study_id", "Unknown")
    metrics = eval_result.get("metrics", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_heading(doc, f"D5 Eval Report — {study_id}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Eval Date: {eval_result.get('eval_date', 'N/A')}\n")
    meta.add_run(
        f"Generator Version: {eval_result.get('generator_version', 'N/A')}\n"
    )
    meta.add_run(
        f"Framework Version: {eval_result.get('eval_framework_version', '1.0')}\n"
    )
    meta.add_run(
        f"Ground Truth: {eval_result.get('ground_truth_source', 'N/A')}\n"
    )
    overall = eval_result.get("overall_status", "N/A")
    run = meta.add_run(f"Overall Status: {overall}")
    run.bold = True
    run.font.color.rgb = _status_color(overall)

    _add_heading(doc, "Metric Scorecard", level=1)
    _add_metric_table(doc, metrics)

    m1_global_detail = eval_result.get("m1_global_detail", [])
    m1_country_detail = eval_result.get("m1_country_detail", [])
    if m1_global_detail or m1_country_detail:
        _add_heading(doc, "M1 — Domain Recall Detail", level=1)
        _add_m1_detail_table(doc, "Global Changes", m1_global_detail)
        _add_m1_detail_table(doc, "Country-Specific Changes", m1_country_detail)

    m5 = metrics.get("m5", {})
    if m5.get("score") is not None:
        _add_heading(doc, "M5 — Confidence Calibration", level=1)
        _add_m5_detail_table(doc, m5)

    analysis = eval_result.get("analysis", {})
    gaps = analysis.get("coverage_gaps", [])
    if gaps:
        _add_heading(doc, "Coverage Gaps", level=1)
        for g in gaps:
            doc.add_paragraph(g, style="List Bullet")

    go_no_go = analysis.get("go_no_go", "")
    if go_no_go:
        _add_heading(doc, "Go / No-Go Assessment", level=1)
        p = doc.add_paragraph()
        run = p.add_run(go_no_go)
        run.bold = True

    actions = analysis.get("improvement_actions", [])
    _add_improvement_actions(doc, actions)

    doc.save(str(output_path))
    return output_path
