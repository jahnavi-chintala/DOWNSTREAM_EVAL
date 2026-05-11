import sys, io
from docx import Document

for fpath in sys.argv[1:]:
    doc = Document(fpath)
    out = fpath.replace(".docx", "_extracted.txt")
    with io.open(out, "w", encoding="utf-8") as f:
        f.write(f"=== {fpath} ===\n\n")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                f.write(f"P{i}: {p.text}\n")
        for i, t in enumerate(doc.tables):
            f.write(f"\n=== TABLE {i} ===\n")
            for ri, row in enumerate(t.rows):
                cells = [c.text.replace("\n", " | ") for c in row.cells]
                f.write(f"R{ri}: " + " || ".join(cells) + "\n")
    print(f"Extracted: {out}")
