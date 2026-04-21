"""
DMP eval Word report — layout aligned to ``reference_specs/DMP_Eval_Report_B7981027.docx``.

Structure: banner table → document score row (3 cols) → sections **1–7** as Heading 2,
same table shapes and column headers as the reference (only cell text varies by study).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:
    Document = None  # type: ignore

_FONT = "Arial"


def _apply_doc_defaults(doc: "Document") -> None:
    """Arial 10pt, narrow margins."""
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)


def _thin_table_borders(table) -> None:
    """Light-gray 0.5 pt borders on all sides."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblB.append(el)
    tblPr.append(tblB)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _banner_table(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    p = c.paragraphs[0]
    p.text = ""
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _document_score_row(
    doc: Document,
    *,
    ta_phase_slug: str,
    score: float,
    passed: bool,
    threshold: float,
    target: float,
) -> None:
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    left, mid, right = t.rows[0].cells
    p0 = left.paragraphs[0]
    p0.text = ""
    r0 = p0.add_run(f"DOCUMENT SCORE {ta_phase_slug}")
    r0.bold = True
    r0.font.size = Pt(10)
    pm = mid.paragraphs[0]
    pm.text = ""
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = "PASS ✓" if passed else "FAIL ✗"
    rm = pm.add_run(f"{score:.1f} / 100 {st}")
    rm.bold = True
    rm.font.size = Pt(11)
    pr = right.paragraphs[0]
    pr.text = ""
    rr = pr.add_run(f"Threshold: {threshold:.0f} | Target: {target:.0f}")
    rr.font.size = Pt(10)
    doc.add_paragraph()


def _add_heading2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _data_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    *,
    header_shade: str = "D9E2F3",
) -> None:
    if not headers:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    _thin_table_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.name = _FONT
        r.font.size = Pt(9)
        _set_cell_shading(c, header_shade)
    for i, row in enumerate(rows):
        for j in range(len(headers)):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            rn = cell.paragraphs[0].add_run(str(row[j]) if j < len(row) else "")
            rn.font.name = _FONT
            rn.font.size = Pt(9)
    doc.add_paragraph()


def write_dmp_eval_docx(output_path: Path, report: Dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError("python-docx required: pip install python-docx")

    meta = report.get("eval_metadata") or {}
    summ = report.get("summary_metrics") or {}
    sid = str(meta.get("study_id", "") or report.get("study_id", ""))

    doc = Document()
    doc.core_properties.title = f"DMP Eval Report — {sid}"
    _apply_doc_defaults(doc)

    ta = str(meta.get("therapeutic_area") or "—")
    ph = str(meta.get("phase") or "—")
    ta_slug = (meta.get("ta_slug") or meta.get("config_label") or "").strip()
    if ta_slug:
        ta_phase = f"{ta} / {ph} — {ta_slug}"
    else:
        ta_phase = f"{ta} / {ph}"

    _banner_table(doc, f"PFIZER PROTOCOL INTELLIGENCE PLATFORM DMP Generator — Eval Report {sid}")

    ds = float(report.get("document_score", 0) or 0)
    dp = bool(report.get("document_pass"))
    thr = float(report.get("document_pass_threshold", 75) or 75)
    tgt = float(report.get("document_target", 80) or 80)
    _document_score_row(
        doc,
        ta_phase_slug=ta_phase,
        score=ds,
        passed=dp,
        threshold=thr,
        target=tgt,
    )

    _add_heading2(doc, "1. Summary Metrics")
    _data_table(
        doc,
        ["Metric", "Score", "Target", "Pass/Fail", "Source"],
        [
            ["M1 S5 System Accuracy", f"{summ.get('m1_s5_system_accuracy', 0):.1%}", f"{summ.get('m1_target', 0):.0%}", "PASS ✓" if summ.get("m1_pass") else "FAIL ✗", "dmp_ground_truth S5"],
            ["M2 S6.2 Vendor Recall", f"{summ.get('m2_s6_vendor_recall', 0):.1%}", f"{summ.get('m2_target', 0):.0%}", "PASS ✓" if summ.get("m2_pass") else "FAIL ✗", "SDS Non-CRF CSV / JSON fallback"],
            ["M3 S8 Module Recall", f"{summ.get('m3_s8_module_recall', 0):.1%}", f"{summ.get('m3_target', 0):.0%}", "PASS ✓" if summ.get("m3_pass") else "FAIL ✗", "dmp_ground_truth S8"],
            ["M4 S11.4 Flags", f"{summ.get('m4_reconciliation_accuracy', 0):.1%}", f"{summ.get('m4_target', 0):.0%}", "PASS ✓" if summ.get("m4_pass") else "FAIL ✗", "dmp_ground_truth S11"],
            ["M4b Hallucinations (null source_tag)", str(summ.get("m4_hallucinations", 0)), str(summ.get("m4_hallucination_target", 0)), "PASS ✓" if summ.get("m4_hallucination_pass") else "FAIL ✗", "S5/S6/S8 traceability"],
        ],
    )

    _add_heading2(doc, "2. Section Scorecard")
    sc = report.get("section_scores") or {}
    sr: List[List[str]] = []
    order = ["s5_systems", "s8_critical_data", "s6_vendors", "s11_reconciliation"]
    labels = {
        "s5_systems": "S5 Systems",
        "s8_critical_data": "S8 Critical Data",
        "s6_vendors": "S6.2 Vendors",
        "s11_reconciliation": "S11.4 Reconciliation",
    }
    metrics_lbl = {
        "s5_systems": "M1",
        "s8_critical_data": "M3",
        "s6_vendors": "M2",
        "s11_reconciliation": "M4",
    }
    for key in order:
        blk = sc.get(key) or {}
        st = "PASS ✓" if (
            (key == "s5_systems" and summ.get("m1_pass"))
            or (key == "s6_vendors" and summ.get("m2_pass"))
            or (key == "s8_critical_data" and summ.get("m3_pass"))
            or (key == "s11_reconciliation" and summ.get("m4_pass"))
        ) else "FAIL ✗"
        sr.append(
            [
                labels.get(key, key),
                metrics_lbl.get(key, "—"),
                f"{float(blk.get('weight', 0)):.2f}",
                str(blk.get("score", "")),
                str(blk.get("weighted_contribution", "")),
                st,
            ],
        )
    _data_table(doc, ["Section", "Metric", "Weight", "Score", "Weighted", "Status"], sr)

    _add_heading2(doc, "3. S5 Systems & Tools Detail")
    s5 = report.get("s5_systems") or []
    s5_rows: List[List[str]] = []
    for r in s5:
        if not isinstance(r, dict):
            continue
        ms = r.get("match_status", "")
        icon = "✓" if ms == "verbatim" else ("~" if ms == "near_miss" else "✗")
        tier = str(r.get("attributes", {}).get("s5_confidence_tier", {}).get("generated", "") or "")
        note = str(r.get("ground_truth_name", ""))[:200]
        s5_rows.append(
            [
                str(r.get("system_type", "")),
                str(r.get("generated_name") or "—"),
                str(r.get("item_score", "")),
                f"{icon} {ms}",
                tier,
                note,
            ]
        )
    _data_table(doc, ["System Type", "Generated", "Score", "Match", "Tier", "Notes"], s5_rows[:80])

    _add_heading2(doc, "4. S8 Critical Data Detail")
    s8 = report.get("s8_critical_data") or []
    s8_rows: List[List[str]] = []
    for r in s8:
        if not isinstance(r, dict):
            continue
        ms = r.get("match_status", "")
        icon = "✓" if ms == "verbatim" else ("~" if ms == "near_miss" else "✗")
        if ms == "extra":
            icon = "✗"
        atr = r.get("attributes") or {}
        tier = str((atr.get("s8_tier") or {}).get("ground_truth", "") or "")[:80]
        module_cell = str(r.get("generated_module") or "—")
        layer_cell = str(r.get("layer") or "—")
        notes = str(r.get("note", ""))[:220]
        if r.get("ground_truth_module"):
            gt_hint = str(r.get("ground_truth_module"))
            if notes:
                notes = f"GT: {gt_hint} | {notes}"
            else:
                notes = f"GT: {gt_hint}"
        s8_rows.append(
            [
                module_cell,
                layer_cell,
                str(r.get("item_score", "")),
                f"{icon} {ms}",
                tier,
                notes,
            ]
        )
    _data_table(doc, ["Module", "Layer", "Score", "Match", "Tier", "Notes"], s8_rows[:80])

    _add_heading2(doc, "5. S6.2 eSource/eData Vendor Table Detail")
    s6 = report.get("s6_vendors") or []
    s6_rows: List[List[str]] = []
    for r in s6:
        if not isinstance(r, dict) or r.get("item_score") == "excluded":
            continue
        atr = r.get("attributes") or {}
        dt = str(atr.get("s6_data_type", {}).get("ground_truth", "") or "")[:80]
        tier = str(atr.get("s6_data_review_tier", {}).get("ground_truth", "") or "")
        s6_rows.append(
            [
                str(r.get("ground_truth_vendor") or "—"),
                str(r.get("generated_vendor") or "—"),
                dt,
                tier,
                str(r.get("item_score", "")),
                str(r.get("match_status", "")),
            ]
        )
    _data_table(
        doc,
        ["Ground Truth Vendor", "Generated Vendor", "Data Type", "Tier", "Score", "Notes"],
        s6_rows[:80],
    )

    _add_heading2(doc, "6. S11.4 Reconciliation Flags Detail")
    s11 = report.get("s11_reconciliation") or []
    s11_rows: List[List[str]] = []
    for r in s11:
        if not isinstance(r, dict):
            continue
        atr = r.get("attributes") or {}
        s11_rows.append(
            [
                str(r.get("subsection_label", "")),
                str(r.get("ground_truth_flag", "")),
                str(r.get("generated_flag", "")),
                str(r.get("item_score", "")),
                str((atr.get("s11_inference_rule") or {}).get("rule") or ""),
            ]
        )
    _data_table(doc, ["Subsection", "Ground Truth", "Generated", "Score", "Inference Rule"], s11_rows)

    _add_heading2(doc, "7. Improvement Actions")
    ia = report.get("improvement_actions") or []
    ia_rows: List[List[str]] = []
    for x in ia:
        if not isinstance(x, dict):
            continue
        ia_rows.append(
            [
                str(x.get("priority", "")),
                str(x.get("section", "")),
                str(x.get("type", "")),
                str(x.get("action", ""))[:400],
                str(x.get("fix_location", ""))[:240],
            ]
        )
    _data_table(doc, ["Priority", "Section", "Type", "Action", "Fix Location"], ia_rows)

    em = report.get("eval_metadata") or {}
    oos = em.get("out_of_scope_sections_excluded") or []
    oos_note = (
        "Out of scope — not penalised: "
        + (", ".join(oos) if oos else "S1 (Revision History), S3 (DM Process Documentation)")
        + "."
    )
    fp = doc.add_paragraph()
    r = fp.add_run(oos_note)
    r.font.size = Pt(9)
    r.italic = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
    except PermissionError:
        alt = output_path.parent / f"{output_path.stem}_{datetime.now():%Y%m%d_%H%M%S}{output_path.suffix}"
        doc.save(str(alt))
