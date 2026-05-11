"""
Rewrite an eval-report Markdown body to mirror the structure of
optional reference Word outline via ``PIPD_EVAL_REFERENCE_DOCX`` using OpenAI.

Metrics and counts are taken from the rule-based Markdown you pass in; the model
only aligns headings, section flow, and table shaping to the reference outline.
Requires OPENAI_API_KEY. Optional: PIPD_EVAL_REFERENCE_DOCX, OPENAI_REF_LAYOUT_MAX_TOKENS.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Optional, Tuple

_PKG_DIR = Path(__file__).resolve().parent


def extract_eval_docx_outline(docx_path: Path, *, max_chars: int = 14_000) -> str:
    """Flatten reference Word into paragraph + table row text for use as a structure guide."""
    from docx import Document

    d = Document(str(docx_path))
    parts: list[str] = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for ti, table in enumerate(d.tables):
        parts.append(f"\n--- reference_table_{ti + 1} ---")
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[outline truncated]"
    return text


def resolve_reference_eval_docx() -> Path:
    env = (os.environ.get("PIPD_EVAL_REFERENCE_DOCX") or "").strip()
    if env:
        p = Path(env).expanduser()
        return p if p.is_absolute() else _PKG_DIR / p
    return Path()


def rewrite_eval_markdown_like_reference(
    markdown_body: str,
    *,
    reference_docx_path: Optional[Path] = None,
) -> Tuple[str, Optional[str]]:
    """
    Call OpenAI to reformat ``markdown_body`` like the reference eval Word outline.

    Returns (markdown, error_message). On any failure, callers should fall back
    to the original markdown.
    """
    from reports.pipd_composite_report import _openai_chat

    ref = reference_docx_path or resolve_reference_eval_docx()
    if not ref.is_file():
        return markdown_body, f"Reference .docx not found: {ref}"

    outline = extract_eval_docx_outline(ref)
    md_in = markdown_body
    if len(md_in) > 110_000:
        md_in = md_in[:110_000] + "\n\n… [source markdown truncated for model context]"

    system = (
        "You reformat PIPD evaluation reports as Markdown only.\n"
        "You are given (1) a structural outline extracted from a reference Word report "
        "(headings and table rows — use it for section order and table *shape*/column labels) "
        "and (2) a source Markdown report whose numbers are authoritative.\n\n"
        "Output a single Markdown document that follows the reference flow: org banner, "
        "DOCUMENT SCORE area, then sections such as Summary Metrics (pipe table with "
        "Metric / Score / Target / Pass/Fail / Source), Category Scorecard, Category Detail, "
        "and any improvement / appendix blocks present in the source.\n\n"
        "Hard rules:\n"
        "- Copy every percentage, count, study ID, date fragment, threshold, and PASS/FAIL line "
        "exactly from the source Markdown. Do not recalculate, re-round, or substitute numbers.\n"
        "- Keep metric and category scorecard table rows value-identical to the source. "
        "Only formatting/layout changes are allowed.\n"
        "- Do not remove tables or rows from the source; you may reorder sections only if the "
        "reference outline clearly uses a different order — if unsure, keep the source order.\n"
        "- Use standard pipe tables (| col |). No HTML. No wrapping the full answer in ``` fences.\n"
        "- Do not invent metrics or categories not in the source.\n"
        "- If the source includes a section titled `## Protocol & ground-truth sourced fields (AI, excerpt-based)` "
        "or similar, preserve its substantive content (you may adjust heading level to fit the reference outline)."
    )
    user = (
        "### Reference outline (structure only; numbers here may be stale)\n\n"
        + outline
        + "\n\n### Source Markdown (preserve all data exactly)\n\n"
        + md_in
    )
    max_tok = int(os.environ.get("OPENAI_REF_LAYOUT_MAX_TOKENS", "12000"))
    text, err = _openai_chat(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        max_tokens=max_tok,
    )
    if err:
        return markdown_body, err
    if not text or not text.strip():
        return markdown_body, "Empty OpenAI response"
    out = text.strip()
    if out.startswith("```"):
        out = out.removeprefix("```markdown").removeprefix("```md").removeprefix("```").strip()
        if out.endswith("```"):
            out = out[: -3].strip()
    ok, why = _preserves_eval_rows(markdown_body, out)
    if not ok:
        return markdown_body, f"OpenAI layout rewrite rejected: {why}"
    return out, None


def _norm_md_row(row: str) -> str:
    return re.sub(r"\s+", " ", row.strip())


def _extract_guard_rows(md: str) -> List[str]:
    """Rows whose numeric values must never change in layout rewriting."""
    rows: List[str] = []
    for ln in md.splitlines():
        s = ln.strip()
        if not (s.startswith("|") and s.endswith("|")):
            continue
        # Guard the metric rows (M1..M6) and category scorecard rows (Cat n.)
        if re.search(r"\|\s*M[1-6]\b", s) or re.search(r"\|\s*Cat\s+\d+\.", s):
            rows.append(_norm_md_row(s))
    return rows


def _preserves_eval_rows(source_md: str, rewritten_md: str) -> Tuple[bool, Optional[str]]:
    guarded = _extract_guard_rows(source_md)
    if not guarded:
        return True, None
    rewritten_norm = {_norm_md_row(ln) for ln in rewritten_md.splitlines()}
    missing = [r for r in guarded if r not in rewritten_norm]
    if missing:
        return False, f"missing or changed guarded rows, first: {missing[0]}"
    return True, None
