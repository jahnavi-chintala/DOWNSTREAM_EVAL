#!/usr/bin/env python3
"""
Generate PIPD Generator Eval Report (.docx) from eval_data.yml or .json (python-docx).
YAML needs: pip install pyyaml

  python generate_pipd_report.py [input.yml|json] [output.docx] [--write-json]

No args: first existing file among eval_data.yml / .yaml / eval_data.json.
``--write-json`` writes ``<input_stem>.json`` pretty-printed from the loaded data
(useful to keep YAML and JSON in sync).
Default output: eval_outputs/PIPD_Eval_Report_{study_id}.docx

**Scenario 1 layout** follows the same section order as the eval pipeline
(``pipd_eval_report_reference.build_reference_eval_markdown``), visually aligned to
``reference_spec/PIPD_Eval_Report_B7981027.docx``. The older informal name
``reference report.docx`` is not read by code — reports are generated Markdown→Word, so any
hand-maintained .docx can drift. Category weights and eval rules for S1 are in
``reference_spec/pipd_eval_config.yaml`` (see also ``reference.yaml`` — different schema from report data).

All substantive text and numbers come from the report input file. Optional root key
``template`` supplies configurable labels; optional ``footer.end_line`` sets the closing
sentence verbatim.

  template:
    none_identified_note: "..."   # default note when a category has none_identified
    near_miss_heading: "Near-miss"
    misses_heading: "Misses (ground truth not generated)"
    references_heading: "References / sources"
    section_metrics: "1. Summary Metrics"      # override any section title if needed
    ...

Per category you may set ``scorecard_status`` and ``scorecard_status_level`` (ok | warn |
error) instead of auto-derived status text. ``none_identified_note`` on a category
overrides the template default for that category.

Category score “bars” are rendered as one embedded matplotlib chart (not a table of
block characters). Requires matplotlib for the figure; without it, scores are listed
as plain text.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.eval_scenario1 import gsop_set_from_value

C_H2 = RGBColor(0x1F, 0x38, 0x64)
C_H3 = RGBColor(0x2E, 0x75, 0xB6)
C_HEADER_BG = "BDD7EE"
C_PASS = RGBColor(0x37, 0x56, 0x23)
C_FAIL = RGBColor(0xC0, 0x00, 0x00)
C_HIGH = RGBColor(0xC0, 0x00, 0x00)
C_MED = RGBColor(0xC5, 0x5A, 0x11)
C_LOW = RGBColor(0x59, 0x59, 0x59)
C_NOTE_BG = "FFFFC0"
C_OOS_BG = "F2F2F2"
C_BORDER = "BFBFBF"
C_STATUS_OK = RGBColor(0x37, 0x56, 0x23)
C_STATUS_WARN = RGBColor(0xC5, 0x5A, 0x11)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table, color: str = C_BORDER, sz: int = 8) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _rgb_run(paragraph, text: str, *, bold: bool = False, size_pt: Optional[float] = None,
             color: Optional[RGBColor] = None, italic: bool = False) -> None:
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    if color is not None:
        r.font.color.rgb = color
    r.font.name = "Calibri"


def _h(doc: Document, text: str, level: int, size_pt: float, color: RGBColor) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size_pt)
        run.font.bold = True
        run.font.color.rgb = color


def _section_margins_letter(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
        sec.page_height = Inches(11)
        sec.page_width = Inches(8.5)


def _template_str(data: Dict[str, Any], key: str, default: str) -> str:
    t = data.get("template")
    if isinstance(t, dict):
        v = t.get(key)
        if isinstance(v, str) and v.strip():
            return v
    return default


def _section_title(data: Dict[str, Any], key: str, default: str) -> str:
    return _template_str(data, key, default)


def _cat_sort_key(cat: Dict[str, Any]) -> int:
    try:
        return int(cat.get("id"))
    except (TypeError, ValueError):
        return 0


def _severity_rgb(level: Optional[str]) -> RGBColor:
    u = (level or "warn").strip().lower()
    if u in ("ok", "pass", "good"):
        return C_STATUS_OK
    if u in ("error", "fail", "bad"):
        return C_FAIL
    return C_STATUS_WARN


def _row_gsop_display(rr: Dict[str, Any]) -> str:
    """Prefer ``gsop_codes`` (eval JSON) then legacy ``gsop`` string."""
    raw = rr.get("gsop_codes")
    if raw is None:
        raw = rr.get("gsop")
    codes = gsop_set_from_value(raw)
    return ", ".join(sorted(codes)) if codes else "—"


def _map_usdm_conf(v: Any) -> str:
    s = str(v or "").strip().lower()
    if s in ("resolved", "ok", "yes", "true", "\u2713", "pass"):
        return "\u2713"
    if s in ("warn", "review", "pending", "~", "low_confidence", "partial"):
        return "~"
    return "\u2717"


def _skip_pure_informational_zero(cat: Dict[str, Any]) -> bool:
    if float(cat.get("score") or 0) != 0:
        return False
    if cat.get("misses") or cat.get("near_misses") or cat.get("none_identified"):
        return False
    st = int(cat.get("subcats_total") or 0)
    return st == 0 and not cat.get("rows")


def _include_scorecard(cat: Dict[str, Any]) -> bool:
    if _skip_pure_informational_zero(cat):
        return False
    if cat.get("misses") or cat.get("near_misses") or cat.get("none_identified"):
        return True
    if int(cat.get("subcats_total") or 0) > 0:
        return True
    return float(cat.get("score") or 0) > 0


def _include_detail(cat: Dict[str, Any]) -> bool:
    return bool(cat.get("rows") or cat.get("misses") or cat.get("near_misses")
                or cat.get("none_identified"))


def _status_line(data: Dict[str, Any], cat: Dict[str, Any]) -> Tuple[str, RGBColor]:
    custom = cat.get("scorecard_status")
    if isinstance(custom, str) and custom.strip():
        return custom.strip(), _severity_rgb(cat.get("scorecard_status_level"))

    sm = int(cat.get("subcats_matched") or 0)
    st = int(cat.get("subcats_total") or 0)
    see = _template_str(data, "see_detail_pointer", "").strip()

    if cat.get("none_identified") and st == 0:
        note = cat.get("none_identified_summary")
        if isinstance(note, str) and note.strip():
            return note.strip(), _severity_rgb(cat.get("scorecard_status_level") or "ok")
        return f"none_identified; subcats {sm}/{st}", _severity_rgb("ok")

    if cat.get("misses"):
        msg = f"{len(cat['misses'])} miss(es)"
        if see:
            msg = f"{msg} ({see})"
        return msg, _severity_rgb("warn")

    if cat.get("near_misses"):
        return f"{len(cat['near_misses'])} near-miss(es)", _severity_rgb("warn")

    if st > 0 and sm >= st:
        return f"{sm}/{st} matched", _severity_rgb("ok")
    if st > 0:
        tail = f" ({see})" if see else ""
        return f"{sm}/{st}{tail}", _severity_rgb("warn")

    tail = f" ({see})" if see else ""
    return f"subcats {sm}/{st}{tail}", _severity_rgb("warn")


def _add_shaded_note_table(
    doc: Document,
    text: str,
    *,
    fill: str = C_NOTE_BG,
    border_sz: int = 6,
    size_pt: float = 10,
) -> None:
    t = doc.add_table(rows=1, cols=1)
    _set_table_borders(t, sz=border_sz)
    cell = t.rows[0].cells[0]
    _set_cell_shading(cell, fill)
    lines = (text or "").strip().split("\n")
    if not lines or not lines[0]:
        _rgb_run(cell.paragraphs[0], "—", size_pt=size_pt)
    else:
        _rgb_run(cell.paragraphs[0], lines[0], size_pt=size_pt)
        for line in lines[1:]:
            _rgb_run(cell.add_paragraph(), line, size_pt=size_pt)
    doc.add_paragraph()


def _near_miss_cell_text(data: Dict[str, Any], nm: Dict[str, Any]) -> str:
    prefix = _template_str(data, "near_miss_line_prefix", "NEAR MISS (semantic):")
    g = str(nm.get("generated") or "").strip()
    gt = str(nm.get("gt") or "").strip()
    cos = nm.get("cosine")
    body = f'{prefix} Generated: "{g}"  vs  GT: "{gt}"'
    if cos is not None:
        body += f"  [cosine: {cos}]"
    return body


def _miss_cell_text(data: Dict[str, Any], m: Any) -> str:
    prefix = _template_str(data, "miss_line_prefix", "MISS:")
    if isinstance(m, dict):
        mt = str(m.get("text") or m.get("issue") or "").strip()
        rs = str(m.get("reason") or "").strip()
    else:
        mt, rs = str(m).strip(), ""
    if rs:
        return f"{prefix} {mt}  → {rs}"
    return f"{prefix} {mt}"


def _out_of_scope_cell_text(data: Dict[str, Any], oos: Dict[str, Any]) -> str:
    custom = oos.get("report_text")
    if isinstance(custom, str) and custom.strip():
        return custom.strip()
    intro = _template_str(data, "out_of_scope_intro", "Out of scope — not penalised")
    cnt = oos.get("count")
    reason = str(oos.get("reason") or "").strip()
    detail = str(oos.get("detail") or "").strip()
    mid = ""
    if cnt is not None and reason:
        mid = f"{cnt} rows excluded ({reason})"
    elif cnt is not None:
        mid = f"{cnt} rows excluded"
    elif reason:
        mid = reason
    if detail:
        mid = f"{mid}. {detail}" if mid else detail
    return f"{intro}\n{mid}" if mid else intro


def _footer_end_line(data: Dict[str, Any]) -> Optional[str]:
    foot = data.get("footer") or {}
    el = foot.get("end_line")
    if isinstance(el, str) and el.strip():
        return el.strip()
    sj = foot.get("source_json")
    cfg = foot.get("config")
    if sj or cfg:
        rest = []
        if sj:
            rest.append(f"Source JSON: {sj}")
        if cfg:
            rest.append(f"Config: {cfg}")
        return "End of eval report. " + "  |  ".join(rest)
    return None


def _add_category_scores_chart(doc: Document, cats_sc: List[Dict[str, Any]], study_id: str) -> None:
    if not cats_sc:
        return
    labels = [f"Cat {c.get('id')}" for c in cats_sc]
    scores = [max(0.0, min(100.0, float(c.get("score") or 0))) for c in cats_sc]
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        doc.add_paragraph("Category scores (matplotlib not installed; chart skipped):")
        for c, sc in zip(cats_sc, scores):
            p = doc.add_paragraph()
            _rgb_run(p, f"Cat {c.get('id')} — {float(sc):.1f}", size_pt=10)
        return

    h = max(2.2, 0.38 * len(cats_sc) + 1.1)
    fig, ax = plt.subplots(figsize=(6.5, min(h, 8.0)))
    ax.barh(
        labels[::-1],
        scores[::-1],
        color="#0073C8",
        height=0.62,
        edgecolor="white",
        linewidth=0.5,
    )
    ax.set_facecolor("white")
    ax.grid(axis="x", linestyle=":", alpha=0.6)
    ax.set_xlabel("Score (0–100)")
    ax.set_xlim(0, 100)
    ax.set_title("")
    fig.tight_layout()
    fd, tmp_path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    fig.savefig(tmp_path, format="png", dpi=120)
    plt.close(fig)
    doc.add_picture(tmp_path, width=Inches(6.3))
    doc.add_paragraph()
    try:
        os.unlink(tmp_path)
    except OSError:
        pass


def _priority_color(p: str) -> RGBColor:
    u = (p or "").upper()
    if u == "HIGH":
        return C_HIGH
    if u == "MEDIUM":
        return C_MED
    return C_LOW


def build_report(data: Dict[str, Any], output_path: Path) -> None:
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10)
    _section_margins_letter(doc)

    h = data.get("header") or {}
    study_id = str(h.get("study_id") or "UNKNOWN")
    study_name = str(h.get("study_name") or "")
    line3 = f"{study_id} — {study_name}" if study_name else study_id

    t0 = doc.add_table(rows=1, cols=1)
    _set_table_borders(t0, sz=8)
    c0 = t0.rows[0].cells[0]
    c0.paragraphs[0].paragraph_format.space_after = Pt(2)
    _rgb_run(c0.paragraphs[0], "PFIZER PROTOCOL INTELLIGENCE PLATFORM", bold=True, size_pt=11)
    _rgb_run(c0.add_paragraph(), "PIPD Generator — Eval Report", bold=True, size_pt=11)
    _rgb_run(c0.add_paragraph(), line3, size_pt=10)
    meta = (
        f"Eval date: {h.get('eval_date', '—')} | Generator: {h.get('generator_version', '—')} | "
        f"Config: {h.get('config', '—')}"
    )
    _rgb_run(c0.add_paragraph(), meta, size_pt=9)
    doc.add_paragraph()

    ds = data.get("document_score") or {}
    study_type = str(h.get("study_type") or "")
    score_val = ds.get("score")
    score_str = f"{score_val} / 100" if isinstance(score_val, (int, float)) else str(score_val)
    passed = bool(ds.get("pass", False))
    pass_txt = "PASS \u2713" if passed else "FAIL \u2717"

    t1 = doc.add_table(rows=1, cols=3)
    _set_table_borders(t1, sz=4)
    r1 = t1.rows[0].cells
    _rgb_run(r1[0].paragraphs[0], "DOCUMENT SCORE", bold=True, size_pt=11)
    _rgb_run(r1[0].add_paragraph(), study_type, bold=True, size_pt=10)
    mid = r1[1].paragraphs[0]
    mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rgb_run(mid, score_str, bold=True, size_pt=14)
    pm = r1[1].add_paragraph()
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rgb_run(pm, f"Threshold: {ds.get('threshold', '—')} / Target: {ds.get('target', '—')}", size_pt=8)
    right = r1[2].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rgb_run(right, pass_txt, bold=True, size_pt=12, color=C_PASS if passed else C_FAIL)
    doc.add_paragraph()

    _h(doc, _section_title(data, "section_metrics", "1. Summary Metrics"), 2, 12, C_H2)
    metrics: List[Dict[str, Any]] = list(data.get("metrics") or [])
    tm = doc.add_table(rows=1 + len(metrics), cols=5)
    _set_table_borders(tm, sz=4)
    for ci, htxt in enumerate(["Metric", "Score", "Target", "Pass/Fail", "Source"]):
        cell = tm.rows[0].cells[ci]
        _set_cell_shading(cell, C_HEADER_BG)
        _rgb_run(cell.paragraphs[0], htxt, bold=True, size_pt=10)

    for ri, m in enumerate(metrics, start=1):
        ok = bool(m.get("pass", False))
        row = tm.rows[ri].cells
        name = str(m.get("name") or m.get("id") or "—")
        sc = str(m.get("score") or "—")
        tg = str(m.get("target") or "—")
        pf = "PASS \u2713" if ok else "FAIL \u2717"
        b = ok
        _rgb_run(row[0].paragraphs[0], name, bold=b, size_pt=10)
        _rgb_run(row[1].paragraphs[0], sc, bold=b, size_pt=10)
        _rgb_run(row[2].paragraphs[0], tg, size_pt=10)
        _rgb_run(row[3].paragraphs[0], pf, bold=True, size_pt=10, color=C_PASS if ok else C_FAIL)
        _rgb_run(row[4].paragraphs[0], str(m.get("source") or "—"), size_pt=10)
    doc.add_paragraph()

    _h(doc, _section_title(data, "section_scorecard", "2. Category Scorecard"), 2, 12, C_H2)
    cats_all: List[Dict[str, Any]] = list(data.get("categories") or [])
    cats_sc = sorted([c for c in cats_all if _include_scorecard(c)], key=_cat_sort_key)
    if not cats_sc:
        doc.add_paragraph("(No categories selected for scorecard.)")
    else:
        ts = doc.add_table(rows=1 + len(cats_sc), cols=6)
        _set_table_borders(ts, sz=4)
        for ci, htxt in enumerate(["Category", "Weight", "Score", "Weighted", "Subcats", "Status"]):
            cell = ts.rows[0].cells[ci]
            _set_cell_shading(cell, C_HEADER_BG)
            _rgb_run(cell.paragraphs[0], htxt, bold=True, size_pt=10)

        for ri, cat in enumerate(cats_sc, start=1):
            row = ts.rows[ri].cells
            cid, cname = cat.get("id", ""), str(cat.get("name") or "")
            w = float(cat.get("weight") or 0) * 100
            sc = float(cat.get("score") or 0)
            wt = cat.get("weighted")
            wt_s = f"{float(wt):.2f}" if isinstance(wt, (int, float)) else str(wt)
            sm = int(cat.get("subcats_matched") or 0)
            st = int(cat.get("subcats_total") or 0)
            sub_custom = cat.get("subcats_display")
            if isinstance(sub_custom, str) and sub_custom.strip():
                sub_txt = sub_custom.strip()
            elif st == 0 and cat.get("none_identified"):
                sub_txt = _template_str(data, "empty_subcats_label", "Empty ✓")
            else:
                sub_txt = f"{sm}/{st}"
            stxt, srgb = _status_line(data, cat)
            _rgb_run(row[0].paragraphs[0], f"Cat {cid}. {cname}", size_pt=10)
            _rgb_run(row[1].paragraphs[0], f"{w:.0f}%", size_pt=10)
            _rgb_run(row[2].paragraphs[0], f"{sc:.1f}", size_pt=10)
            _rgb_run(row[3].paragraphs[0], wt_s, size_pt=10)
            _rgb_run(row[4].paragraphs[0], sub_txt, size_pt=10)
            _rgb_run(row[5].paragraphs[0], stxt, bold=True, size_pt=10, color=srgb)

        cap = _template_str(
            data,
            "scorecard_chart_caption",
            "Category scores (weighted contribution)",
        )
        _rgb_run(doc.add_paragraph(), cap, size_pt=10)
        if data.get("embed_category_score_chart", True) is not False:
            _add_category_scores_chart(doc, cats_sc, study_id)

    _h(doc, _section_title(data, "section_detail", "3. Category Detail"), 2, 12, C_H2)
    for cat in sorted([c for c in cats_all if _include_detail(c)], key=_cat_sort_key):
        cid, cname = cat.get("id"), str(cat.get("name") or "")
        _h(doc, f"Category {cid} — {cname}", 3, 11, C_H3)
        st = int(cat.get("subcats_total") or 0)
        if cat.get("none_identified") and st == 0 and not cat.get("rows"):
            nn = cat.get("none_identified_note")
            if not (isinstance(nn, str) and nn.strip()):
                nn = _template_str(
                    data,
                    "none_identified_note",
                    "No subcategories in scope for this category (none_identified).",
                )
            _add_shaded_note_table(doc, nn.strip())
            continue

        rows_data: List[Dict[str, Any]] = list(cat.get("rows") or [])
        if rows_data:
            td = doc.add_table(rows=1 + len(rows_data), cols=7)
            _set_table_borders(td, sz=4)
            dh = ["Subcategory text", "Score", "Match", "YES/NO", "USDM", "Conf.", "GSOP"]
            for ci, htxt in enumerate(dh):
                cell = td.rows[0].cells[ci]
                _set_cell_shading(cell, C_HEADER_BG)
                _rgb_run(cell.paragraphs[0], htxt, bold=True, size_pt=10)

            for ri, rr in enumerate(rows_data, start=1):
                row = td.rows[ri].cells
                is_extra = bool(rr.get("hallucination") or rr.get("extra"))
                if is_extra:
                    sym = "—"
                elif bool(rr.get("yes_no")):
                    sym = "\u2713"
                else:
                    sym = "\u2717"
                txt = str(rr.get("text") or "")
                _rgb_run(row[0].paragraphs[0], txt, size_pt=9)
                _rgb_run(row[1].paragraphs[0], str(rr.get("score", "—")), bold=True, size_pt=10)
                _rgb_run(row[2].paragraphs[0], str(rr.get("match") or "—"), size_pt=10)
                _rgb_run(row[3].paragraphs[0], sym, size_pt=10)
                _rgb_run(row[4].paragraphs[0], _map_usdm_conf(rr.get("usdm")), size_pt=10)
                _rgb_run(row[5].paragraphs[0], _map_usdm_conf(rr.get("conf")), size_pt=10)
                _rgb_run(row[6].paragraphs[0], _row_gsop_display(rr), size_pt=10)
            doc.add_paragraph()

        for nm in cat.get("near_misses") or []:
            if not isinstance(nm, dict):
                continue
            _add_shaded_note_table(doc, _near_miss_cell_text(data, nm))

        misses = cat.get("misses") or []
        for m in misses:
            _add_shaded_note_table(doc, _miss_cell_text(data, m))

    _h(doc, _section_title(data, "section_actions", "4. Improvement Actions"), 2, 12, C_H2)
    intro4 = _template_str(
        data,
        "section_actions_intro",
        "Actions derived from eval scoring. Prioritised for generator developer.",
    )
    _rgb_run(doc.add_paragraph(), intro4, size_pt=10)
    actions: List[Dict[str, Any]] = list(data.get("improvement_actions") or [])
    if not actions:
        doc.add_paragraph("(No improvement actions listed.)")
    else:
        ta = doc.add_table(rows=1 + len(actions), cols=5)
        _set_table_borders(ta, sz=4)
        for ci, htxt in enumerate(["Priority", "Category", "Type", "Action", "Fix location"]):
            cell = ta.rows[0].cells[ci]
            _set_cell_shading(cell, C_HEADER_BG)
            _rgb_run(cell.paragraphs[0], htxt, bold=True, size_pt=9)
        for ri, a in enumerate(actions, start=1):
            row = ta.rows[ri].cells
            pr = str(a.get("priority") or "—")
            _rgb_run(row[0].paragraphs[0], pr, bold=True, size_pt=9, color=_priority_color(pr))
            _rgb_run(row[1].paragraphs[0], str(a.get("category") or "—"), size_pt=9)
            _rgb_run(row[2].paragraphs[0], str(a.get("type") or "—"), size_pt=9)
            _rgb_run(row[3].paragraphs[0], str(a.get("action") or "—"), size_pt=9)
            _rgb_run(row[4].paragraphs[0], str(a.get("fix_location") or "—"), size_pt=9)
        doc.add_paragraph()

    oos = data.get("out_of_scope")
    if isinstance(oos, dict) and (
        oos.get("count") is not None or oos.get("reason") or oos.get("report_text")
    ):
        _add_shaded_note_table(
            doc,
            _out_of_scope_cell_text(data, oos),
            fill=C_OOS_BG,
            border_sz=6,
            size_pt=10,
        )

    end_ln = _footer_end_line(data)
    if end_ln:
        _rgb_run(doc.add_paragraph(), end_ln, size_pt=9)

    doc.save(output_path)


def load_eval_data(path: Path) -> Dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    if path.suffix.lower() in (".yml", ".yaml"):
        try:
            import yaml  # type: ignore
        except ImportError as e:
            raise SystemExit("YAML input requires: pip install pyyaml") from e
        return yaml.safe_load(text) or {}
    return json.loads(text)


def main(argv: List[str]) -> int:
    root = Path(__file__).resolve().parent
    candidates = [root / "eval_data.yml", root / "eval_data.yaml", root / "eval_data.json"]
    in_path: Optional[Path] = None
    out_path: Optional[Path] = None
    write_json_sidecar = False
    pos: List[str] = []
    i = 1
    while i < len(argv):
        a = argv[i]
        if a == "--write-json":
            write_json_sidecar = True
            i += 1
            continue
        if not a.startswith("-"):
            pos.append(a)
        i += 1
    if len(pos) >= 1:
        in_path = Path(pos[0]).expanduser().resolve()
    if len(pos) >= 2:
        out_path = Path(pos[1]).expanduser().resolve()
    if in_path is None:
        for p in candidates:
            if p.is_file():
                in_path = p
                break
        if in_path is None:
            print("No input file. Provide eval_data.yml or place it in the script directory.", file=sys.stderr)
            return 1
    data = load_eval_data(in_path)
    if write_json_sidecar:
        js = in_path.with_suffix(".json")
        js.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        print(f"Wrote {js}")
    if out_path is None:
        sid = str((data.get("header") or {}).get("study_id") or "UNKNOWN")
        out_dir = root / "eval_outputs"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"PIPD_Eval_Report_{sid}.docx"
    build_report(data, out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
