"""
CMP Eval Report Word layout — stakeholder spec (blocks 1–7).

Imported by ``cmp_report_docx``; keeps the main module smaller.
"""

from __future__ import annotations

from datetime import date
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

# ── Theme — Pfizer reference palette ─────────────────────────────────────────
HEADER_BG = "003087"       # Pfizer blue — banner and table header backgrounds
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
YELLOW_NOTE = "FFF2CC"
BLUE_NOTE = "D9E2F3"
GREY_OOS = "D9D9D9"
GREEN_BANNER = "C6EFCE"
GREEN_OK = RGBColor(0x1A, 0x73, 0x40)   # pass text
RED_HIGH = RGBColor(0xB9, 0x1C, 0x1C)   # fail text

_BODY = 10.5
_TITLE = 15.0
_SECTION = 12.0
_SCORE_BIG = 22.0


def set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def clear_cell(cell) -> Any:
    for child in list(cell._tc):
        if child.tag.endswith("}tcPr"):
            continue
        cell._tc.remove(child)
    return cell.add_paragraph()


def run_text(
    p,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = _BODY,
    rgb: Optional[RGBColor] = None,
) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    if rgb is not None:
        r.font.color.rgb = rgb


def add_dark_header_table(
    doc: Document,
    line1: str,
    line2: str,
    line3: str,
    line4: str,
) -> None:
    """Block 1 — single cell, dark background, white text."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, HEADER_BG)
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(p, line1 + "\n", bold=True, size_pt=_TITLE, rgb=WHITE)
    run_text(p, line2 + "\n", bold=True, size_pt=_TITLE - 1, rgb=WHITE)
    run_text(p, line3 + "\n", bold=False, size_pt=_BODY, rgb=WHITE)
    run_text(p, line4, bold=False, size_pt=_BODY - 0.5, rgb=WHITE)
    doc.add_paragraph()


def add_document_score_banner(
    doc: Document,
    *,
    ta_phase_bucket: str,
    score: float,
    passed: bool,
    threshold: float,
    target: float,
) -> None:
    """Block 2 — 3 columns × 1 row."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    left, mid, right = t.rows[0].cells
    widths = [Inches(2.35), Inches(1.85), Inches(2.3)]
    for c, w in zip((left, mid, right), widths):
        c.width = w

    p = clear_cell(left)
    run_text(p, "DOCUMENT SCORE", bold=True, size_pt=_SECTION)
    p = left.add_paragraph()
    run_text(p, ta_phase_bucket, bold=True, size_pt=_BODY)

    pm = clear_cell(mid)
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(pm, f"{score:.1f}", bold=True, size_pt=_SCORE_BIG)
    pm2 = mid.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(pm2, "/ 100", bold=True, size_pt=_SECTION)

    pr = clear_cell(right)
    st = "PASS ✓" if passed else "FAIL ✗"
    run_text(pr, st, bold=True, size_pt=_SECTION, rgb=GREEN_OK if passed else RED_HIGH)
    pr2 = right.add_paragraph()
    run_text(pr2, f"Threshold: {threshold:.0f}", bold=True, size_pt=_BODY)
    pr3 = right.add_paragraph()
    run_text(pr3, f"Target: {target:.0f}", bold=True, size_pt=_BODY)
    doc.add_paragraph()


def text_bar_for_score(score: float, *, width_chars: int = 40) -> str:
    """~1 █ per 2.5 pts on 0–100 scale."""
    s = max(0.0, min(100.0, float(score)))
    n = int(round(s / 2.5))
    n = min(n, width_chars)
    bar = "█" * n + "░" * (width_chars - n)
    return f"{bar} {s:.1f}"


# Section “weighted contribution” horizontal bars (stakeholder figure)
_BAR_FILL_GOOD = "C6EFCE"  # light green — score ≥ 90
_BAR_FILL_WARN = "FFF2CC"  # light yellow — score < 90
_BAR_TRACK = "F2F2F2"  # remainder of 0–100 track


def _strip_cell_body_keep_tcpr(cell) -> None:
    """Remove paragraphs/tables from a cell; keep ``tcPr`` (e.g. width, shading)."""
    tc = cell._tc
    for child in list(tc):
        if child.tag.endswith("}tcPr"):
            continue
        tc.remove(child)


def _section_bar_score_display(score: float) -> str:
    """92 / 78.5 / 96.5 style (no trailing .0)."""
    r = max(0.0, min(100.0, float(score)))
    if abs(r - round(r)) < 0.001:
        return str(int(round(r)))
    s = f"{r:.1f}"
    if s.endswith(".0"):
        return s[:-2]
    return s


def add_section_weighted_contribution_bar_table(doc: Document, srows: Sequence[Sequence[str]]) -> None:
    """
    One bordered table: **Label (weight%)** | proportional bar | **score**.

    Bar fill: green when score ≥ 90, else yellow (matches stakeholder mock-up).
    """
    if not srows:
        return
    n = len(srows)
    t = doc.add_table(rows=n, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    bar_total = Inches(3.72)
    min_seg = Inches(0.04)

    for ri, row in enumerate(srows):
        parts = list(row)
        name = str(parts[0]) if parts else "—"
        wp = str(parts[1]) if len(parts) > 1 else ""
        label = f"{name} ({wp})" if wp else name
        try:
            sc = float(parts[2]) if len(parts) > 2 else 0.0
        except (TypeError, ValueError):
            sc = 0.0
        fill = _BAR_FILL_GOOD if sc >= 90.0 else _BAR_FILL_WARN

        c0, c1, c2 = t.rows[ri].cells
        c0.width = Inches(2.12)
        c1.width = bar_total
        c2.width = Inches(0.68)

        p0 = clear_cell(c0)
        run_text(p0, label, bold=False, size_pt=_BODY)

        _strip_cell_body_keep_tcpr(c1)
        sc_clamped = max(0.0, min(100.0, sc))
        if sc_clamped <= 0.0:
            inner = c1.add_table(rows=1, cols=1)
            inner.style = "Table Grid"
            inner.autofit = False
            inner.rows[0].cells[0].width = bar_total
            set_cell_shading(inner.rows[0].cells[0], _BAR_TRACK)
            clear_cell(inner.rows[0].cells[0])
        elif sc_clamped >= 100.0:
            inner = c1.add_table(rows=1, cols=1)
            inner.style = "Table Grid"
            inner.autofit = False
            inner.rows[0].cells[0].width = bar_total
            set_cell_shading(inner.rows[0].cells[0], fill)
            clear_cell(inner.rows[0].cells[0])
        else:
            inner = c1.add_table(rows=1, cols=2)
            inner.style = "Table Grid"
            inner.autofit = False
            frac = sc_clamped / 100.0
            w_fill = bar_total * frac
            w_rest = bar_total - w_fill
            if w_fill < min_seg:
                w_fill = min_seg
                w_rest = bar_total - w_fill
            if w_rest < min_seg:
                w_rest = min_seg
                w_fill = bar_total - w_rest
            inner.rows[0].cells[0].width = w_fill
            inner.rows[0].cells[1].width = w_rest
            set_cell_shading(inner.rows[0].cells[0], fill)
            set_cell_shading(inner.rows[0].cells[1], _BAR_TRACK)
            clear_cell(inner.rows[0].cells[0])
            clear_cell(inner.rows[0].cells[1])

        p2 = clear_cell(c2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_text(p2, _section_bar_score_display(sc), bold=True, size_pt=_BODY)
    doc.add_paragraph()


def add_heading_numbered(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)
    doc.add_paragraph()


def add_two_col_bar_row(doc: Document, label: str, score: float, *, label_bold: bool = True) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.rows[0].cells[0].width = Inches(2.4)
    t.rows[0].cells[1].width = Inches(4.1)
    c0 = clear_cell(t.rows[0].cells[0])
    run_text(c0, label, bold=label_bold, size_pt=_BODY)
    c1 = clear_cell(t.rows[0].cells[1])
    run_text(c1, text_bar_for_score(score), bold=True, size_pt=_BODY)
    doc.add_paragraph()


def add_shaded_callout(doc: Document, text: str, fill: str) -> None:
    tb = doc.add_table(rows=1, cols=1)
    tb.style = "Table Grid"
    cell = tb.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = clear_cell(cell)
    run_text(p, text, bold=False, size_pt=_BODY)
    doc.add_paragraph()


def _fmt_thr_pct(attr: Optional[dict]) -> str:
    if not isinstance(attr, dict):
        return "—"
    g, h = attr.get("generated"), attr.get("ground_truth")
    if g is not None and h is not None:
        return f"{g} / {h}"
    if g is not None:
        return str(g)
    if h is not None:
        return str(h)
    return "—"


def _thr_pct_display(attr: Optional[dict]) -> str:
    if not isinstance(attr, dict):
        return "—"
    dp = attr.get("delta_pct")
    if dp is not None:
        try:
            return f"{float(dp):.1f}%"
        except (TypeError, ValueError):
            return str(dp)
    return _fmt_thr_pct(attr)


def _tier_cell(attrs: dict) -> str:
    """KRI Tier: 1 = study-standard (appears in study GT), 2 = TA/phase-inferred."""
    tier = attrs.get("confidence_tier") if isinstance(attrs, dict) else None
    if not isinstance(tier, dict):
        # generator may store tier as a plain int/str at a higher level
        return "—"
    g = tier.get("generated")
    if g is not None:
        label = "Study-standard" if str(g).strip() == "1" else f"TA-inferred (T{g})"
        return label
    return "—"


def _match_miss_cell(attr: dict, *, na_label: str = "N/A") -> str:
    """Convert a scored attribute dict to a simple Match ✓ / Miss ✗ / N/A label."""
    if not isinstance(attr, dict):
        return "—"
    score = attr.get("score")
    if score is None:
        return "—"
    score = float(score)
    gen = attr.get("generated")
    gt  = attr.get("ground_truth")
    gt_blank = gt in (None, "", "nan", "NaN") or (isinstance(gt, float) and str(gt) == "nan")
    gen_blank = gen in (None, "", "nan", "NaN") or (isinstance(gen, float) and str(gen) == "nan")
    if gt_blank and gen_blank:
        return na_label
    if score >= 1.0:
        return "Match \u2713"
    return "Miss \u2717"


def _iqmp_status_cell(iq: dict) -> str:
    """Show Match ✓ / Miss ✗ / Extra ! / N/A for IQMP risk ID."""
    if not isinstance(iq, dict):
        return "—"
    score = iq.get("score")
    is_miss = bool(iq.get("is_missing"))
    is_extra = bool(iq.get("is_hallucination"))
    gt = iq.get("ground_truth")
    gen = iq.get("generated")
    gt_blank = gt in (None, "", "nan", "NaN") or (isinstance(gt, str) and gt.lower() == "nan")
    gen_blank = gen in (None, "", "nan", "NaN") or (isinstance(gen, str) and gen.lower() == "nan")
    if gt_blank and gen_blank:
        return "N/A"
    if is_extra:
        return "Extra !"
    if is_miss or (score is not None and float(score) == 0.0 and not gen_blank):
        return "Miss \u2717"
    if score is not None and float(score) == 0.0:
        return "Miss \u2717"
    return "Match \u2713"


def kri_rows_eight_column(items: List[dict]) -> List[List[str]]:
    """
    KRI Label | Score | Match | Mod Threshold | High Threshold | IQMP ID | KRI Tier | Weight
    Mod/High threshold and Weight show Match/Miss status, not raw values.
    IQMP ID shows Match/Miss/Extra/N/A status.
    KRI Tier: 1=Study-standard, 2=TA-inferred (generator-assigned).
    """
    rows: List[List[str]] = []
    for item in items:
        ms = (item.get("match_status") or "").lower()
        is_miss = ms == "miss"
        gen = item.get("generated_label")
        glab = item.get("ground_truth_label") or "—"
        label = str(glab if is_miss and not gen else (gen if gen is not None else glab)) or "—"
        if item.get("classification_error"):
            label = "\u26a0 " + label + " \u2014 classification error"
        attrs = item.get("attributes") or {}
        if not isinstance(attrs, dict):
            attrs = {}
        if is_miss:
            sc_miss = float(item.get("kri_score", 0) or 0)
            rows.append([label, f"{sc_miss:.0f}", "MISS", "—", "—", "—", "—", "—"])
            continue
        mod  = attrs.get("moderate_threshold")
        high = attrs.get("high_threshold")
        iq   = attrs.get("iqmp_risk_id")
        wgt  = attrs.get("weight_field")
        lbl  = attrs.get("kri_label") or {}
        mt   = str((lbl.get("match_type") if isinstance(lbl, dict) else "") or "").lower()
        ms2  = (item.get("match_status") or "").lower()
        if ms2 == "verbatim" or mt == "verbatim":
            match_disp = "Verbatim"
        elif "semantic" in ms2 or "semantic" in mt:
            match_disp = "Semantic ~"
        else:
            match_disp = (item.get("match_status") or "—").title()
        sc = float(item.get("kri_score", 0) or 0)
        rows.append(
            [
                label,
                f"{sc:.0f}",
                match_disp,
                _match_miss_cell(mod if isinstance(mod, dict) else {}, na_label="N/A (relative)"),
                _match_miss_cell(high if isinstance(high, dict) else {}, na_label="N/A (relative)"),
                _iqmp_status_cell(iq if isinstance(iq, dict) else {}),
                _tier_cell(attrs),
                _match_miss_cell(wgt if isinstance(wgt, dict) else {}),
            ]
        )
    return rows


def _qtl_expect_tol_cells(q: dict) -> Tuple[str, str]:
    attrs = q.get("attributes") or {}
    ep = attrs.get("expectation_pct") if isinstance(attrs.get("expectation_pct"), dict) else None
    tp = attrs.get("tolerance_limit_pct") if isinstance(attrs.get("tolerance_limit_pct"), dict) else None
    if isinstance(ep, dict) and (ep.get("generated") is not None or ep.get("ground_truth") is not None):
        ex = _fmt_thr_pct(ep)
    else:
        ex = _pct(q.get("expectation_score"))
    if isinstance(tp, dict) and (tp.get("generated") is not None or tp.get("ground_truth") is not None):
        tol = _fmt_thr_pct(tp)
    else:
        tol = _pct(q.get("tolerance_score"))
    return ex, tol


def _qtl_sister_cell(q: dict) -> str:
    attrs = q.get("attributes") or {}
    sk = attrs.get("sister_kri") if isinstance(attrs.get("sister_kri"), dict) else None
    if isinstance(sk, dict) and sk.get("generated"):
        return str(sk.get("generated"))
    return "—"


def qtl_is_near_miss(q: dict) -> bool:
    nm = q.get("near_miss") if isinstance(q.get("near_miss"), dict) else {}
    if nm.get("cosine_similarity") is not None:
        return True
    ms = str(q.get("match_status") or "").lower()
    if ms in ("semantic", "near_match", "near_miss"):
        return True
    return False


def qtl_name_cell(q: dict) -> str:
    gn = str(q.get("generated_name") or "—")
    gtn = str(q.get("ground_truth_name") or "")
    nm = q.get("near_miss") if isinstance(q.get("near_miss"), dict) else {}
    if qtl_is_near_miss(q) and gtn:
        inner = nm.get("generated") if nm.get("generated") is not None else gn
        return f"{gtn} (generated: '{inner}')"
    return gn


def qtl_near_miss_box_text(q: dict) -> str:
    nm = q.get("near_miss") if isinstance(q.get("near_miss"), dict) else {}
    g = nm.get("generated") or q.get("generated_name") or "—"
    t = nm.get("ground_truth") or q.get("ground_truth_name") or "—"
    cos = nm.get("cosine_similarity")
    cos_s = f"{float(cos):.2f}" if cos is not None else "—"
    fix = nm.get("action") or "Align QTL naming to ground truth."
    fl = nm.get("fix_file") or "cmp_generator"
    loc = nm.get("fix_location") or "QTL mapping"
    return (
        f"NEAR MISS (semantic): Generated: '{g}' vs GT: '{t}' [cosine: {cos_s}]  "
        f"Fix: {fix}. {fl} → {loc}."
    )


def qtl_row_cells(q: dict) -> List[str]:
    """One data row: QTL Name | Score | Match | Expect% | Tolerance% | Sister KRI"""
    gn = qtl_name_cell(q)
    ms = str(q.get("match_status") or "").lower()
    if ms in ("semantic", "near_match", "near_miss"):
        match_disp = "Semantic ~"
    elif ms in ("verbatim", "matched"):
        match_disp = "Verbatim"
    else:
        match_disp = str(q.get("match_status") or "—").title()
    ex, tol = _qtl_expect_tol_cells(q)
    return [
        gn,
        f"{float(q.get('qtl_score', 0)):.0f}",
        match_disp,
        ex,
        tol,
        _qtl_sister_cell(q),
    ]


def add_qtl_table_with_near_miss_notes(doc: Document, items: List[dict]) -> None:
    """Single 6-column grid; near-miss rows followed by merged blue note row."""
    headers = ["QTL Name", "Score", "Match", "Expect%", "Tolerance%", "Sister KRI"]
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    w = 6.5 / 6
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        p = clear_cell(cell)
        run_text(p, h, bold=True, size_pt=_BODY)
        set_cell_shading(cell, "D9E2F3")
        cell.width = Inches(w)
    for q in items:
        row = t.add_row()
        cells = qtl_row_cells(q)
        for ci, val in enumerate(cells):
            c = row.cells[ci]
            p = clear_cell(c)
            bold = ci == 1 or (ci == 2 and str(val) == "Semantic ~")
            run_text(p, str(val), bold=bold, size_pt=_BODY)
            c.width = Inches(w)
        if qtl_is_near_miss(q):
            note_row = t.add_row()
            a = note_row.cells[0]
            b = note_row.cells[5]
            a.merge(b)
            set_cell_shading(a, BLUE_NOTE)
            p = clear_cell(a)
            run_text(p, qtl_near_miss_box_text(q), size_pt=_BODY - 0.5)
    doc.add_paragraph()


def study_specific_miss_fix_note(item: dict) -> str:
    lab = str(item.get("ground_truth_label") or item.get("generated_label") or "—")
    usdm = str(item.get("usdm_location") or "USDM — see protocol JSON")
    fix_f = str(item.get("fix_file") or "cmp_generator")
    fix_l = str(item.get("fix_location") or "study_specific_kris")
    return f"MISS: {lab} → KRI not generated. {usdm}. Fix: {fix_f} → {fix_l}."


def add_kri_eight_table_with_callouts(
    doc: Document,
    items: List[dict],
    *,
    yellow_miss_notes: bool = False,
) -> None:
    """8-column KRI grid; optional merged yellow row after each MISS (study-specific)."""
    headers = ["KRI Label", "Score", "Match", "Mod Threshold", "High Threshold", "IQMP ID", "KRI Tier", "Weight"]
    rows_data = kri_rows_eight_column(items)
    t = doc.add_table(rows=1, cols=8)
    t.style = "Table Grid"
    w = 6.5 / 8
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        p = clear_cell(cell)
        run_text(p, h, bold=True, size_pt=_BODY)
        set_cell_shading(cell, "D9E2F3")
        cell.width = Inches(w)
    for i, row_cells in enumerate(rows_data):
        item = items[i]
        row = t.add_row()
        for ci, val in enumerate(row_cells):
            c = row.cells[ci]
            p = clear_cell(c)
            bold = ci == 1 or (ci == 2 and str(val) == "MISS")
            run_text(p, str(val), bold=bold, size_pt=_BODY)
            c.width = Inches(w)
        if yellow_miss_notes and (item.get("match_status") or "").lower() == "miss":
            note_row = t.add_row()
            a = note_row.cells[0]
            b = note_row.cells[7]
            a.merge(b)
            set_cell_shading(a, YELLOW_NOTE)
            p = clear_cell(a)
            run_text(p, study_specific_miss_fix_note(item), size_pt=_BODY - 0.5)
    doc.add_paragraph()


def add_section_scorecard_table(doc: Document, srows: List[List[str]]) -> None:
    """Section | Weight | Score | Weighted | KRIs/QTLs | Status — bold section names & status."""
    headers = ["Section", "Weight", "Score", "Weighted", "KRIs/QTLs", "Status"]
    t = doc.add_table(rows=1 + len(srows), cols=6)
    t.style = "Table Grid"
    w = 6.5 / 6
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        p = clear_cell(cell)
        run_text(p, h, bold=True, size_pt=_BODY, rgb=WHITE)
        set_cell_shading(cell, HEADER_BG)
        cell.width = Inches(w)
    for ri, row in enumerate(srows):
        for ci in range(6):
            val = row[ci] if ci < len(row) else ""
            c = t.rows[ri + 1].cells[ci]
            p = clear_cell(c)
            bold = ci == 0 or ci == 2 or ci == 5
            rgb = None
            if ci == 5:
                if "PASS" in str(val):
                    rgb = GREEN_OK
                elif "FAIL" in str(val):
                    rgb = RED_HIGH
            run_text(p, str(val), bold=bold, size_pt=_BODY, rgb=rgb)
            c.width = Inches(w)
    doc.add_paragraph()


def _pct(v: Any) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
        if 0 <= f <= 1.0:
            return f"{f * 100:.0f}%"
        return f"{f:.0f}%"
    except (TypeError, ValueError):
        return str(v)


def add_metrics_table_formatted(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> Table:
    """Metric names, scores, Pass/Fail column bold; PASS ✓ / FAIL ✗."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = False
    w = 6.5 / max(ncols, 1)
    widths = [Inches(w)] * ncols
    hdr = table.rows[0].cells
    for ci, h in enumerate(headers):
        p = clear_cell(hdr[ci])
        run_text(p, str(h), bold=True, size_pt=_BODY, rgb=WHITE)
        set_cell_shading(hdr[ci], HEADER_BG)
        hdr[ci].width = widths[ci]
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci in range(ncols):
            val = row_data[ci] if ci < len(row_data) else ""
            cell = row.cells[ci]
            p = clear_cell(cell)
            bold = ci in (0, 1) or (ci == 3 and ("PASS" in str(val) or "FAIL" in str(val)))
            rgb = None
            if ci == 3:
                if "PASS" in str(val):
                    rgb = GREEN_OK
                elif "FAIL" in str(val):
                    rgb = RED_HIGH
            run_text(p, str(val), bold=bold, size_pt=_BODY, rgb=rgb)
            cell.width = widths[ci]
    doc.add_paragraph()
    return table


def improvement_rows_sorted(actions: List[dict]) -> List[List[str]]:
    order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}

    def key(a: dict) -> Tuple[int, str]:
        p = str(a.get("priority", "")).upper()
        return (order.get(p, 9), str(a.get("action", "")))

    rows = []
    for a in sorted(actions, key=key):
        pr = str(a.get("priority", ""))
        rows.append(
            [
                pr,
                str(a.get("section", "")),
                str(a.get("type", "")),
                str(a.get("action", "") or ""),
                str(a.get("fix_location", "") or ""),
            ]
        )
    return rows


def add_improvement_table(doc: Document, rows: List[List[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.style = "Table Grid"
    hdr = ["Priority", "Section", "Type", "Action", "Fix Location"]
    for ci, h in enumerate(hdr):
        p = clear_cell(tbl.rows[0].cells[ci])
        run_text(p, h, bold=True, size_pt=_BODY, rgb=WHITE)
        set_cell_shading(tbl.rows[0].cells[ci], HEADER_BG)
    for ri, r in enumerate(rows):
        for ci in range(5):
            val = r[ci] if ci < len(r) else ""
            cell = tbl.rows[ri + 1].cells[ci]
            p = clear_cell(cell)
            is_high = ci == 0 and str(val).upper() == "HIGH"
            rgb = RED_HIGH if is_high else None
            run_text(p, str(val), bold=is_high, size_pt=_BODY, rgb=rgb)
    doc.add_paragraph()


def footer_cmp(doc: Document, study: str, cfg_name: str, *, version: str = "v1.0") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_text(
        p,
        f"End of eval report. Source JSON: cmp_eval_{study}.json | Config: {cfg_name}",
        bold=False,
        size_pt=9,
    )
    for r in p.runs:
        r.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mo = date.today().strftime("%B %Y")
    run_text(
        p2,
        f"D3 CMP Generator — Eval Report {version} | {study} | Pfizer Protocol Intelligence Platform | {mo}",
        size_pt=9,
    )
    for r in p2.runs:
        r.italic = True
