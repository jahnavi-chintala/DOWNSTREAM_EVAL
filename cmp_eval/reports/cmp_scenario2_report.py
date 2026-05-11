"""
cmp_scenario2_report.py
-----------------------
CMP Scenario 2 (no ground truth) — Word document report.

Takes the JSON dict produced by ``cmp_eval/eval_scenario2.py::run_scenario2_eval``
and writes a styled .docx file matching the Pfizer brand palette used by the
CMP Scenario 1 report.

Public API
~~~~~~~~~~
    write_cmp_scenario2_docx(s2_result, output_path)
    write_cmp_scenario2_docx_from_json(json_path, output_path)
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Pfizer brand palette ────────────────────────────────────────────────────
_HEADER_BG  = "003087"
_GREEN_BG   = "C6EFCE"
_AMBER_BG   = "FFF2CC"
_RED_BG     = "FFDCE0"
_BLUE_BG    = "D9E2F3"
_GREY_BG    = "F2F2F2"
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_FG   = RGBColor(0x1A, 0x73, 0x40)
_AMBER_FG   = RGBColor(0x92, 0x60, 0x00)
_RED_FG     = RGBColor(0xB9, 0x1C, 0x1C)
_BLUE_FG    = RGBColor(0x00, 0x30, 0x87)
_DARK_FG    = RGBColor(0x26, 0x26, 0x26)
_FONT       = "Arial"
_BODY_PT    = 10.0
_CELL_PT    = 9.0       # table cell text — matches S1 eval
_HEAD_PT    = 11.0
_TITLE_PT   = 13.0
_BIG_PT     = 22.0      # score % in verdict box — matches S1 eval


# ── XML helpers ──────────────────────────────────────────────────────────────

def _shd(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _color(run, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement("w:color")
    el.set(qn("w:val"), hex_color)
    rPr.append(el)


def _borders(table) -> None:
    """Light-gray single-pt borders for all sides."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _run(para, text: str, *, bold: bool = False, italic: bool = False,
         size: float = _BODY_PT, rgb: Optional[RGBColor] = None,
         hex_color: Optional[str] = None) -> None:
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = _FONT
    r.font.size = Pt(size)
    if rgb is not None:
        r.font.color.rgb = rgb
    if hex_color is not None:
        _color(r, hex_color)


def _clear(cell):
    for child in list(cell._tc):
        if child.tag.endswith("}tcPr"):
            continue
        cell._tc.remove(child)
    return cell.add_paragraph()


# ── Compound builders ─────────────────────────────────────────────────────────

def _apply_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)


def _section_heading(doc: Document, text: str) -> None:
    """Full-width blue banner — matches S1 eval report section heading style."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shd(cell, _HEADER_BG)
    p = _clear(cell)
    r = p.add_run(text)
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(_HEAD_PT)
    r.font.color.rgb = _WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_banner(doc: Document, study_id: str, title: str, eval_date: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shd(cell, _HEADER_BG)
    p = _clear(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CMP Evaluation Report — Scenario 2\n", bold=True, size=_TITLE_PT,
         rgb=_WHITE)
    _run(p, f"Study: {study_id}  ·  {title}\n", bold=True, size=_HEAD_PT,
         rgb=_WHITE)
    _run(p, "No Ground Truth · Proxy Quality Signals  ·  CONFIDENTIAL",
         size=_BODY_PT, rgb=_WHITE)
    doc.add_paragraph()


def _add_verdict_box(doc: Document, verdict: str, health: Dict[str, Any]) -> None:
    verdict_upper = (verdict or "").upper()
    go_no_go = "GO" if verdict_upper in ("GREEN", "AMBER") else "NO-GO"
    pct = round(health.get("percent", 0.0), 1)
    pass_n = health.get("pass", 0)
    warn_n = health.get("warn", 0)
    fail_n = health.get("fail", 0)
    total  = health.get("total", 0)

    if verdict_upper == "GREEN":
        bg, fg_hex = _GREEN_BG, "1A7340"
    elif verdict_upper == "AMBER":
        bg, fg_hex = _AMBER_BG, "926000"
    else:
        bg, fg_hex = _RED_BG, "B91C1C"

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    widths = [Inches(2.3), Inches(1.8), Inches(2.4)]
    for ci, w in enumerate(widths):
        t.rows[0].cells[ci].width = w

    # Left: label
    c0 = t.rows[0].cells[0]
    _shd(c0, _BLUE_BG)
    p0 = _clear(c0)
    _run(p0, "SCENARIO 2 SCORE\n", bold=True, size=_HEAD_PT, hex_color=_HEADER_BG)
    p0b = c0.add_paragraph()
    _run(p0b, "CMP Quality Signals", italic=True, size=_BODY_PT, hex_color=_HEADER_BG)

    # Mid: score %
    c1 = t.rows[0].cells[1]
    _shd(c1, bg)
    p1 = _clear(c1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, f"{pct}%", bold=True, size=_BIG_PT, hex_color=fg_hex)
    p1b = c1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1b, "Signal Health", size=_BODY_PT - 1, hex_color="555555")

    # Right: verdict
    c2 = t.rows[0].cells[2]
    p2 = _clear(c2)
    _run(p2, f"{verdict_upper}  ·  {go_no_go}\n", bold=True, size=_HEAD_PT,
         hex_color=fg_hex)
    p2b = c2.add_paragraph()
    _run(p2b,
         f"{pass_n} PASS · {warn_n} WARN · {fail_n} FAIL  (of {total})\n",
         size=_BODY_PT - 0.5, hex_color="555555")
    p2c = c2.add_paragraph()
    _run(p2c, "Scoring: PASS=1.0 · WARN=0.5 · FAIL=0.0", italic=True,
         size=_BODY_PT - 1, hex_color="555555")
    doc.add_paragraph()


def _status_colors(status: str) -> Tuple[str, str]:
    """Returns (bg_hex, fg_hex) for a signal status cell."""
    s = (status or "").upper()
    if s == "PASS":
        return _GREEN_BG, "1A7340"
    if s == "WARN":
        return _AMBER_BG, "926000"
    if s == "FAIL":
        return _RED_BG, "B91C1C"
    return _GREY_BG, "262626"


_CMP_SIG_DESC = {
    "S1": "Structure Validity — CMP JSON satisfies schema / field structure",
    "S2": "KRI & QTL Count — global KRIs, study KRIs, QTLs within expected range",
    "S3": "Threshold Completeness — moderate / high thresholds populated",
    "S4": "Protocol Provenance — each KRI has iqmp_risk_id traceability",
    "S5": "Confidence Distribution — < 40% low/review confidence items",
    "S6": "KRI Label Uniqueness — no duplicate KRI labels",
    "S7": "Analysis Frequency Completeness — analysis_frequency populated",
    "S8": "USDM Traceability — key USDM anchor classes present in protocol",
}


def _add_score_breakdown(doc: Document, signals: Dict[str, Any], health: Dict[str, Any]) -> None:
    """Insert a 'How the Score Was Calculated' table after the verdict box."""
    _section_heading(doc, "Score Breakdown — How the Score Was Calculated")
    total = health.get("total", len(signals))
    p = doc.add_paragraph()
    _run(p, f"Formula: total points ÷ {total} checks × 100   |   PASS = 1.0 pt  ·  WARN = 0.5 pt  ·  FAIL = 0.0 pt",
         italic=True, size=_CELL_PT, hex_color="555555")

    weights = {"PASS": 1.0, "WARN": 0.5, "FAIL": 0.0}
    headers = ["#", "Quality Check", "What It Checks", "Result", "Points"]
    col_widths = [Inches(0.3), Inches(1.4), Inches(2.6), Inches(0.8), Inches(0.9)]

    sig_keys = list(signals.keys())
    t = doc.add_table(rows=1 + len(sig_keys) + 1, cols=5)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_BODY_PT - 0.5, rgb=_WHITE)

    running = 0.0
    for ri, sk in enumerate(sig_keys, 1):
        sv = signals.get(sk) or {}
        st = str(sv.get("status") or "").upper()
        pts = weights.get(st, 0.0)
        running += pts
        bg, fg = _status_colors(st)
        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w
        _clear(row.cells[0]).add_run(str(ri)).font.size = Pt(_BODY_PT - 0.5)
        p_name = _clear(row.cells[1])
        _run(p_name, sk, bold=True, size=_BODY_PT - 0.5)
        p_desc = _clear(row.cells[2])
        _run(p_desc, _CMP_SIG_DESC.get(sk, str(sv.get("description") or "")), size=_BODY_PT - 1)
        c_st = row.cells[3]
        _shd(c_st, bg)
        p_st = _clear(c_st)
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_st, st, bold=True, size=_BODY_PT - 0.5, hex_color=fg)
        p_pts = _clear(row.cells[4])
        p_pts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_pts, f"{pts:.1f} / 1.0", size=_BODY_PT - 0.5)

    # Total row
    total_row = t.rows[-1]
    for ci, w in enumerate(col_widths):
        total_row.cells[ci].width = w
    _shd(total_row.cells[0], _BLUE_BG)
    _shd(total_row.cells[1], _BLUE_BG)
    _shd(total_row.cells[2], _BLUE_BG)
    _shd(total_row.cells[3], _BLUE_BG)
    _shd(total_row.cells[4], _BLUE_BG)
    p_tot = _clear(total_row.cells[1])
    _run(p_tot, f"TOTAL  ·  {running:.1f} ÷ {total} × 100", bold=True, size=_BODY_PT - 0.5, hex_color=_HEADER_BG)
    p_res = _clear(total_row.cells[4])
    p_res.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_res, f"= {round(running / total * 100, 1)}%", bold=True, size=_BODY_PT, hex_color=_HEADER_BG)
    doc.add_paragraph()


def _add_kri_history(doc: Document, s2: Dict[str, Any]) -> None:
    """Show each generated KRI vs. historical benchmark presence."""
    import os
    import pandas as pd

    _section_heading(doc, "KRI Historical Evidence")
    _note_para = doc.add_paragraph()
    _run(_note_para,
         "Each KRI checked against benchmark data from historical Pfizer studies.  "
         "Novel KRIs (never seen before) are flagged for review.",
         italic=True, size=_CELL_PT, hex_color="555555")

    # Try to load KRI ground truth CSV
    _report_dir = Path(__file__).resolve().parent.parent
    gt_candidates = [
        str(_report_dir / "data" / "cmp_kri_ground_truth.csv"),
        r"C:\Users\jahna\OneDrive\Desktop\Pfizer\cmp_eval\data\cmp_kri_ground_truth.csv",
    ]
    df = pd.DataFrame()
    for path in gt_candidates:
        if os.path.exists(path):
            try:
                df = pd.read_csv(path, dtype=str).fillna("")
                break
            except Exception:
                pass

    # Collect all generated KRIs
    gen_kris: List[Dict[str, Any]] = []
    for bucket in ("global_kris", "study_specific_kris", "kris"):
        for item in (s2.get(bucket) or []):
            if isinstance(item, dict):
                gen_kris.append({"bucket": bucket.replace("_kris", "").replace("kris", "global"), "item": item})

    # Fallback to signals data if no embedded KRIs
    if not gen_kris:
        s2_counts = (s2.get("signals") or {}).get("S2") or {}
        counts = s2_counts.get("counts") or {}
        doc.add_paragraph(
            f"KRI detail not embedded in Scenario 2 JSON. "
            f"Global KRIs: {counts.get('global_kris', '—')}  "
            f"Study KRIs: {counts.get('study_specific_kris', '—')}  "
            f"QTLs: {counts.get('qtls', '—')}"
        ).italic = True
        return

    headers = ["Bucket", "KRI Label", "Confidence", "Studies with this KRI", "Evidence Status"]
    col_widths = [Inches(0.8), Inches(2.8), Inches(0.9), Inches(1.2), Inches(1.3)]

    total_kris = len(gen_kris)
    t = doc.add_table(rows=1 + total_kris, cols=5)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_BODY_PT - 0.5, rgb=_WHITE)

    novel_count = 0
    for ri, entry in enumerate(gen_kris, 1):
        item = entry["item"]
        bucket = entry["bucket"]
        label = str(item.get("kri_label") or "—")
        conf = str(item.get("confidence") or item.get("tier") or "—")
        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

        # Historical lookup
        hist_count = 0
        if not df.empty and "kri_label" in df.columns:
            matches = df[df["kri_label"].str.strip().str.lower() == label.strip().lower()]
            hist_count = matches["study_id"].nunique() if "study_id" in matches.columns else len(matches)

        if hist_count > 0:
            ev = "Confirmed"
            ev_bg, ev_fg = _GREEN_BG, "1A7340"
        else:
            ev = "Novel — Review"
            ev_bg, ev_fg = _AMBER_BG, "926000"
            novel_count += 1

        _clear(row.cells[0]).add_run(bucket).font.size = Pt(_BODY_PT - 0.5)
        p_lbl = _clear(row.cells[1])
        _run(p_lbl, label[:100], size=_BODY_PT - 0.5)
        _clear(row.cells[2]).add_run(conf).font.size = Pt(_BODY_PT - 0.5)
        hist_txt = f"{hist_count} / {df['study_id'].nunique() if not df.empty and 'study_id' in df.columns else '?'} studies"
        p_hist = _clear(row.cells[3])
        _run(p_hist, hist_txt, size=_BODY_PT - 0.5)
        c_ev = row.cells[4]
        _shd(c_ev, ev_bg)
        p_ev = _clear(c_ev)
        _run(p_ev, ev, bold=True, size=_BODY_PT - 0.5, hex_color=ev_fg)

    doc.add_paragraph()
    p_sum = doc.add_paragraph()
    _run(p_sum,
         f"Summary: {total_kris - novel_count} / {total_kris} KRIs confirmed in historical data · "
         f"{novel_count} novel (review recommended)",
         bold=True, size=_BODY_PT - 0.5, hex_color=_HEADER_BG)
    doc.add_paragraph()


def _note_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT - 0.5)
    _color(r, "555555")
    p.paragraph_format.space_after = Pt(4)


def _add_scorecard(doc: Document, signals: Dict[str, Any]) -> None:
    _section_heading(doc, "1. Signal Scorecard")
    _signal_order = ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8"]
    headers = ["Signal", "Description", "Status", "Key Detail"]
    col_widths = [Inches(1.4), Inches(2.6), Inches(0.85), Inches(1.65)]

    t = doc.add_table(rows=1 + len(_signal_order), cols=4)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    # Header row
    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_BODY_PT, rgb=_WHITE)

    # Data rows
    for ri, sig_id in enumerate(_signal_order, 1):
        blk = signals.get(sig_id) or {}
        status = str(blk.get("status") or "—").upper()
        name = str(blk.get("name") or sig_id)
        desc = str(blk.get("description") or "")
        detail = _signal_detail(sig_id, blk)
        bg_hex, fg_hex = _status_colors(status)

        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

        # Signal name
        c_sig = row.cells[0]
        p_sig = _clear(c_sig)
        _run(p_sig, f"{sig_id} — {name}", bold=True, size=_BODY_PT - 0.5)

        # Description
        c_desc = row.cells[1]
        p_desc = _clear(c_desc)
        _run(p_desc, desc, size=_BODY_PT - 0.5)

        # Status with color
        c_st = row.cells[2]
        _shd(c_st, bg_hex)
        p_st = _clear(c_st)
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_st, status, bold=True, size=_BODY_PT - 0.5, hex_color=fg_hex)

        # Key detail
        c_det = row.cells[3]
        p_det = _clear(c_det)
        _run(p_det, detail, size=_BODY_PT - 0.5)

    doc.add_paragraph()


def _signal_detail(sig_id: str, blk: Dict[str, Any]) -> str:
    """Produce a concise one-line detail string for the scorecard."""
    if sig_id == "S1":
        errs  = blk.get("error_count", 0)
        warns = blk.get("warning_count", 0)
        return f"{errs} error(s), {warns} warning(s)"
    if sig_id == "S2":
        counts = blk.get("counts") or {}
        return (
            f"Global KRIs: {counts.get('global_kris', '—')}  "
            f"Study KRIs: {counts.get('study_specific_kris', '—')}  "
            f"QTLs: {counts.get('qtls', '—')}"
        )
    if sig_id == "S3":
        n = blk.get("violation_count", 0)
        return f"{n} threshold violation(s)"
    if sig_id == "S4":
        n = blk.get("missing_count", 0)
        return f"{n} KRI(s) missing iqmp_risk_id"
    if sig_id == "S5":
        total = blk.get("total_with_tier", 0)
        low_n = blk.get("low_or_review_count", 0)
        rate  = blk.get("low_or_review_rate", 0.0)
        return f"{low_n}/{total} low/review ({rate*100:.0f}%)"
    if sig_id == "S6":
        n = blk.get("duplicate_count", 0)
        return f"{n} duplicate KRI label(s)"
    if sig_id == "S7":
        return str(blk.get("description") or "—")[:80]
    if sig_id == "S8":
        anchors = blk.get("anchors_present") or {}
        found = [k for k, v in anchors.items() if v]
        return f"Anchors present: {found or 'none'}"
    return "—"


def _add_detail_section(doc: Document, signals: Dict[str, Any]) -> None:
    """Detailed tables for signals that have lists of issues."""
    _section_heading(doc, "2. Signal Detail")

    # S1 — warnings
    s1 = signals.get("S1") or {}
    warns = (s1.get("sample_warnings") or []) + (s1.get("sample_errors") or [])
    if warns:
        _sub_heading(doc, "S1 — Structure Validity  ·  Warnings")
        rows: List[List[str]] = [["Level", "Code", "Message", "Path"]]
        for w in warns[:30]:
            rows.append([
                str(w.get("level") or ""),
                str(w.get("code") or ""),
                str(w.get("message") or "")[:120],
                str(w.get("path") or ""),
            ])
        _simple_table(doc, rows, col_widths=[Inches(0.8), Inches(1.6), Inches(2.8), Inches(1.3)])

    # S4 — missing iqmp_risk_id
    s4 = signals.get("S4") or {}
    missing = s4.get("missing") or []
    if missing:
        _sub_heading(doc, f"S4 — Protocol Provenance  ·  {len(missing)} KRI(s) missing iqmp_risk_id")
        doc.add_paragraph(
            "NOTE: iqmp_risk_id links each KRI back to the risk register. Missing IDs reduce "
            "downstream traceability but do not invalidate the KRI itself. Populate during "
            "CMP review with the corresponding IQMP risk reference.",
            style="Normal"
        ).italic = True
        rows = [["Bucket", "KRI ID", "KRI Label", "Issue"]]
        for m in missing[:50]:
            rows.append([
                str(m.get("bucket") or ""),
                str(m.get("kri_id") or ""),
                str(m.get("kri_label") or "")[:80],
                str(m.get("issue") or ""),
            ])
        _simple_table(doc, rows, col_widths=[Inches(1.3), Inches(1.0), Inches(3.2), Inches(1.0)])

    # S5 — low/review confidence items
    s5 = signals.get("S5") or {}
    rev_items = s5.get("review_items") or []
    if rev_items:
        _sub_heading(doc, f"S5 — Confidence Distribution  ·  {len(rev_items)} low/review item(s)")
        rows = [["Bucket", "KRI ID", "KRI Label", "Confidence"]]
        for item in rev_items[:30]:
            rows.append([
                str(item.get("bucket") or ""),
                str(item.get("kri_id") or ""),
                str(item.get("kri_label") or "")[:80],
                str(item.get("confidence") or ""),
            ])
        _simple_table(doc, rows, col_widths=[Inches(1.3), Inches(1.0), Inches(3.2), Inches(1.0)])

    # S8 — USDM anchors
    s8 = signals.get("S8") or {}
    anchors = s8.get("anchors_present") or {}
    if anchors:
        _sub_heading(doc, "S8 — USDM Traceability  ·  Anchor Classes")
        rows = [["USDM Class", "Found in Protocol"]]
        for cls, present in anchors.items():
            rows.append([cls, "Yes" if present else "No"])
        _simple_table(doc, rows, col_widths=[Inches(2.5), Inches(2.0)])
        if s8.get("usdm_types_sample"):
            doc.add_paragraph(
                f"Protocol instance types (sample): "
                + ", ".join(s8["usdm_types_sample"][:12]) + " …",
                style="Normal"
            ).italic = True


def _add_kri_field_check(doc: Document, signals: Dict[str, Any]) -> None:
    """Show a KRI field presence summary table using S2, S4 signal data."""
    s2_sig = signals.get("S2") or {}
    s4_sig = signals.get("S4") or {}
    counts  = s2_sig.get("counts") or {}
    missing = s4_sig.get("missing") or []

    _section_heading(doc, "KRI Field Structure Check")
    p = doc.add_paragraph()
    _run(p, "Verifying key fields are present for all KRIs.  "
         "IQMP Risk ID links each KRI back to the risk register — missing IDs need manual population.",
         italic=True, size=_CELL_PT, hex_color="555555")

    # Summary counts table
    rows = [["Bucket", "Count", "iqmp_risk_id missing", "Status"]]
    total_missing = len([m for m in missing if m.get("bucket") == "global_kris"])
    total_ss_missing = len([m for m in missing if m.get("bucket") == "study_specific_kris"])
    rows.append(["Global KRIs",         str(counts.get("global_kris", "?")),
                 str(total_missing),     "NEEDS REVIEW" if total_missing else "OK ✓"])
    rows.append(["Study-Specific KRIs", str(counts.get("study_specific_kris", "?")),
                 str(total_ss_missing),  "NEEDS REVIEW" if total_ss_missing else "OK ✓"])
    rows.append(["QTLs",               str(counts.get("qtls", "?")), "0", "OK ✓"])
    _simple_table(doc, rows, col_widths=[Inches(1.8), Inches(0.9), Inches(1.8), Inches(2.0)])

    # Full detail of KRIs with missing provenance (all of them)
    if missing:
        _sub_heading(doc, f"All {len(missing)} KRIs requiring iqmp_risk_id — populate before submission")
        rows2 = [["Bucket", "KRI ID", "KRI Label", "Action Required"]]
        for m in missing:
            rows2.append([
                str(m.get("bucket") or "").replace("_kris","").replace("_", " ").title(),
                str(m.get("kri_id") or ""),
                str(m.get("kri_label") or "")[:70],
                "Add IQMP risk ID",
            ])
        _simple_table(doc, rows2, col_widths=[Inches(1.2), Inches(0.9), Inches(3.0), Inches(1.4)])


def _add_review_list(doc: Document, review_list: List[Dict[str, Any]]) -> None:
    if not review_list:
        return
    _section_heading(doc, f"3. Items Requiring Review  ({len(review_list)})")
    for item in review_list:
        sig = str(item.get("signal_id") or "")
        reason = str(item.get("reason") or "")
        p = doc.add_paragraph(style="List Bullet")
        _run(p, f"{sig}: ", bold=True, size=_BODY_PT)
        _run(p, reason, size=_BODY_PT)


def _sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT + 0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def _simple_table(doc: Document, rows: List[List[str]],
                  col_widths: Optional[List] = None) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)
    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            text = row_data[ci] if ci < len(row_data) else ""
            cell = t.rows[ri].cells[ci]
            if col_widths and ci < len(col_widths):
                cell.width = col_widths[ci]
            p = _clear(cell)
            is_header = ri == 0
            if is_header:
                _shd(cell, _HEADER_BG)
                _run(p, text, bold=True, size=_BODY_PT - 0.5, rgb=_WHITE)
            else:
                _run(p, text, size=_BODY_PT - 0.5)
    doc.add_paragraph()


def _compute_health(signals: Dict[str, Any]) -> Dict[str, Any]:
    weights = {"PASS": 1.0, "WARN": 0.5, "FAIL": 0.0}
    pass_n = warn_n = fail_n = 0
    weighted = 0.0
    for blk in signals.values():
        st = str(blk.get("status") or "").upper()
        pass_n += st == "PASS"
        warn_n += st == "WARN"
        fail_n += st == "FAIL"
        weighted += weights.get(st, 0.0)
    total = pass_n + warn_n + fail_n
    pct = round((weighted / total * 100) if total else 0.0, 1)
    return {"percent": pct, "pass": pass_n, "warn": warn_n, "fail": fail_n, "total": total}


# ── Public API ────────────────────────────────────────────────────────────────

def write_cmp_scenario2_docx(s2: Dict[str, Any], output_path: str) -> str:
    """Generate the CMP Scenario 2 Word report and return the resolved output path."""
    study_id = str(s2.get("study_id") or "Unknown")
    ts = str(s2.get("timestamp") or "")[:10] or datetime.now().strftime("%Y-%m-%d")
    verdict = str(s2.get("verdict") or "AMBER")
    signals = s2.get("signals") or {}
    review_list = s2.get("review_list") or []

    health = _compute_health(signals)

    doc = Document()
    _apply_defaults(doc)

    _add_banner(doc, study_id, "Clinical Monitoring Plan", ts)
    _add_verdict_box(doc, verdict, health)
    _add_score_breakdown(doc, signals, health)
    _add_kri_history(doc, s2)
    _add_scorecard(doc, signals)
    _add_detail_section(doc, signals)
    _add_kri_field_check(doc, signals)
    _add_review_list(doc, review_list)

    # Footer
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    _run(p_foot,
         f"Protocol Digitalization Platform — USDM 4.0  ·  {ts}  ·  CONFIDENTIAL",
         italic=True, size=_CELL_PT - 1, hex_color="777777")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
        return str(out.resolve())
    except PermissionError:
        alt = out.with_name(f"{out.stem}_{datetime.now():%Y%m%d_%H%M%S}{out.suffix}")
        doc.save(str(alt))
        return str(alt.resolve())


def write_cmp_scenario2_docx_from_json(json_path: str, output_path: str) -> str:
    """Load a Scenario 2 JSON file and write the Word report."""
    with open(json_path, encoding="utf-8") as fh:
        s2 = json.load(fh)
    return write_cmp_scenario2_docx(s2, output_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python cmp_scenario2_report.py <s2_json> <output.docx>")
        sys.exit(1)
    out = write_cmp_scenario2_docx_from_json(sys.argv[1], sys.argv[2])
    print(f"Wrote: {out}")
