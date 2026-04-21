"""
Convert Markdown (subset) to .docx.

Supports: #–#### headings, bullets (-/*), paragraphs, fenced ``` blocks,
simple GFM tables (| col | col |).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = text.strip()
            if ri == 0:
                _set_cell_shading(cell, "D9E2F3")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def _split_table_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    core = s.replace("|", "")
    return bool(re.match(r"^[\s:\-]+$", core))


def write_docx_from_markdown(markdown_text: str, output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph("\n".join(code_lines) if code_lines else "")
                para.paragraph_format.left_indent = Pt(12)
                for r in para.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: List[List[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_separator(row_line):
                    i += 1
                    continue
                table_rows.append(_split_table_row(row_line))
                i += 1
            _add_table(doc, table_rows)
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            doc.add_paragraph(stripped)
        i += 1

    doc.save(str(path))
