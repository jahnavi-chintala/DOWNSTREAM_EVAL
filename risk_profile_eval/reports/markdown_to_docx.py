"""
Convert Markdown (subset) to .docx.

Supports: #–#### headings, bullets (-/*), paragraphs, fenced ``` blocks,
simple GFM tables (| col | col |).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Pfizer brand constants — match S1 eval reports ───────────────────────────
_FONT       = "Arial"
_BODY_PT    = 10.0
_CELL_PT    = 9.0
_HEAD_PT    = 11.0
_HEADER_BG  = "003087"   # Pfizer blue
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _thin_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblB.append(el)
    tblPr.append(tblB)


def _strip_md(text: str) -> str:
    """Remove **bold**, *italic*, _italic_ markdown markers from text."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"_(.*?)_",       r"\1", text)
    text = re.sub(r"`(.*?)`",       r"\1", text)
    return text


def _add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    _thin_borders(table)
    for ri, row in enumerate(rows):
        for ci in range(cols):
            raw = row[ci] if ci < len(row) else ""
            text = _strip_md(raw).strip()
            cell = table.rows[ri].cells[ci]
            # Clear default paragraph
            for p in list(cell.paragraphs):
                for run in list(p.runs):
                    run.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = _FONT
            r.font.size = Pt(_CELL_PT)
            if ri == 0:
                _set_cell_shading(cell, _HEADER_BG)
                r.bold = True
                r.font.color.rgb = _WHITE
    doc.add_paragraph()


def _split_table_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    core = s.replace("|", "")
    return bool(re.match(r"^[\s:\-]+$", core))


def _heading_as_blue_bar(doc: Document, text: str, level: int) -> None:
    """Render ## / ### headings as full-width Pfizer-blue bars (matches S1 style)."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    bg = _HEADER_BG if level <= 2 else "D9E2F3"
    fg = _WHITE if level <= 2 else RGBColor(0x00, 0x30, 0x87)
    _set_cell_shading(cell, bg)
    for p in list(cell.paragraphs):
        for r in list(p.runs):
            r.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(_strip_md(text))
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(_HEAD_PT if level <= 2 else _BODY_PT)
    r.font.color.rgb = fg
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def write_docx_from_markdown(markdown_text: str, output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # Match S1 eval: Arial 10pt, narrow margins
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)

    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph("\n".join(code_lines) if code_lines else "")
                para.paragraph_format.left_indent = Pt(12)
                for r in para.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: List[List[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_separator(row_line):
                    i += 1
                    continue
                table_rows.append(_split_table_row(row_line))
                i += 1
            _add_table(doc, table_rows)
            continue

        if stripped.startswith("#### "):
            _heading_as_blue_bar(doc, stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            _heading_as_blue_bar(doc, stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            _heading_as_blue_bar(doc, stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            _heading_as_blue_bar(doc, stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(_strip_md(stripped[2:].strip()))
            r.font.name = _FONT
            r.font.size = Pt(_BODY_PT)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(_strip_md(re.sub(r"^\d+\.\s", "", stripped)))
            r.font.name = _FONT
            r.font.size = Pt(_BODY_PT)
        elif stripped.startswith("_") and stripped.endswith("_"):
            # italic note line
            p = doc.add_paragraph()
            r = p.add_run(_strip_md(stripped))
            r.italic = True
            r.font.name = _FONT
            r.font.size = Pt(_CELL_PT)
        else:
            p = doc.add_paragraph()
            r = p.add_run(_strip_md(stripped))
            r.font.name = _FONT
            r.font.size = Pt(_BODY_PT)
        i += 1

    doc.save(str(path))
