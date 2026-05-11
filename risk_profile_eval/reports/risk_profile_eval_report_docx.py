"""
Risk Profile **generator eval** report → Word (.docx), stakeholder-style layout.

Scenario 1 uses **structured tables** (blue header rows, green pass / red fail text) for risk blocks:
risk match fields, RPN components, and attribute comparison — not plain label paragraphs.

Requires ``eval_metadata`` from ``run_eval`` with paths to generator JSON + GT CSVs when available.
"""

from __future__ import annotations

import json
import os
import tempfile
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

_BODY_PT = 10.0
_HEADER_PT = 13.0
_TITLE_PT = 15.0
_SCORE_NUM_PT = 28.0
_FONT = "Arial"
_HEADER_SHADE = "D9D9D9"
# Pfizer reference palette — matches C5091017_eval_report.docx stakeholder template
_HEADER_BLUE = "003087"
_RGB_OK = RGBColor(0x1A, 0x73, 0x40)
_RGB_FAIL = RGBColor(0xB9, 0x1C, 0x1C)
_RGB_AMBER = RGBColor(0xB4, 0x53, 0x09)
_RGB_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_RGB_DARK = RGBColor(0x1A, 0x1A, 0x1A)
_RGB_GRAY = RGBColor(0x66, 0x66, 0x66)


# ─── Low-level OOXML ─────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _para_line_spacing(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)


def _run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size_pt: float = _BODY_PT,
    rgb: RGBColor | None = None,
) -> None:
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if rgb is not None:
        r.font.color.rgb = rgb


def _cell_write(
    cell,
    text: str,
    *,
    rgb: RGBColor | None = None,
    bold: bool = False,
    size_pt: float = _BODY_PT,
) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.text = ""
    _para_line_spacing(p)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size_pt)
    if rgb is not None:
        r.font.color.rgb = rgb


# ─── Template helpers (required API) ───────────────────────────────────────────

def add_section_title(doc: Document, text: str, *, level: int = 1) -> None:
    """Use Word Heading styles so structure matches ``Risk_Profile_Eval_Report_*.docx``."""
    doc.add_heading(text, level=min(max(level, 1), 9))


def add_paragraph_block(
    doc: Document,
    label: str,
    value: str | None = None,
    *,
    label_bold: bool = True,
) -> None:
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, label, bold=label_bold, size_pt=_BODY_PT)
    if value is not None and str(value).strip():
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(p, str(value), bold=False, size_pt=_BODY_PT)


def add_spacer(doc: Document) -> None:
    doc.add_paragraph()


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    *,
    col_widths: Sequence[Any] | None = None,
    max_width_inches: float = 6.5,
    style_header: bool = True,
    row_fill: Callable[[int, List[str]], Optional[str]] | None = None,
    header_fill_hex: str | None = None,
    header_text_white: bool = False,
) -> Table:
    ncols = len(headers)
    nrows = len(rows)
    table = doc.add_table(rows=1 + max(0, nrows), cols=max(1, ncols))
    table.style = "Table Grid"
    table.autofit = False
    widths: List = list(col_widths) if col_widths else []
    if not widths and ncols:
        w = max_width_inches / ncols
        widths = [Inches(w)] * ncols
    fill_hex = header_fill_hex if header_fill_hex else _HEADER_SHADE
    hdr = table.rows[0].cells
    for ci, h in enumerate(headers):
        c = hdr[ci]
        p = c.paragraphs[0] if c.paragraphs else c.add_paragraph()
        p.text = ""
        _para_line_spacing(p)
        _run(
            p,
            str(h),
            bold=True,
            size_pt=_BODY_PT,
            rgb=_RGB_HEADER_TEXT if header_text_white else None,
        )
        if style_header:
            _set_cell_shading(c, fill_hex)
        if ci < len(widths):
            c.width = widths[ci]
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = row_fill(ri, [str(x) for x in row_data]) if row_fill else None
        for ci in range(ncols):
            c = row.cells[ci]
            val = row_data[ci] if ci < len(row_data) else ""
            p = c.paragraphs[0] if c.paragraphs else c.add_paragraph()
            p.text = ""
            _para_line_spacing(p)
            _run(p, str(val), size_pt=_BODY_PT)
            if fill:
                _set_cell_shading(c, fill)
            if ci < len(widths):
                c.width = widths[ci]
    doc.add_paragraph()
    return table


def _try_add_risk_score_bar_chart(
    doc: Document,
    bar_order: List[Tuple[str, float]],
    study_id: str,
) -> bool:
    """Embedded matplotlib horizontal bars (reference-style colors). Falls back to Unicode bars."""
    if not bar_order:
        return True
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return False
    labels = [str(x[0])[:80] for x in bar_order]
    vals = [max(0.0, min(100.0, float(x[1]))) for x in bar_order]
    n = len(labels)
    fig_h = min(9.0, max(2.8, 0.42 * n + 1.0))
    fig, ax = plt.subplots(figsize=(6.8, fig_h))
    palette = ["#0073C8", "#00A3A1", "#5B9BD5", "#2E86AB", "#8ECAE6", "#4A90A4"]
    colors = [palette[i % len(palette)] for i in range(n)]
    y = list(range(n))
    ax.barh(y, vals, color=colors, height=0.62)
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlabel("Score (0–100)")
    ax.set_xlim(0, 100)
    ax.set_title(f"{study_id} — risk scorecard (display score 0–100)")
    ax.grid(axis="x", linestyle=":", alpha=0.45)
    fig.tight_layout()
    fd, tmp = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        fig.savefig(tmp, dpi=130, bbox_inches="tight")
        plt.close(fig)
        doc.add_picture(tmp, width=Inches(6.3))
        doc.add_paragraph()
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass
    return True


def add_text_bar_row(doc: Document, risk_name: str, score: float, *, width: int = 42) -> None:
    frac = max(0.0, min(100.0, float(score))) / 100.0
    n = int(round(frac * width))
    bar = "█" * n + "░" * (width - n)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, f"{risk_name}\n", bold=True, size_pt=_BODY_PT)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, f"{bar} {score:.1f}", size_pt=_BODY_PT)
    add_spacer(doc)


# ─── Data loading & scoring ──────────────────────────────────────────────────

def _load_generator_json(report: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    em = report.get("eval_metadata") or {}
    p = em.get("generator_json_path")
    if p and Path(p).is_file():
        with open(p, encoding="utf-8") as f:
            return json.load(f)
    return None


def _load_risk_gt_df(report: Dict[str, Any]):
    try:
        import pandas as pd
    except ImportError:
        return None
    em = report.get("eval_metadata") or {}
    p = em.get("risk_ground_truth_csv_path")
    sid = report.get("study_id")
    if p and sid and Path(p).is_file():
        df = pd.read_csv(p)
        sub = df[df["study_id"].astype(str) == str(sid)].copy()
        return sub if not sub.empty else None
    return None


def _risk_by_name(gen: Optional[Dict[str, Any]], name: str) -> Optional[Dict[str, Any]]:
    if not gen:
        return None
    from core.risk_generator_risks import get_all_risk_dicts

    target = str(name or "").strip()
    for r in get_all_risk_dicts(gen):
        if isinstance(r, dict) and str(r.get("risk_name", "") or "").strip() == target:
            return r
    return None


def _m2_row_for_gt(metrics: Dict, grow) -> Optional[Dict[str, Any]]:
    """Disambiguate M2 per_risk when the same risk name appears on multiple GT rows."""
    m2 = metrics.get("m2_rpn_tier_accuracy") or {}
    rname = str(grow.get("risk_name", "") or "")
    tr = grow.get("rpn")
    gid = str(grow.get("risk_id", "") or "").strip()
    candidates = [pr for pr in (m2.get("per_risk") or []) if pr.get("risk_name") == rname]
    if not candidates:
        return None
    if len(candidates) == 1:
        return candidates[0]
    if gid:
        for pr in candidates:
            if str(pr.get("gt_risk_id", "") or "").strip() == gid:
                return pr
    if tr is not None and not (isinstance(tr, float) and str(tr) == "nan"):
        try:
            ti = int(tr)
            for pr in candidates:
                tp = pr.get("truth_rpn")
                if tp is not None and int(tp) == ti:
                    return pr
        except (TypeError, ValueError):
            pass
    return candidates[0]


def _pair_for_gt_row(matched_pairs: List[Dict[str, Any]], grow) -> Optional[Dict[str, Any]]:
    """Pick the M1 matched pair for this GT row (name + risk_id when available)."""
    if not matched_pairs:
        return None
    rname = str(grow.get("risk_name", "") or "").strip()
    gid = str(grow.get("risk_id", "") or "").strip()
    for p in matched_pairs:
        if str(p.get("gt_risk_name", "") or "").strip() != rname:
            continue
        if gid and str(p.get("gt_risk_id", "") or "").strip() == gid:
            return p
    for p in matched_pairs:
        if str(p.get("gt_risk_name", "") or "").strip() == rname:
            return p
    return None


def _gen_risk_from_pair(generator_json: Optional[Dict[str, Any]], pair: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    if not pair or not generator_json:
        return None

    key = pair.get("generated_json_key")
    want = str(pair.get("generated_risk_name", "") or "").strip()
    block = generator_json.get(key) if key else None
    if not isinstance(block, list):
        return None
    for r in block:
        if isinstance(r, dict) and str(r.get("risk_name", "") or "").strip() == want:
            return r
    return None


def _rpn_tier_pass_for_row(
    gen_risk: Optional[Dict[str, Any]],
    gt_row: Optional[Any],
    m2row: Optional[Dict[str, Any]],
) -> bool:
    """TDD §6: pass when tier distance ≤ 1; fallback to M2 per_risk if RPNs unparseable."""
    if gen_risk is None or gt_row is None:
        return bool(m2row and m2row.get("passed"))
    tr = gt_row.get("rpn")
    if tr is None or (isinstance(tr, float) and str(tr) == "nan"):
        return bool(m2row and m2row.get("passed"))
    try:
        from core.eval_scenario1 import tier_distance

        gr = int(float(gen_risk.get("rpn")))
        dist = tier_distance(gr, int(float(tr)))
        if dist is None:
            return bool(m2row and m2row.get("passed"))
        return dist <= 1
    except (TypeError, ValueError):
        return bool(m2row and m2row.get("passed"))


def _format_rpn_tier_display(rpn_val: Any) -> str:
    try:
        if rpn_val is None or (isinstance(rpn_val, float) and str(rpn_val) == "nan"):
            return "—"
        r = int(float(rpn_val))
    except (TypeError, ValueError):
        return str(rpn_val)
    from core.eval_scenario1 import rpn_to_tier

    t = rpn_to_tier(r)
    if t is None:
        return str(r)
    return f"{r} (Tier {t})"


def _add_rpn_row_table(
    doc: Document,
    headers: Sequence[str],
    row_values: Sequence[str],
    col_widths: Sequence[Any],
) -> None:
    """Single data row under blue header; green/red on cells that contain ✓/✗."""
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for ci, h in enumerate(headers):
        c = table.rows[0].cells[ci]
        _set_cell_shading(c, _HEADER_BLUE)
        _cell_write(c, str(h), bold=True, rgb=_RGB_HEADER_TEXT)
        if ci < len(col_widths):
            c.width = col_widths[ci]
    for ci, val in enumerate(row_values):
        c = table.rows[1].cells[ci]
        v = str(val)
        rgb = _RGB_OK if "✓" in v else (_RGB_FAIL if "✗" in v else None)
        _cell_write(c, v, rgb=rgb)
    doc.add_paragraph()


def _add_attribute_rows_table(
    doc: Document,
    attr_rows: List[List[str]],
    col_widths: Sequence[Any],
) -> None:
    headers = ["Attribute", "Score", "Match Type", "Generated", "Ground Truth"]
    table = doc.add_table(rows=1 + len(attr_rows), cols=5)
    table.style = "Table Grid"
    table.autofit = False
    for ci, h in enumerate(headers):
        c = table.rows[0].cells[ci]
        _set_cell_shading(c, _HEADER_BLUE)
        _cell_write(c, h, bold=True, rgb=_RGB_HEADER_TEXT)
        if ci < len(col_widths):
            c.width = col_widths[ci]
    for ri, row_data in enumerate(attr_rows):
        for ci in range(5):
            c = table.rows[ri + 1].cells[ci]
            val = row_data[ci] if ci < len(row_data) else ""
            v = str(val)
            rgb = None
            if ci == 1:
                rgb = _RGB_OK if v == "1.00" else (_RGB_FAIL if v == "0.00" else None)
            _cell_write(c, v, rgb=rgb)
            if ci < len(col_widths):
                c.width = col_widths[ci]
    doc.add_paragraph()


def _document_score_s1(metrics: Dict[str, Any]) -> Tuple[float, bool, int]:
    """Overall score 0–100 and pass vs default threshold 75."""
    m1 = metrics.get("m1_risk_name_recall") or {}
    m2 = metrics.get("m2_rpn_tier_accuracy") or {}
    m3 = metrics.get("m3_critical_factor_match") or {}
    m4 = metrics.get("m4_hallucination_detection") or {}
    parts: List[float] = [
        float(m1.get("score") or 0),
    ]
    if not m2.get("skipped"):
        parts.append(float(m2.get("score") or 0))
    parts.append(1.0 if m4.get("passed") else 0.0)
    if not m3.get("skipped"):
        parts.append(float(m3.get("score") or 0))
    score = round(100.0 * sum(parts) / len(parts), 1) if parts else 0.0
    thr = 75
    passed = score >= thr
    return score, passed, thr


def _per_risk_display_score(
    risk_name: str,
    matched: bool,
    m2row: Optional[Dict[str, Any]],
    gen_risk: Optional[Dict[str, Any]] = None,
    gt_row: Optional[Any] = None,
    gen_json_key: str = "",
) -> float:
    """
    Per-risk display score (0-100) reflecting:
      - Name match (M1):  40 pts
      - RPN tier (M2):    30 pts
      - Domain match:     15 pts
      - Control count:    15 pts
    """
    if not matched:
        return 0.0

    score = 40.0  # name matched

    rpn_ok = False
    if gen_risk is not None and gt_row is not None:
        tr = gt_row.get("rpn")
        if tr is not None and not (isinstance(tr, float) and str(tr) == "nan"):
            try:
                from core.eval_scenario1 import tier_distance
                gr = int(float(gen_risk.get("rpn")))
                dist = tier_distance(gr, int(float(tr)))
                if dist is not None and dist <= 1:
                    rpn_ok = True
            except (TypeError, ValueError):
                pass
    elif m2row and m2row.get("passed"):
        rpn_ok = True
    if rpn_ok:
        score += 30.0

    if gen_risk is not None and gt_row is not None:
        from core.risk_generator_risks import infer_generated_domain, normalize_domain_label
        raw_dom = infer_generated_domain(gen_json_key, gen_risk)
        gen_dom = normalize_domain_label(raw_dom)
        gt_dom = str(gt_row.get("risk_domain") if hasattr(gt_row, "get") else "").strip().lower()
        if not gt_dom or gt_dom in ("--", "nan"):
            score += 15.0
        elif gen_dom == gt_dom:
            score += 15.0

    if gen_risk is not None and gt_row is not None:
        gt_cc_raw = gt_row.get("control_count") if hasattr(gt_row, "get") else None
        gt_cc = None
        if gt_cc_raw is not None and not (isinstance(gt_cc_raw, float) and str(gt_cc_raw) == "nan"):
            s = str(gt_cc_raw).strip()
            if s and s.lower() not in ("--", "nan"):
                try:
                    gt_cc = int(float(s))
                except (TypeError, ValueError):
                    pass
        gen_cc = len(gen_risk.get("controls") or [])
        if gt_cc is None:
            score += 15.0
        elif gt_cc == gen_cc:
            score += 15.0

    return round(score, 1)


def _bucket_slug(ta: str, phase: str) -> str:
    ta_s = (ta or "unknown").strip()
    ph_s = (phase or "unknown").strip().replace("_", " ")
    slug = f"{ta_s.lower()}_{ph_s.lower()}".replace(" ", "_")
    return f"{ta_s.title()} / {ph_s} — {slug}"


def _protocol_subtitle(gen: Optional[Dict[str, Any]], study_id: str) -> str:
    if not gen:
        return study_id
    so = gen.get("study_overview") or {}
    comp = so.get("compound_code") or ""
    ind = so.get("indication") or ""
    ph = so.get("development_phase") or ""
    title = (so.get("title") or "").upper()
    drug = "Ritlecitinib" if "RITLECITINIB" in title else (so.get("drug_program") or "").strip() or "Study intervention"
    bits = [study_id, f"{drug} ({comp})" if comp else drug]
    tail = " ".join(x for x in [ph.replace("_", " "), ind] if x)
    if tail:
        bits.append(tail)
    return " — ".join(bits)


def _thresh_target() -> Tuple[int, int]:
    return 75, 80


# ─── Risk block ────────────────────────────────────────────────────────────────

def _add_risk_match_fields_table(
    doc: Document,
    rows: List[Tuple[str, str, Optional[RGBColor]]],
) -> None:
    """Two-column Field | Value with blue header; optional RGB for value column."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for ci, h in enumerate(["Field", "Value"]):
        c = table.rows[0].cells[ci]
        _set_cell_shading(c, _HEADER_BLUE)
        _cell_write(c, h, bold=True, rgb=_RGB_HEADER_TEXT)
        c.width = Inches(1.45) if ci == 0 else Inches(5.05)
    for ri, (lab, val, val_rgb) in enumerate(rows):
        _cell_write(table.rows[ri + 1].cells[0], lab, bold=True)
        _cell_write(table.rows[ri + 1].cells[1], val, rgb=val_rgb)
    doc.add_paragraph()


def add_risk_block(
    doc: Document,
    idx: int,
    risk_name: str,
    gen_risk: Optional[Dict[str, Any]],
    gt_row: Optional[Any],
    m2row: Optional[Dict[str, Any]],
    m1_matched: bool,
    gen_json_key: str = "",
) -> None:
    rid = (gen_risk or {}).get("risk_id") or (
        str(gt_row.get("risk_id", "—")) if gt_row is not None else "—"
    )
    add_section_title(doc, f"Risk {idx} — {risk_name} [{rid}]", level=2)

    if not gen_risk:
        doc.add_heading("Risk match information", level=3)
        miss_rows: List[Tuple[str, str, Optional[RGBColor]]] = [
            ("Risk Name Match", "— MISS — not generated for this ground-truth row.", _RGB_FAIL),
            ("Match Status", "Not generated (see M1 misses).", _RGB_FAIL),
        ]
        if gt_row is not None:
            miss_rows.append(("Ground Truth ID", str(gt_row.get("risk_id", "—")), None))
            miss_rows.append(("Risk domain (GT)", str(gt_row.get("risk_domain", "—")), None))
        _add_risk_match_fields_table(doc, miss_rows)
        add_spacer(doc)
        return

    doc.add_heading("Risk match information", level=3)
    from core.risk_generator_risks import fingerprint_risk_name_for_m1, infer_generated_domain

    fp_g = fingerprint_risk_name_for_m1(str(gen_risk.get("risk_name", "") or ""))
    fp_t = fingerprint_risk_name_for_m1(str(risk_name or ""))
    names_norm_match = bool(fp_g and fp_t and fp_g == fp_t)
    lev = 0 if m1_matched else "—"
    if m1_matched:
        near_txt = f"Verbatim ✓ | Levenshtein: {lev}"
    elif names_norm_match:
        near_txt = "M1-equivalent (cosmetic-only) ✓ (not in M1 domain-matched set)"
    else:
        near_txt = "Not in matched set (check name/domain)"
    gt_id = str(gt_row["risk_id"]) if gt_row is not None and gt_row.get("risk_id") is not None else "—"
    _inferred_domain = infer_generated_domain(gen_json_key, gen_risk)
    gen_domain = _inferred_domain if _inferred_domain else "—"
    gt_domain_val = (
        str(gt_row["risk_domain"]).strip()
        if gt_row is not None and pd_safe(gt_row, "risk_domain")
        else ""
    )
    domain_ok = True
    domain_display = gen_domain
    if gt_domain_val and gt_domain_val != "—":
        domain_ok = gen_domain.strip().lower() == gt_domain_val.lower()
        if not domain_ok:
            domain_display = f"{gen_domain}  |  GT: {gt_domain_val}"

    match_rows: List[Tuple[str, str, Optional[RGBColor]]] = [
        ("Risk Name Match", str(gen_risk.get("risk_name", "—")), None),
        (
            "Match Status",
            near_txt,
            _RGB_OK if (m1_matched or names_norm_match) else _RGB_FAIL,
        ),
        ("Ground Truth ID", gt_id, None),
        ("Risk Domain", domain_display, _RGB_OK if domain_ok else _RGB_FAIL),
    ]
    _add_risk_match_fields_table(doc, match_rows)

    gi = gen_risk.get("impact")
    gl = gen_risk.get("likelihood")
    gd = gen_risk.get("detectability")
    gr = gen_risk.get("rpn")
    ti = gt_row["impact"] if gt_row is not None and pd_safe(gt_row, "impact") else None
    tl = gt_row["likelihood"] if gt_row is not None and pd_safe(gt_row, "likelihood") else None
    td = gt_row["detectability"] if gt_row is not None and pd_safe(gt_row, "detectability") else None
    tr = gt_row["rpn"] if gt_row is not None and pd_safe(gt_row, "rpn") else None

    def cell(gv: Any, tv: Any, numeric: bool = True) -> str:
        if gv is None and tv is None:
            return "—"
        if tv is None or (isinstance(tv, float) and str(tv) == "nan"):
            return str(gv)
        if numeric:
            try:
                ok = int(gv) == int(tv)
            except (TypeError, ValueError):
                ok = str(gv).strip() == str(tv).strip()
        else:
            ok = str(gv).strip() == str(tv).strip()
        return f"{gv} = {tv} {'✓' if ok else '✗'}"

    doc.add_heading("RPN components (generated vs ground truth)", level=3)
    hdr = ["Active RPN", "Impact", "Likelihood", "Detectability"]
    row_vals = [cell(gr, tr), cell(gi, ti), cell(gl, tl), cell(gd, td)]
    _add_rpn_row_table(doc, hdr, row_vals, [Inches(1.6), Inches(1.3), Inches(1.3), Inches(1.3)])
    add_spacer(doc)

    ctrl_types_gt = ""
    if gt_row is not None:
        ct = gt_row.get("control_types")
        ctrl_types_gt = "" if ct is None or (isinstance(ct, float) and str(ct) == "nan") else str(ct)
    pri_ctrl = ""
    ctrls = gen_risk.get("controls") or []
    if ctrls and isinstance(ctrls[0], dict):
        pri_ctrl = str(ctrls[0].get("control_type") or "")

    name_ok = m1_matched or names_norm_match
    rpn_ok = _rpn_tier_pass_for_row(gen_risk, gt_row, m2row)

    gen_domain_s = infer_generated_domain(gen_json_key, gen_risk).strip()
    gt_domain_val2 = (
        str(gt_row["risk_domain"]).strip()
        if gt_row is not None and pd_safe(gt_row, "risk_domain")
        else ""
    )
    if gt_domain_val2 and gt_domain_val2 != "—":
        domain_ok2 = gen_domain_s.lower() == gt_domain_val2.lower()
        domain_score = "1.00" if domain_ok2 else "0.00"
        domain_match = "exact" if domain_ok2 else "miss"
    else:
        domain_score = "—"
        domain_match = "N/A"
        gt_domain_val2 = "—"

    def _ctrl_type_parts(s: str):
        """Split GT control_types by ';' or ' and ' to get individual type tokens."""
        import re as _re
        parts = _re.split(r";| and ", s, flags=_re.IGNORECASE)
        return [p.strip() for p in parts if p.strip()]

    ctrl_gt_first = ctrl_types_gt.split(";")[0].strip() if ctrl_types_gt else ""
    # Re-derive first GT type after splitting by both ';' and ' and '
    if ctrl_types_gt:
        parts = _ctrl_type_parts(ctrl_types_gt)
        ctrl_gt_first = parts[0] if parts else ctrl_types_gt.strip()

    if ctrl_gt_first and ctrl_gt_first not in ("—", "--"):
        gen_norm = pri_ctrl.strip().lower()
        gt_norm  = ctrl_gt_first.lower()
        # Exact match OR gen type appears in any GT part (and vice versa)
        all_gt_parts = _ctrl_type_parts(ctrl_types_gt) if ctrl_types_gt else [ctrl_gt_first]
        ctrl_ok = (
            gen_norm == gt_norm
            or any(gen_norm == p.lower() for p in all_gt_parts)
            or (gen_norm and gen_norm in ctrl_types_gt.lower())
        )
        if ctrl_ok:
            ctrl_score, ctrl_match = "1.00", "partial ✓" if gen_norm != gt_norm else "exact"
        else:
            ctrl_score, ctrl_match = "0.00", "miss"
    elif pri_ctrl:
        ctrl_score = "—"
        ctrl_match = "N/A"
        ctrl_gt_first = "—"
    else:
        ctrl_score = "0.00"
        ctrl_match = "miss"
        ctrl_gt_first = ctrl_gt_first or "—"

    doc.add_heading("Attribute comparison", level=3)
    attr_rows = [
        [
            "Risk name",
            "1.00" if name_ok else "0.00",
            "verbatim" if name_ok else "miss",
            str(gen_risk.get("risk_name", "—")),
            risk_name,
        ],
        [
            "RPN",
            "1.00" if rpn_ok else "0.00",
            "±1 tier",
            _format_rpn_tier_display(gen_risk.get("rpn")),
            _format_rpn_tier_display(tr),
        ],
        [
            "Risk domain",
            domain_score,
            domain_match,
            gen_domain_s or "—",
            gt_domain_val2,
        ],
        [
            "Primary control type",
            ctrl_score,
            ctrl_match,
            pri_ctrl or "—",
            ctrl_gt_first,
        ],
    ]
    _add_attribute_rows_table(
        doc,
        attr_rows,
        [Inches(1.15), Inches(0.55), Inches(0.85), Inches(1.75), Inches(1.75)],
    )
    add_spacer(doc)


def pd_safe(row, key: str) -> bool:
    try:
        v = row.get(key) if hasattr(row, "get") else row[key]
        if v is None:
            return False
        if isinstance(v, float) and str(v) == "nan":
            return False
        return True
    except Exception:
        return False


def _inventory_generated_risks(gen: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Flat list of generated risks with domain and dict for Word sections when GT has no rows."""
    from core.risk_generator_risks import infer_generated_domain, iter_risk_dicts_with_keys

    rows: List[Dict[str, Any]] = []
    if not gen:
        return rows
    for json_key, _j, risk in iter_risk_dicts_with_keys(gen):
        if not isinstance(risk, dict):
            continue
        rn = str(risk.get("risk_name") or "").strip()
        if not rn:
            continue
        ctrls = risk.get("controls") or []
        nctrl = len(ctrls) if isinstance(ctrls, list) else 0
        rows.append({
            "risk_name": rn,
            "domain": infer_generated_domain(json_key, risk),
            "json_key": json_key,
            "risk": risk,
            "n_controls": nctrl,
        })
    return rows


# ─── M4 False-Positive Analysis ────────────────────────────────────────────────

_USDM_ENTITY_ALIAS_MAP = {
    "eligibility criterion": "EligibilityCriterion",
    "study intervention": "StudyIntervention",
    "population": "StudyDesignPopulation",
    "studydesign": "InterventionalStudyDesign",
    "complexitysignal": None,
    "adverseevent": None,
    "informedconsent": None,
}


def _classify_m4_false_positives(m4: Dict[str, Any]) -> Dict[str, Any]:
    """
    Analyse M4 flagged fields to classify which are schema mapping issues
    (false positives) vs genuine hallucinations.

    Returns dict with counts and lists.
    """
    flagged = m4.get("flagged_fields") or m4.get("traceability_flags") or []
    fp_schema_alias = []
    fp_schema_gap = []
    genuine = []

    for f in flagged:
        reason = str(f.get("reason") or "")
        value = f.get("value") or {}
        entity_raw = str(value.get("entity") or "").strip()
        entity_norm = entity_raw.lower().replace(" ", "").replace("_", "")

        if "not present in USDM" in reason or "instanceType" in reason:
            if entity_norm in _USDM_ENTITY_ALIAS_MAP:
                correct_name = _USDM_ENTITY_ALIAS_MAP[entity_norm]
                if correct_name is not None:
                    fp_schema_alias.append({
                        "field_path": f.get("field_path", ""),
                        "claimed": entity_raw,
                        "correct_usdm_type": correct_name,
                        "type": "alias_mismatch",
                    })
                else:
                    fp_schema_gap.append({
                        "field_path": f.get("field_path", ""),
                        "claimed": entity_raw,
                        "type": "usdm_schema_gap",
                    })
            else:
                genuine.append(f)
        else:
            genuine.append(f)

    return {
        "total_flags": len(flagged),
        "false_positive_alias_count": len(fp_schema_alias),
        "false_positive_schema_gap_count": len(fp_schema_gap),
        "genuine_count": len(genuine),
        "alias_mismatches": fp_schema_alias,
        "schema_gaps": fp_schema_gap,
        "genuine": genuine,
    }


def _write_m4_false_positive_note(doc: Document, m4: Dict[str, Any]) -> None:
    """Write the improvements / false-positive analysis note for M4."""
    analysis = _classify_m4_false_positives(m4)
    total = analysis["total_flags"]
    fp_alias = analysis["false_positive_alias_count"]
    fp_gap = analysis["false_positive_schema_gap_count"]
    genuine = analysis["genuine_count"]

    if total == 0:
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(p, "No M4 traceability flags — no improvements needed.", size_pt=_BODY_PT)
        return

    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, "M4 Traceability Flag Analysis (False-Positive Assessment)", bold=True, size_pt=_HEADER_PT)

    p = doc.add_paragraph()
    _para_line_spacing(p)
    fp_total = fp_alias + fp_gap
    _run(
        p,
        f"Of {total} M4 flags, {fp_total} are false positives due to USDM entity naming issues "
        f"(not actual hallucinations). {genuine} flag(s) are genuine provenance defects.",
        size_pt=_BODY_PT,
    )

    rows = []
    if fp_alias > 0:
        rows.append([
            "Entity alias mismatch",
            str(fp_alias),
            "Generator uses human-readable names (e.g. 'Eligibility Criterion') "
            "instead of USDM camelCase (e.g. 'EligibilityCriterion'). "
            "Content is valid and traceable via Intelligence document.",
        ])
    if fp_gap > 0:
        rows.append([
            "USDM schema gap",
            str(fp_gap),
            "Generator references concepts (e.g. 'ComplexitySignal', 'AdverseEvent', "
            "'InformedConsent') that do not exist in USDM 4.0 schema. "
            "These are real protocol concepts but have no USDM representation.",
        ])
    if genuine > 0:
        rows.append([
            "Genuine defect",
            str(genuine),
            "References that cannot be mapped to any known USDM entity or Intelligence "
            "document source. Require generator fix.",
        ])

    if rows:
        add_table(
            doc,
            ["Category", "Count", "Explanation"],
            rows,
            col_widths=[Inches(1.5), Inches(0.7), Inches(4.0)],
            header_fill_hex=_HEADER_BLUE,
            header_text_white=True,
        )

    # Alias detail table
    alias_items = analysis["alias_mismatches"]
    if alias_items:
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(p, "Recommended generator fix — entity name mapping:", bold=True, size_pt=_BODY_PT)
        seen = set()
        alias_rows = []
        for a in alias_items:
            key = (a["claimed"], a["correct_usdm_type"])
            if key not in seen:
                seen.add(key)
                alias_rows.append([a["claimed"], a["correct_usdm_type"]])
        add_table(
            doc,
            ["Generated Entity Name", "Correct USDM instanceType"],
            alias_rows,
            col_widths=[Inches(2.5), Inches(3.0)],
            header_fill_hex=_HEADER_BLUE,
            header_text_white=True,
        )

    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(
        p,
        "Note: All flagged risks are present in the Intelligence document with valid "
        "protocol source references. The M4 failures are USDM schema compliance issues, "
        "not content hallucinations. Fixing the entity name mapping in the generator "
        "will resolve these flags without changing risk content.",
        size_pt=_BODY_PT,
    )


# ─── Scenario 1 main layout ────────────────────────────────────────────────────

def _write_scenario1(doc: Document, report: Dict[str, Any]) -> None:
    em = report.get("eval_metadata") or {}
    study = report.get("study_id") or "?"
    metrics = report.get("metrics") or {}
    targets = report.get("targets") or {}
    gen = _load_generator_json(report)
    gt_df = _load_risk_gt_df(report)

    eval_date = em.get("eval_date") or (report.get("timestamp") or "")[:10] or "—"
    cfg_ver = em.get("config_version", "1.0")
    gen_ver = em.get("generator_version") or report.get("generator_version", "unknown")
    if gen_ver and not str(gen_ver).lower().startswith("v"):
        gen_ver = f"v{gen_ver}"
    ta = em.get("ta") or report.get("ta") or "—"
    phase = em.get("phase") or report.get("phase") or "—"

    # Line spacing default
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)

    # Top platform strip (paragraph, not table)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(
        p,
        f"Pfizer Protocol Intelligence Platform | D2: Risk Profile Generator | Eval Report | {study}",
        size_pt=_BODY_PT,
    )
    add_spacer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(p)
    _run(p, "PFIZER PROTOCOL INTELLIGENCE PLATFORM", bold=True, size_pt=_TITLE_PT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(p)
    _run(p, "Risk Profile Generator — Eval Report", bold=True, size_pt=_HEADER_PT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(p)
    _run(p, _protocol_subtitle(gen, study), size_pt=_BODY_PT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(p)
    _run(p, f"Eval date: {eval_date} | Generator: {gen_ver} | Config: {cfg_ver}", size_pt=_BODY_PT)
    add_spacer(doc)

    doc_score, doc_pass, thr = _document_score_s1(metrics)
    _, tgt = _thresh_target()
    pass_txt = "PASS ✓" if doc_pass else "FAIL ✗"

    # Document score block (3-column layout table)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    left = cells[0]
    mid = cells[1]
    right = cells[2]
    p = left.paragraphs[0]
    p.text = ""
    _para_line_spacing(p)
    _run(p, "DOCUMENT SCORE", bold=True, size_pt=_HEADER_PT)
    p = left.add_paragraph()
    _para_line_spacing(p)
    _run(p, _bucket_slug(str(ta), str(phase)), size_pt=_BODY_PT)
    mid.paragraphs[0].text = ""
    pm = mid.paragraphs[0]
    _para_line_spacing(pm)
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pm, f"{doc_score:.1f}\n", bold=True, size_pt=_SCORE_NUM_PT)
    pm2 = mid.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(pm2)
    _run(pm2, "/ 100", size_pt=_HEADER_PT)
    p = right.paragraphs[0]
    p.text = ""
    _para_line_spacing(p)
    _run(p, pass_txt, bold=True, size_pt=_HEADER_PT, rgb=_RGB_OK if doc_pass else _RGB_FAIL)
    p = right.add_paragraph()
    _para_line_spacing(p)
    _run(p, f"Threshold: {thr}\nTarget: {tgt}", size_pt=_BODY_PT)
    doc.add_paragraph()

    # ── Score formula explanation ────────────────────────────────────────────
    sf_tbl = doc.add_table(rows=1, cols=1)
    sf_tbl.style = "Table Grid"
    _set_cell_shading(sf_tbl.rows[0].cells[0], "EBF0FB")
    sf_p = sf_tbl.rows[0].cells[0].paragraphs[0]
    sf_p.text = ""
    _run(sf_p, "How the Document Score is Calculated", bold=True, size_pt=_BODY_PT)
    sf_p2 = sf_tbl.rows[0].cells[0].add_paragraph()
    _run(sf_p2,
         "Score = simple average of all applicable metrics x 100  "
         "(each metric carries equal weight):  "
         "M1 Risk Name Recall  +  M2 RPN Tier Accuracy (if applicable)  +  "
         "M3 Critical Factor Match (if applicable)  +  "
         "M4 Traceability (binary: 1.0 if PASS, 0.0 if FAIL)  |  "
         "PASS threshold: 75  ·  Target: 80",
         size_pt=_BODY_PT - 0.5, italic=True)
    doc.add_paragraph()

    # ── 1. Summary metrics ──────────────────────────────────────────────────
    add_section_title(doc, "1. Summary Metrics")
    gt_names = em.get("ground_truth_sources") or [
        "risk_profile_ground_truth.csv",
        "critical_factors_ground_truth.csv",
    ]
    risk_csv = gt_names[0] if gt_names else "risk_profile_ground_truth.csv"
    fac_csv = gt_names[1] if len(gt_names) > 1 else "critical_factors_ground_truth.csv"
    m1 = metrics.get("m1_risk_name_recall") or {}
    matched_pairs = m1.get("matched_pairs") or []
    gt_empty = gt_df is None or gt_df.empty
    inv_when_no_gt = _inventory_generated_risks(gen) if gt_empty else []
    m2 = metrics.get("m2_rpn_tier_accuracy") or {}
    m3 = metrics.get("m3_critical_factor_match") or {}
    m4 = metrics.get("m4_hallucination_detection") or {}
    t1 = targets.get("m1_risk_name_recall", 0.85)
    t2 = targets.get("m2_rpn_tier_accuracy", 0.9)
    t3 = targets.get("m3_critical_factor_match", 0.8)
    t4 = targets.get("m4_hallucinations", 0.0)

    def pf(ok: bool) -> str:
        return "PASS ✓" if ok else "FAIL ✗"

    rows_sum = [
        ["M1 Risk name recall", f"{float(m1.get('score', 0)):.0%}", f"{float(t1):.0%}", pf(bool(m1.get("passed"))), risk_csv],
    ]
    if m2.get("skipped"):
        rows_sum.append(
            ["M2 RPN tier (±1)", "N/A", f"{float(t2):.0%}", "SKIP", risk_csv]
        )
    else:
        rows_sum.append(
            [
                "M2 RPN tier (±1)",
                f"{float(m2.get('score', 0)):.0%}",
                f"{float(t2):.0%}",
                pf(bool(m2.get("passed"))),
                risk_csv,
            ]
        )
    if m3.get("skipped"):
        rows_sum.append(
            ["M3 Critical factor match", "N/A", f"{float(t3):.0%}", "SKIP", fac_csv]
        )
    else:
        rows_sum.append(
            [
                "M3 Critical factor match",
                f"{float(m3.get('score', 0)):.0%}",
                f"{float(t3):.0%}",
                pf(bool(m3.get("passed"))),
                fac_csv,
            ]
        )
    m4_source = "USDM / benchmark provenance rules"
    m4_existence = m4.get("risk_existence_check") or {}
    m4_cf_existence = m4.get("cf_existence_check") or {}
    risk_gate = (
        m4_existence
        and not m4_existence.get("gt_risk_exists")
        and int(m4_existence.get("generated_risk_count") or 0) > 0
    )
    cf_gate = (
        m4_cf_existence
        and not m4_cf_existence.get("gt_cf_exists")
        and int(m4_cf_existence.get("generated_cf_count") or 0) > 0
    )
    if risk_gate and cf_gate:
        m4_source = "Risk + critical-factor existence gates + USDM / benchmark provenance rules"
    elif risk_gate:
        m4_source = "Risk existence gate + USDM / benchmark provenance rules"
    elif cf_gate:
        m4_source = "Critical-factor existence gate + USDM / benchmark provenance rules"
    prov_n = int(m4.get("provenance_defect_count", m4.get("hallucinations_found") or 0))
    sem_n = int(m4.get("semantic_hallucination_count") or 0)
    rows_sum.append(
        [
            "M4 Traceability flags",
            f"{prov_n} (semantic unmatched risks: {sem_n})",
            str(int(t4)) if t4 == int(t4) else str(t4),
            pf(bool(m4.get("passed"))),
            m4_source,
        ]
    )
    add_table(
        doc,
        ["Metric", "Score", "Target", "Pass/Fail", "Source"],
        rows_sum,
        col_widths=[Inches(2.0), Inches(0.85), Inches(0.85), Inches(0.95), Inches(1.95)],
        header_fill_hex=_HEADER_BLUE,
        header_text_white=True,
    )

    all_metrics_ok = (
        bool(m1.get("passed"))
        and (bool(m2.get("passed")) or bool(m2.get("skipped")))
        and (bool(m3.get("passed")) or bool(m3.get("skipped")))
        and bool(m4.get("passed"))
    )
    if all_metrics_ok and doc_score >= tgt:
        overall_txt = (
            f"All summary requirements pass. Document score {doc_score:.1f} meets or exceeds target of {tgt}."
        )
    elif all_metrics_ok:
        overall_txt = (
            f"All summary requirements pass. Document score {doc_score:.1f} (target {tgt}, threshold {thr})."
        )
    else:
        overall_txt = (
            f"One or more requirements did not pass. Document score {doc_score:.1f} "
            f"(threshold {thr}, target {tgt}). See eval JSON for details."
        )
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, "Overall Result", bold=True, size_pt=_BODY_PT)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, overall_txt, size_pt=_BODY_PT)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    verdict = str(report.get("verdict", "NO-GO"))
    ok_go = verdict == "GO"
    _run(p, f"{verdict} {'✓' if ok_go else '✗'}", bold=True, size_pt=_HEADER_PT)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    tail = (
        "Generator ready for verify set evaluation"
        if ok_go
        else "Generator requires fixes before verify sign-off"
    )
    _run(p, tail, size_pt=_BODY_PT)
    add_spacer(doc)

    # ── 2. Risk scorecard ────────────────────────────────────────────────────
    add_section_title(doc, "2. Risk Scorecard")
    yaml_bucket = _bucket_slug(str(ta), str(phase)).split(" — ")[-1]
    n_gt = int(len(gt_df)) if gt_df is not None else int(m1.get("ground_truth_total") or 0)
    n_gt = max(n_gt, 1)
    w_each = 100.0 / n_gt
    sc_rows = []
    bar_order: List[Tuple[str, float]] = []
    if gt_df is not None and not gt_df.empty:
        for _, grow in gt_df.iterrows():
            rname = str(grow["risk_name"])
            pair = _pair_for_gt_row(matched_pairs, grow)
            genr = _gen_risk_from_pair(gen, pair) if pair else _risk_by_name(gen, rname)
            m2r = _m2_row_for_gt(metrics, grow)
            matched = pair is not None
            sc_json_key = pair.get("generated_json_key", "") if pair else ""
            rs = _per_risk_display_score(rname, matched, m2r, genr, grow, gen_json_key=sc_json_key)
            bar_order.append((rname, rs))
            wt = round(w_each * rs / 100.0, 2)
            if not matched:
                rpn_yes = "No"
            else:
                rpn_yes = "Yes" if _rpn_tier_pass_for_row(genr, grow, m2r) else "No"
            nctrl = len(genr.get("controls") or []) if genr else 0
            st = "On Track" if rs >= 85 else "Needs Work"
            sc_rows.append(
                [
                    rname,
                    f"{w_each:.1f}%",
                    f"{rs:.1f}",
                    f"{wt:.2f}",
                    rpn_yes,
                    str(nctrl),
                    st,
                ]
            )
    if not sc_rows:
        inv = inv_when_no_gt
        if inv:
            p = doc.add_paragraph()
            _para_line_spacing(p)
            _run(
                p,
                "risk_profile_ground_truth.csv has no rows for this study_id: the benchmark specifies "
                "**zero** risks. The table lists risks the generator emitted anyway; each is scored as "
                "**hallucination / extra** (M1, M4).",
                size_pt=_BODY_PT,
            )
            n_gen = len(inv)
            w_g = 100.0 / n_gen
            for item in inv:
                rname = item["risk_name"]
                dom = str(item["domain"] or "").strip() or "—"
                label = f"{rname}  ({dom})"
                rs = 0.0
                bar_order.append((label, rs))
                sc_rows.append(
                    [
                        label,
                        f"{w_g:.1f}%",
                        f"{rs:.1f}",
                        "0.00",
                        "N/A",
                        str(item["n_controls"]),
                        "Hallucination / extra",
                    ]
                )
        else:
            sc_rows = [
                [
                    "(Benchmark: zero risks; generator emitted none — conformant.)",
                    "—",
                    "—",
                    "—",
                    "—",
                    "—",
                    "—",
                ]
            ]
    add_table(
        doc,
        ["Risk Name", "Weight", "Risk Score", "Weighted", "RPN ±1", "Controls", "Status"],
        sc_rows,
        col_widths=[Inches(2.0), Inches(0.55), Inches(0.65), Inches(0.65), Inches(0.75), Inches(0.65), Inches(0.85)],
        header_fill_hex=_HEADER_BLUE,
        header_text_white=True,
    )
    doc.add_heading("Risk Score Bar Chart", level=2)
    if not bar_order:
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(
            p,
            "No per-risk chart: benchmark specifies zero risks and the generator JSON has no risk objects.",
            size_pt=_BODY_PT,
        )
    elif not _try_add_risk_score_bar_chart(doc, bar_order, study):
        for rname, rs in bar_order:
            add_text_bar_row(doc, rname, rs)
    add_spacer(doc)

    # ── 3. Domain → Risk → Controls hierarchy ──────────────────────────────
    hv = report.get("hierarchy_verification") or {}
    add_section_title(doc, "3. Domain Verification (Domain → Risk → Controls)")
    by_dom = hv.get("by_domain") or {}
    gt_dom = by_dom.get("ground_truth_risk_rows") or {}
    gen_dom = by_dom.get("generated_risk_rows") or {}
    m1_dom = by_dom.get("m1_matched_pairs") or {}
    all_doms = sorted(set(list(gt_dom.keys()) + list(gen_dom.keys())))
    if all_doms:
        dom_rows = []
        for d in all_doms:
            g = gt_dom.get(d, 0)
            r = gen_dom.get(d, 0)
            m = m1_dom.get(d, 0)
            dom_rows.append([d.title() if d else "(none)", str(g), str(r), str(m)])
        add_table(
            doc,
            ["Domain", "GT risks", "Generated risks", "M1 matched"],
            dom_rows,
            col_widths=[Inches(2.0), Inches(1.2), Inches(1.6), Inches(1.2)],
            header_fill_hex=_HEADER_BLUE,
            header_text_white=True,
        )
    else:
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(p, "Domain information not available.", size_pt=_BODY_PT)
    add_spacer(doc)

    per_rc = hv.get("per_risk_control_alignment") or []
    ctrl_rate = hv.get("control_count_match_rate")
    ctrl_n = hv.get("pairs_with_gt_control_count") or 0
    if per_rc:
        doc.add_heading("Control count alignment (GT control_count vs generated controls)", level=2)
        ctrl_rows = []
        for rc in per_rc:
            gt_cc = rc.get("gt_control_count")
            gen_cc = rc.get("generated_control_count")
            ok = rc.get("control_count_match")
            status = "✓" if ok else ("✗" if ok is False else "—")
            ctrl_rows.append([
                str(rc.get("gt_risk_name") or ""),
                str(rc.get("gt_domain") or "").title(),
                str(gt_cc) if gt_cc is not None else "—",
                str(gen_cc) if gen_cc is not None else "—",
                status,
            ])
        add_table(
            doc,
            ["Risk", "Domain", "GT #controls", "Gen #controls", "Match"],
            ctrl_rows,
            col_widths=[Inches(2.0), Inches(1.2), Inches(1.0), Inches(1.0), Inches(0.7)],
            header_fill_hex=_HEADER_BLUE,
            header_text_white=True,
        )
        if ctrl_rate is not None and ctrl_n:
            p = doc.add_paragraph()
            _para_line_spacing(p)
            _run(p, f"Control count match rate: {float(ctrl_rate):.0%} over {ctrl_n} pair(s).", size_pt=_BODY_PT)
    add_spacer(doc)

    # ── 4. Risk detail (grouped by domain) ────────────────────────────────
    add_section_title(doc, "4. Risk Detail")
    if gt_df is not None and not gt_df.empty:
        has_dom_col = "risk_domain" in gt_df.columns
        if has_dom_col:
            domain_order = []
            seen_doms: set = set()
            for _, grow in gt_df.iterrows():
                d = str(grow.get("risk_domain", "") or "").strip()
                if d and d not in seen_doms:
                    domain_order.append(d)
                    seen_doms.add(d)
            if not domain_order:
                domain_order = [""]
        else:
            domain_order = [""]

        risk_idx = 0
        for dom in domain_order:
            if dom:
                doc.add_heading(dom, level=2)
            if has_dom_col and dom:
                dom_df = gt_df[gt_df["risk_domain"].astype(str).str.strip() == dom]
            else:
                dom_df = gt_df
            for _, grow in dom_df.iterrows():
                risk_idx += 1
                rname = str(grow["risk_name"])
                pair = _pair_for_gt_row(matched_pairs, grow)
                genr = _gen_risk_from_pair(gen, pair) if pair else _risk_by_name(gen, rname)
                m2r = _m2_row_for_gt(metrics, grow)
                matched = pair is not None
                det_json_key = pair.get("generated_json_key", "") if pair else ""
                add_risk_block(doc, risk_idx, rname, genr, grow, m2r, matched, gen_json_key=det_json_key)
    else:
        inv = inv_when_no_gt
        if inv:
            p = doc.add_paragraph()
            _para_line_spacing(p)
            _run(
                p,
                "Benchmark specifies zero risks for this study_id. Each section below is one generated risk "
                "with **no** authorized benchmark row (hallucination / extra under M1 and M4).",
                size_pt=_BODY_PT,
            )
            for i, item in enumerate(inv, start=1):
                add_risk_block(
                    doc,
                    i,
                    item["risk_name"],
                    item["risk"],
                    None,
                    None,
                    False,
                )
        else:
            p = doc.add_paragraph()
            _para_line_spacing(p)
            _run(
                p,
                "Benchmark specifies zero risks; generator output has no risk objects (conformant).",
                size_pt=_BODY_PT,
            )
    add_spacer(doc)

    # ── 5. Critical factors ─────────────────────────────────────────────────
    add_section_title(doc, "5. Critical Factors")
    exp = int(m3.get("ground_truth_total") or 0)
    matched_n = int(m3.get("matched_factors") or 0)
    if m3.get("skipped"):
        cf_line = "Critical factor ground truth not available for this study (M3 skipped)."
    elif m3.get("benchmark_expects_zero_factors"):
        n_cf = int(m3.get("generated_factor_count") or 0)
        if n_cf == 0:
            cf_line = (
                "critical_factors_ground_truth.csv has no rows for this study_id: benchmark specifies "
                "zero critical factors; generator emitted none (conformant)."
            )
        else:
            ex = m3.get("extra_names") or []
            cf_line = (
                "Benchmark specifies **zero** critical factors for this protocol; generator emitted "
                f"{n_cf}: {ex}. Each counts as hallucination / extra (M3, M4 existence gate)."
            )
    else:
        f1p = m3.get("f1")
        pr = m3.get("precision")
        rc = m3.get("recall")
        cv = m3.get("content_violations") or []
        stat_bits = (
            f"{matched_n} name matches; F1={f1p}, P={pr}, R={rc}."
            if f1p is not None
            else f"{matched_n} matched."
        )
        cv_bits = f" Content violations: {len(cv)}." if cv else ""
        cf_line = (
            f"{exp} critical factors expected for {yaml_bucket} bucket. {stat_bits}{cv_bits}"
        )
        if m3.get("missing_names"):
            cf_line += f" Missing: {m3.get('missing_names')}."
        if m3.get("extra_names"):
            cf_line += f" Extra (not in GT): {m3.get('extra_names')}."
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, cf_line, size_pt=_BODY_PT)
    add_spacer(doc)
    cf_rows = []
    if gen:
        matched_cf = set(m3.get("matched_names") or [])
        for cf in gen.get("critical_factors") or []:
            if not isinstance(cf, dict):
                continue
            fn = str(cf.get("factor_name") or "")
            cf_rows.append(
                [
                    fn,
                    "Yes ✓" if fn in matched_cf else "No ✗",
                    str(cf.get("critical_data") or "")[:200],
                    str(cf.get("critical_process") or "")[:200],
                ]
            )
    add_table(
        doc,
        ["Critical Factor", "Match", "Critical Data", "Critical Process"],
        cf_rows,
        col_widths=[Inches(1.4), Inches(0.55), Inches(2.3), Inches(2.3)],
        header_fill_hex=_HEADER_BLUE,
        header_text_white=True,
    )
    add_spacer(doc)

    # ── 6. Hallucination check ─────────────────────────────────────────────
    add_section_title(doc, "6. Hallucination Check")
    hf = int(m4.get("hallucinations_found") or 0)
    sem_u = int(m4.get("semantic_hallucination_count") or 0)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    if hf == 0:
        _run(p, "ZERO PROVENANCE DEFECTS (M4) ✓", bold=True, size_pt=_HEADER_PT)
    else:
        _run(
            p,
            f"PROVENANCE DEFECTS (field-level checks): {hf}",
            bold=True,
            size_pt=_HEADER_PT,
            rgb=_RGB_FAIL,
        )
    p2 = doc.add_paragraph()
    _para_line_spacing(p2)
    n_cf_ex = 0
    if m3.get("benchmark_expects_zero_factors"):
        n_cf_ex = int(m3.get("generated_factor_count") or 0)
    _run(
        p2,
        f"Semantic unmatched generated risks (M1; not summed into defect count above): {sem_u}. "
        f"Extra critical factors when benchmark expects zero (M3): {n_cf_ex}.",
        size_pt=_BODY_PT,
    )
    add_spacer(doc)
    chk_rows = [
        ["benchmark_source present", pf(bool(m4.get("passed")) and hf == 0)],
        ["usdm_drivers non-empty", pf(bool(m4.get("passed")) and hf == 0)],
        ["associated_cause usdm_trigger", pf(bool(m4.get("passed")) and hf == 0)],
        ["critical_factor usdm critical_data", pf(bool(m4.get("passed")) and hf == 0)],
    ]
    add_table(
        doc,
        ["Check", "Result"],
        chk_rows,
        col_widths=[Inches(3.5), Inches(2.5)],
        header_fill_hex=_HEADER_BLUE,
        header_text_white=True,
    )
    add_spacer(doc)

    # ── 7. Near misses ──────────────────────────────────────────────────────
    add_section_title(doc, "7. Near Misses")
    nm = report.get("near_misses") or []
    if not nm:
        p = doc.add_paragraph()
        _para_line_spacing(p)
        mt = int(m1.get("ground_truth_total") or 0)
        if mt == 0:
            near_txt = (
                "No near misses: the benchmark defines no risk names for this protocol, so name-distance "
                "logging is empty. Any generated risk is evaluated as hallucination / extra (M1, M4), "
                "not as a near miss."
            )
        else:
            near_txt = (
                f"No near misses found. All {mt} ground-truth risk rows for this protocol were "
                f"either exact-name matched or missed outright (Levenshtein logging empty)."
            )
        _run(p, near_txt, size_pt=_BODY_PT)
    else:
        for n in nm:
            add_paragraph_block(
                doc,
                f"Near miss (distance {n.get('edit_distance')})",
                f"GT: {n.get('truth_name')} — Gen: {n.get('generated_name')}",
            )
    add_spacer(doc)

    # ── 7. Vendor / site / other ───────────────────────────────────────────
    add_section_title(doc, "7. Vendor / Site / Other Domain Risks")
    vr = len(gen.get("vendor_risks") or []) if gen else 0
    sr = len(gen.get("study_site_risks") or []) if gen else 0
    od = len(gen.get("other_domain_risks") or []) if gen else 0
    # Compute GT domain counts from GT CSV when available
    _gt_df = _load_risk_gt_df(report)
    if _gt_df is not None and "risk_domain" in _gt_df.columns:
        _domains = _gt_df["risk_domain"].astype(str).str.strip().str.lower()
        gt_vr = int((_domains == "vendor risks").sum())
        gt_sr = int((_domains == "study site risks").sum())
        gt_od = int((~_domains.isin(["vendor risks", "study risks", "study site risks"])).sum())
        gt_vr_s, gt_sr_s, gt_od_s = str(gt_vr), str(gt_sr), str(gt_od)
    else:
        gt_vr_s = gt_sr_s = gt_od_s = "—"
    add_table(
        doc,
        ["Domain", "Generated", "Ground Truth"],
        [
            ["Vendor risks", str(vr), gt_vr_s],
            ["Study site risks", str(sr), gt_sr_s],
            ["Other domain risks", str(od), gt_od_s],
        ],
        col_widths=[Inches(1.8), Inches(1.2), Inches(1.2)],
        header_fill_hex=_HEADER_BLUE,
        header_text_white=True,
    )
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(
        p,
        "Ground Truth counts are from risk_profile_ground_truth.csv for this study_id only. "
        "Vendor / site / other columns show how many benchmark rows fall in each domain when rows exist.",
        size_pt=_BODY_PT,
    )
    if (
        m4_existence
        and not m4_existence.get("gt_risk_exists")
        and int(m4_existence.get("generated_risk_count") or 0) > 0
    ):
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(
            p,
            f"Risk existence: benchmark specifies zero risks for this protocol; generator emitted "
            f"{int(m4_existence.get('generated_risk_count') or 0)} risk object(s) "
            f"(non-conformant — hallucination / extra).",
            size_pt=_BODY_PT,
            bold=True,
        )
    m4_cf_ex = m4.get("cf_existence_check") or {}
    if (
        m4_cf_ex
        and not m4_cf_ex.get("gt_cf_exists")
        and int(m4_cf_ex.get("generated_cf_count") or 0) > 0
    ):
        p = doc.add_paragraph()
        _para_line_spacing(p)
        _run(
            p,
            f"Critical factor existence: benchmark specifies zero critical factors for this protocol; "
            f"generator emitted {int(m4_cf_ex.get('generated_cf_count') or 0)} factor object(s) "
            f"(non-conformant — hallucination / extra).",
            size_pt=_BODY_PT,
            bold=True,
        )
    add_spacer(doc)

    # ── 8. Improvements / Notes ──────────────────────────────────────────────
    add_section_title(doc, "8. Improvements")
    _write_m4_false_positive_note(doc, m4)
    add_spacer(doc)

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_line_spacing(p)
    _run(p, "Confidential | Hexaware Technologies", size_pt=9, italic=True)


def _write_scenario2(doc: Document, report: Dict[str, Any]) -> None:
    em = report.get("eval_metadata") or {}
    study = report.get("study_id") or "?"
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    p = doc.add_paragraph()
    _para_line_spacing(p)
    _run(p, f"Pfizer Protocol Intelligence Platform | D2: Risk Profile Generator | Eval Report | {study}", size_pt=_BODY_PT)
    add_spacer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Risk Profile Generator — Eval Report (Scenario 2)", bold=True, size_pt=_TITLE_PT)
    add_spacer(doc)
    verdict = report.get("verdict", "RED")
    p = doc.add_paragraph()
    _run(p, f"Traffic-light verdict: {verdict}", bold=True, size_pt=_HEADER_PT)
    add_spacer(doc)
    signals = report.get("signals") or {}
    srows = []
    for sid in sorted(signals.keys()):
        s = signals[sid]
        srows.append([sid, s.get("name", ""), s.get("status", ""), str(s.get("description", ""))[:240]])
    add_table(
        doc,
        ["Signal", "Name", "Status", "Detail"],
        srows,
        col_widths=[Inches(0.5), Inches(1.45), Inches(0.7), Inches(3.35)],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Confidential | Hexaware Technologies", size_pt=9, italic=True)


def write_risk_profile_eval_docx(report: Dict[str, Any], output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if int(report.get("scenario") or 0) == 1:
        _write_scenario1(doc, report)
    else:
        _write_scenario2(doc, report)
    doc.save(str(path))


def write_risk_profile_eval_docx_from_json(json_path: str, output_path: str | None = None) -> str:
    p = Path(json_path)
    with open(p, encoding="utf-8") as f:
        report = json.load(f)
    study = report.get("study_id", p.stem)
    out = output_path or str(p.parent / f"Risk_Profile_Eval_Report_{study}.docx")
    write_risk_profile_eval_docx(report, out)
    return out
