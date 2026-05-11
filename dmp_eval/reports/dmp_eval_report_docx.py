"""
DMP eval Word report — layout aligned to ``reference_specs/DMP_Eval_Report_B7981027.docx``.

Structure: banner table → document score row (3 cols) → sections **1–7** as Heading 2,
same table shapes and column headers as the reference (only cell text varies by study).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None  # type: ignore

_FONT = "Arial"
# Pfizer reference palette
_HEADER_BG = "003087"             # Pfizer blue — table header backgrounds
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_OK = RGBColor(0x1A, 0x73, 0x40)
_RED_FAIL = RGBColor(0xB9, 0x1C, 0x1C)
_DARK = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY = RGBColor(0x66, 0x66, 0x66)


def _apply_doc_defaults(doc: "Document") -> None:
    """Arial 10pt, narrow margins."""
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)


def _thin_table_borders(table) -> None:
    """Light-gray 0.5 pt borders on all sides."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblB.append(el)
    tblPr.append(tblB)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _banner_table(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    _set_cell_shading(c, _HEADER_BG)
    p = c.paragraphs[0]
    p.text = ""
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = _WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _document_score_row(
    doc: Document,
    *,
    ta_phase_slug: str,
    score: float,
    passed: bool,
    threshold: float,
    target: float,
) -> None:
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    left, mid, right = t.rows[0].cells
    p0 = left.paragraphs[0]
    p0.text = ""
    r0 = p0.add_run(f"DOCUMENT SCORE {ta_phase_slug}")
    r0.bold = True
    r0.font.size = Pt(10)
    pm = mid.paragraphs[0]
    pm.text = ""
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = "PASS ✓" if passed else "FAIL ✗"
    rm_score = pm.add_run(f"{score:.1f}")
    rm_score.bold = True
    rm_score.font.size = Pt(22)
    rm_score.font.color.rgb = _GREEN_OK if passed else _RED_FAIL
    pm2 = mid.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm_label = pm2.add_run(f"/ 100  {st}")
    rm_label.bold = True
    rm_label.font.size = Pt(10)
    rm_label.font.color.rgb = _GREEN_OK if passed else _RED_FAIL
    pr = right.paragraphs[0]
    pr.text = ""
    rr = pr.add_run(f"Threshold: {threshold:.0f} | Target: {target:.0f}")
    rr.font.size = Pt(10)
    doc.add_paragraph()


def _add_heading2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _data_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    *,
    header_shade: str = _HEADER_BG,
) -> None:
    if not headers:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    _thin_table_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.name = _FONT
        r.font.size = Pt(9)
        r.font.color.rgb = _WHITE
        _set_cell_shading(c, header_shade)
    for i, row in enumerate(rows):
        for j in range(len(headers)):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            val = str(row[j]) if j < len(row) else ""
            rn = cell.paragraphs[0].add_run(val)
            rn.font.name = _FONT
            rn.font.size = Pt(9)
            # Colorize PASS/FAIL cells
            if "PASS" in val and "FAIL" not in val:
                rn.font.color.rgb = _GREEN_OK
                rn.bold = True
            elif "FAIL" in val:
                rn.font.color.rgb = _RED_FAIL
                rn.bold = True
    doc.add_paragraph()


def _append_s8_score_note(doc: Document, s8_rows_raw: List[Dict[str, Any]]) -> None:
    """
    Add an explanatory note when S8 rows have matched module names but reduced scores.
    Most commonly this is due to layer metadata mismatch (e.g., generated L1 vs expected L2).
    """
    reduced_due_layer = 0
    reduced_other = 0
    for r in s8_rows_raw:
        if not isinstance(r, dict):
            continue
        if str(r.get("match_status") or "") not in ("verbatim", "near_miss"):
            continue
        try:
            item_score = float(r.get("item_score") or 0)
        except (TypeError, ValueError):
            item_score = 0.0
        if item_score >= 99.99:
            continue

        atr = r.get("attributes") or {}
        layer_sc = float(((atr.get("s8_layer") or {}).get("score")) or 0)
        if layer_sc < 1.0:
            reduced_due_layer += 1
        else:
            reduced_other += 1

    if reduced_due_layer == 0 and reduced_other == 0:
        return

    p = doc.add_paragraph()
    run = p.add_run("Score note: ")
    run.bold = True
    run.font.size = Pt(9)

    if reduced_due_layer > 0:
        txt = (
            f"{reduced_due_layer} matched S8 row(s) scored below 100 due to "
            "metadata mismatch in `s8_layer` (for example generated L1 vs expected L2), "
            "not because module names were wrong."
        )
    else:
        txt = "Some matched S8 rows scored below 100 due to non-name attribute mismatches."
    r2 = p.add_run(txt + (f" Additional rows with other non-name deductions: {reduced_other}." if reduced_other else ""))
    r2.font.size = Pt(9)
    r2.italic = True


def write_dmp_eval_docx(output_path: Path, report: Dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError("python-docx required: pip install python-docx")

    meta = report.get("eval_metadata") or {}
    summ = report.get("summary_metrics") or {}
    sid = str(meta.get("study_id", "") or report.get("study_id", ""))

    doc = Document()
    doc.core_properties.title = f"DMP Eval Report — {sid}"
    _apply_doc_defaults(doc)

    ta = str(meta.get("therapeutic_area") or "—")
    ph = str(meta.get("phase") or "—")
    ta_slug = (meta.get("ta_slug") or meta.get("config_label") or "").strip()
    if ta_slug:
        ta_phase = f"{ta} / {ph} — {ta_slug}"
    else:
        ta_phase = f"{ta} / {ph}"

    _banner_table(doc, f"PFIZER PROTOCOL INTELLIGENCE PLATFORM DMP Generator — Eval Report {sid}")

    ds = float(report.get("document_score", 0) or 0)
    dp = bool(report.get("document_pass"))
    thr = float(report.get("document_pass_threshold", 75) or 75)
    tgt = float(report.get("document_target", 80) or 80)
    _document_score_row(
        doc,
        ta_phase_slug=ta_phase,
        score=ds,
        passed=dp,
        threshold=thr,
        target=tgt,
    )

    # Score formula explanation
    sf_tbl = doc.add_table(rows=1, cols=1)
    sf_tbl.style = "Table Grid"
    _set_cell_shading(sf_tbl.rows[0].cells[0], "EBF0FB")
    sf_p = sf_tbl.rows[0].cells[0].paragraphs[0]
    sf_p.text = ""
    r = sf_p.add_run("How the Document Score is Calculated")
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(10)
    sf_p2 = sf_tbl.rows[0].cells[0].add_paragraph()
    r2 = sf_p2.add_run(
        "Weighted section average:  "
        "M1 S5 System Accuracy (25%)  +  M2 S6.2 Vendor Recall (25%)  +  "
        "M3 S8 Module Recall (40%)  +  M4 S11.4 Reconciliation (10%)  "
        "=  Document Score (0–100)  |  Pass threshold: 75  ·  Target: 80"
    )
    r2.italic = True
    r2.font.name = _FONT
    r2.font.size = Pt(9)
    doc.add_paragraph()

    _add_heading2(doc, "1. Summary Metrics")
    _data_table(
        doc,
        ["Metric", "Score", "Target", "Pass/Fail", "Source"],
        [
            ["M1 S5 System Accuracy", f"{summ.get('m1_s5_system_accuracy', 0):.1%}", f"{summ.get('m1_target', 0):.0%}", "PASS ✓" if summ.get("m1_pass") else "FAIL ✗", "dmp_ground_truth S5"],
            ["M2 S6.2 Vendor Recall", f"{summ.get('m2_s6_vendor_recall', 0):.1%}", f"{summ.get('m2_target', 0):.0%}", "PASS ✓" if summ.get("m2_pass") else "FAIL ✗", "SDS Non-CRF CSV / JSON fallback"],
            ["M3 S8 Module Recall", f"{summ.get('m3_s8_module_recall', 0):.1%}", f"{summ.get('m3_target', 0):.0%}", "PASS ✓" if summ.get("m3_pass") else "FAIL ✗", "dmp_ground_truth S8"],
            ["M4 S11.4 Flags", f"{summ.get('m4_reconciliation_accuracy', 0):.1%}", f"{summ.get('m4_target', 0):.0%}", "PASS ✓" if summ.get("m4_pass") else "FAIL ✗", "dmp_ground_truth S11"],
            ["M4b Hallucinations (null source_tag)", str(summ.get("m4_hallucinations", 0)), str(summ.get("m4_hallucination_target", 0)), "PASS ✓" if summ.get("m4_hallucination_pass") else "FAIL ✗", "S5/S6/S8 traceability"],
        ],
    )

    _add_heading2(doc, "2. Section Scorecard")
    sc = report.get("section_scores") or {}
    sr: List[List[str]] = []
    order = ["s5_systems", "s8_critical_data", "s6_vendors", "s11_reconciliation"]
    labels = {
        "s5_systems": "S5 Systems",
        "s8_critical_data": "S8 Critical Data",
        "s6_vendors": "S6.2 Vendors",
        "s11_reconciliation": "S11.4 Reconciliation",
    }
    metrics_lbl = {
        "s5_systems": "M1",
        "s8_critical_data": "M3",
        "s6_vendors": "M2",
        "s11_reconciliation": "M4",
    }
    for key in order:
        blk = sc.get(key) or {}
        st = "PASS ✓" if (
            (key == "s5_systems" and summ.get("m1_pass"))
            or (key == "s6_vendors" and summ.get("m2_pass"))
            or (key == "s8_critical_data" and summ.get("m3_pass"))
            or (key == "s11_reconciliation" and summ.get("m4_pass"))
        ) else "FAIL ✗"
        sr.append(
            [
                labels.get(key, key),
                metrics_lbl.get(key, "—"),
                f"{float(blk.get('weight', 0)):.2f}",
                str(blk.get("score", "")),
                str(blk.get("weighted_contribution", "")),
                st,
            ],
        )
    _data_table(doc, ["Section", "Metric", "Weight", "Score", "Weighted", "Status"], sr)

    _add_heading2(doc, "3. S5 Systems & Tools Detail")
    s5 = report.get("s5_systems") or []
    s5_rows: List[List[str]] = []
    for r in s5:
        if not isinstance(r, dict):
            continue
        ms = r.get("match_status", "")
        icon = "✓" if ms == "verbatim" else ("~" if ms == "near_miss" else "✗")
        tier = str(r.get("attributes", {}).get("s5_confidence_tier", {}).get("generated", "") or "")
        note = str(r.get("ground_truth_name", ""))[:200]
        s5_rows.append(
            [
                str(r.get("system_type", "")),
                str(r.get("generated_name") or "—"),
                str(r.get("item_score", "")),
                f"{icon} {ms}",
                tier,
                note,
            ]
        )
    _data_table(doc, ["System Type", "Generated", "Score", "Match", "Tier", "Notes"], s5_rows[:80])

    _add_heading2(doc, "4. S8 Critical Data Detail")
    s8 = report.get("s8_critical_data") or []
    s8_rows: List[List[str]] = []
    for r in s8:
        if not isinstance(r, dict):
            continue
        ms = r.get("match_status", "")
        icon = "✓" if ms == "verbatim" else ("~" if ms == "near_miss" else "✗")
        if ms == "extra":
            icon = "✗"
        atr = r.get("attributes") or {}
        tier = str((atr.get("s8_tier") or {}).get("ground_truth", "") or "")[:80]
        module_cell = str(r.get("generated_module") or "—")
        layer_cell = str(r.get("layer") or "—")
        notes = str(r.get("note", ""))[:220]
        # Explain reduced scores even when module text is matched (e.g., 90 due to layer mismatch).
        atr_layer = (atr.get("s8_layer") or {})
        try:
            item_score = float(r.get("item_score") or 0)
        except (TypeError, ValueError):
            item_score = 0.0
        if ms in ("verbatim", "near_miss") and item_score < 100:
            lay_sc = float(atr_layer.get("score") or 0)
            if lay_sc < 1.0:
                g = str(atr_layer.get("generated") or "—")
                e = str(atr_layer.get("expected") or "—")
                why = f"Score reduced: layer mismatch (generated {g}, expected {e})."
                notes = f"{notes} {why}".strip()
        if r.get("ground_truth_module"):
            gt_hint = str(r.get("ground_truth_module"))
            if notes:
                notes = f"GT: {gt_hint} | {notes}"
            else:
                notes = f"GT: {gt_hint}"
        s8_rows.append(
            [
                module_cell,
                layer_cell,
                str(r.get("item_score", "")),
                f"{icon} {ms}",
                tier,
                notes,
            ]
        )
    _data_table(doc, ["Module", "Layer", "Score", "Match", "Tier", "Notes"], s8_rows[:80])
    _append_s8_score_note(doc, s8)

    _add_heading2(doc, "5. S6.2 eSource/eData Vendor Table Detail")
    s6 = report.get("s6_vendors") or []
    s6_rows: List[List[str]] = []
    for r in s6:
        if not isinstance(r, dict) or r.get("item_score") == "excluded":
            continue
        atr = r.get("attributes") or {}
        dt = str(atr.get("s6_data_type", {}).get("ground_truth", "") or "")[:80]
        t_det = atr.get("s6_data_review_tier", {}) or {}
        t_gt_raw = str(t_det.get("ground_truth") or "").strip()
        t_gen_raw = str(t_det.get("generated") or "").strip()
        t_gt = t_gt_raw or "Not provided"
        t_gen = t_gen_raw or "Not provided"
        tier = f"GT: {t_gt} | Gen: {t_gen}"
        ms = str(r.get("match_status", "") or "")
        notes = ms

        # Row-level deduction reason (same style as S8 notes).
        if ms in ("mismatch", "near_miss"):
            reasons = []
            v_det = atr.get("s6_vendor_name", {}) or {}
            d_det = atr.get("s6_data_type", {}) or {}
            s_det = atr.get("s6_source_tag", {}) or {}

            if float(v_det.get("score") or 0) < 1.0:
                gv = str(r.get("ground_truth_vendor") or "—")
                gn = str(r.get("generated_vendor") or "—")
                reasons.append(f"vendor mismatch (GT {gv}, Gen {gn})")
            if float(d_det.get("score") or 0) < 1.0:
                d_gt = str(d_det.get("ground_truth") or "—")
                d_gn = str(d_det.get("generated") or "—")
                reasons.append(f"data type mismatch (GT {d_gt}, Gen {d_gn})")
            if float(t_det.get("score") or 0) < 1.0:
                reasons.append(f"tier mismatch ({tier})")
            if float(s_det.get("score") or 0) < 1.0:
                reasons.append("source tag missing/invalid")

            if reasons:
                notes = f"{ms} | Score reduced: " + "; ".join(reasons) + "."
        elif not t_gt_raw and not t_gen_raw:
            notes = f"{ms} | Tier not provided in source for this row; treated as neutral match."

        s6_rows.append(
            [
                str(r.get("ground_truth_vendor") or "—"),
                str(r.get("generated_vendor") or "—"),
                dt,
                tier,
                str(r.get("item_score", "")),
                notes,
            ]
        )
    _data_table(
        doc,
        ["Ground Truth Vendor", "Generated Vendor", "Data Type", "Tier", "Score", "Notes"],
        s6_rows[:80],
    )

    _add_heading2(doc, "6. S11.4 Reconciliation Flags Detail")
    s11 = report.get("s11_reconciliation") or []
    s11_rows: List[List[str]] = []
    for r in s11:
        if not isinstance(r, dict):
            continue
        atr = r.get("attributes") or {}
        s11_rows.append(
            [
                str(r.get("subsection_label", "")),
                str(r.get("ground_truth_flag", "")),
                str(r.get("generated_flag", "")),
                str(r.get("item_score", "")),
                str((atr.get("s11_inference_rule") or {}).get("rule") or ""),
            ]
        )
    _data_table(doc, ["Subsection", "Ground Truth", "Generated", "Score", "Inference Rule"], s11_rows)

    _add_heading2(doc, "7. Improvement Actions")
    ia = report.get("improvement_actions") or []
    ia_rows: List[List[str]] = []
    for x in ia:
        if not isinstance(x, dict):
            continue
        ia_rows.append(
            [
                str(x.get("priority", "")),
                str(x.get("section", "")),
                str(x.get("type", "")),
                str(x.get("action", ""))[:400],
                str(x.get("fix_location", ""))[:240],
            ]
        )
    _data_table(doc, ["Priority", "Section", "Type", "Action", "Fix Location"], ia_rows)

    em = report.get("eval_metadata") or {}
    oos = em.get("out_of_scope_sections_excluded") or []
    oos_note = (
        "Out of scope — not penalised: "
        + (", ".join(oos) if oos else "S1 (Revision History), S3 (DM Process Documentation)")
        + "."
    )
    fp = doc.add_paragraph()
    r = fp.add_run(oos_note)
    r.font.size = Pt(9)
    r.italic = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
    except PermissionError:
        alt = output_path.parent / f"{output_path.stem}_{datetime.now():%Y%m%d_%H%M%S}{output_path.suffix}"
        doc.save(str(alt))
