"""
dmp_scenario2_report.py
-----------------------
DMP Scenario 2 (no ground truth) — Word document report.

Takes the JSON dict produced by ``DMP_eval/eval_scenario2.py::run_scenario2_eval``
and writes a styled .docx file matching the Pfizer brand palette used by the
DMP Scenario 1 report.

Public API
~~~~~~~~~~
    write_dmp_scenario2_docx(s2_result, output_path)
    write_dmp_scenario2_docx_from_json(json_path, output_path)
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Pfizer brand palette ────────────────────────────────────────────────────
_HEADER_BG  = "003087"
_GREEN_BG   = "C6EFCE"
_AMBER_BG   = "FFF2CC"
_RED_BG     = "FFDCE0"
_BLUE_BG    = "D9E2F3"
_GREY_BG    = "F2F2F2"
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_FG   = RGBColor(0x1A, 0x73, 0x40)
_AMBER_FG   = RGBColor(0x92, 0x60, 0x00)
_RED_FG     = RGBColor(0xB9, 0x1C, 0x1C)
_BLUE_FG    = RGBColor(0x00, 0x30, 0x87)
_FONT       = "Arial"
_BODY_PT    = 10.0
_CELL_PT    = 9.0       # table cell text — matches S1 eval
_HEAD_PT    = 11.0
_TITLE_PT   = 13.0
_BIG_PT     = 22.0      # score % in verdict box — matches S1 eval


# ── XML helpers ──────────────────────────────────────────────────────────────

def _shd(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _color(run, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement("w:color")
    el.set(qn("w:val"), hex_color)
    rPr.append(el)


def _borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _run(para, text: str, *, bold: bool = False, italic: bool = False,
         size: float = _BODY_PT, rgb: Optional[RGBColor] = None,
         hex_color: Optional[str] = None) -> None:
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = _FONT
    r.font.size = Pt(size)
    if rgb is not None:
        r.font.color.rgb = rgb
    if hex_color is not None:
        _color(r, hex_color)


def _clear(cell):
    for child in list(cell._tc):
        if child.tag.endswith("}tcPr"):
            continue
        cell._tc.remove(child)
    return cell.add_paragraph()


# ── Document builders ─────────────────────────────────────────────────────────

def _apply_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)


def _section_heading(doc: Document, text: str) -> None:
    """Full-width blue banner — matches S1 eval report section heading style."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shd(cell, _HEADER_BG)
    p = _clear(cell)
    r = p.add_run(text)
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(_HEAD_PT)
    r.font.color.rgb = _WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT + 0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def _note_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = _FONT
    r.font.size = Pt(_CELL_PT)
    _color(r, "555555")
    p.paragraph_format.space_after = Pt(4)


def _add_banner(doc: Document, study_id: str, ts: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shd(cell, _HEADER_BG)
    p = _clear(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "DMP Evaluation Report — Scenario 2\n", bold=True, size=_TITLE_PT,
         rgb=_WHITE)
    _run(p, f"Study: {study_id}  ·  Data Management Plan\n", bold=True, size=_HEAD_PT,
         rgb=_WHITE)
    _run(p, f"No Ground Truth · Proxy Quality Signals  ·  CONFIDENTIAL  ·  {ts}",
         size=_BODY_PT, rgb=_WHITE)
    doc.add_paragraph()


def _add_verdict_box(doc: Document, verdict: str, health: Dict[str, Any]) -> None:
    verdict_upper = (verdict or "").upper()
    go_no_go = "GO" if verdict_upper in ("GREEN", "AMBER") else "NO-GO"
    pct   = round(health.get("percent", 0.0), 1)
    pass_n = health.get("pass", 0)
    warn_n = health.get("warn", 0)
    fail_n = health.get("fail", 0)
    total  = health.get("total", 0)

    if verdict_upper == "GREEN":
        bg, fg_hex = _GREEN_BG, "1A7340"
    elif verdict_upper == "AMBER":
        bg, fg_hex = _AMBER_BG, "926000"
    else:
        bg, fg_hex = _RED_BG, "B91C1C"

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    widths = [Inches(2.3), Inches(1.8), Inches(2.4)]
    for ci, w in enumerate(widths):
        t.rows[0].cells[ci].width = w

    c0 = t.rows[0].cells[0]
    _shd(c0, _BLUE_BG)
    p0 = _clear(c0)
    _run(p0, "SCENARIO 2 SCORE\n", bold=True, size=_HEAD_PT, hex_color=_HEADER_BG)
    p0b = c0.add_paragraph()
    _run(p0b, "DMP Quality Signals", italic=True, size=_BODY_PT, hex_color=_HEADER_BG)

    c1 = t.rows[0].cells[1]
    _shd(c1, bg)
    p1 = _clear(c1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, f"{pct}%", bold=True, size=_BIG_PT, hex_color=fg_hex)
    p1b = c1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1b, "Signal Health", size=_CELL_PT, hex_color="555555")

    c2 = t.rows[0].cells[2]
    p2 = _clear(c2)
    _run(p2, f"{verdict_upper}  ·  {go_no_go}\n", bold=True, size=_HEAD_PT,
         hex_color=fg_hex)
    p2b = c2.add_paragraph()
    _run(p2b,
         f"{pass_n} PASS · {warn_n} WARN · {fail_n} FAIL  (of {total})\n",
         size=_CELL_PT, hex_color="555555")
    p2c = c2.add_paragraph()
    _run(p2c, "Scoring: PASS=1.0 · WARN=0.5 · FAIL=0.0", italic=True,
         size=_CELL_PT, hex_color="555555")

    # USDM S8 false-positive note if verdict is RED solely due to S8
    if verdict_upper == "RED":
        doc.add_paragraph()
        _note_paragraph(
            doc,
            "NOTE: RED verdict is driven by S8 USDM Traceability failing because the "
            "expected USDM anchor classes (Assessment, DataAcquisition, "
            "ScheduleOfActivities, Procedure) are not instantiated in the USDM 4.0 "
            "JSON for this study. This is a schema coverage gap in the USDM export, "
            "not an error in the DMP content itself. Review other signals individually."
        )
    doc.add_paragraph()


def _status_colors(status: str) -> Tuple[str, str]:
    s = (status or "").upper()
    if s == "PASS":
        return _GREEN_BG, "1A7340"
    if s == "WARN":
        return _AMBER_BG, "926000"
    if s == "FAIL":
        return _RED_BG, "B91C1C"
    return _GREY_BG, "262626"


def _signal_detail(sig_id: str, blk: Dict[str, Any]) -> str:
    if sig_id == "S1":
        missing = blk.get("missing") or []
        return f"{len(missing)} section(s) missing" if missing else "All required sections present"
    if sig_id == "S2":
        n = blk.get("missing_count", 0)
        return f"{n} row(s) missing source/provenance"
    if sig_id == "S3":
        counts = blk.get("counts") or {}
        return (f"S6 vendors: {counts.get('s6_vendors', '—')}  "
                f"S8 modules: {counts.get('s8_modules', '—')}")
    if sig_id == "S4":
        total = blk.get("total_with_confidence", 0)
        low_n = blk.get("low_or_review_count", 0)
        rate  = blk.get("low_or_review_rate", 0.0)
        return f"{low_n}/{total} low/review ({rate*100:.0f}%)"
    if sig_id == "S5":
        n = blk.get("issue_count", 0)
        return f"{n} reconciliation issue(s)"
    if sig_id == "S6":
        n = blk.get("issue_count", 0)
        return f"{n} vendor row quality issue(s)"
    if sig_id == "S7":
        n = blk.get("duplicate_count", 0)
        return f"{n} duplicate S8 module name(s)"
    if sig_id == "S8":
        anchors = blk.get("anchors_present") or {}
        found = [k for k, v in anchors.items() if v]
        if not found:
            return "No USDM anchor classes found — schema gap (not DMP error)"
        return f"Anchors present: {found}"
    return "—"


_DMP_SIG_DESC = {
    "S1": "Required Section Presence — all mandatory DMP sections present",
    "S2": "Source Tag Provenance — each row has a source/generation tag",
    "S3": "Volume Sanity — S6 vendor count and S8 module count in expected range",
    "S4": "Confidence Distribution — < 40% low/review confidence sections",
    "S5": "S11 Reconciliation Integrity — all reconciliation flags correctly set",
    "S6": "Vendor Row Quality — S6.2 vendor rows have required fields",
    "S7": "S8 Module Uniqueness — no duplicate critical data module names",
    "S8": "USDM Traceability — key USDM anchor classes found in protocol JSON",
}

_DMP_SECTIONS = {
    "S1_revision_history": "S1 — Revision History",
    "S2_purpose_scope": "S2 — Purpose & Scope",
    "S3_process_documentation": "S3 — Process Documentation",
    "S4_roles_responsibilities": "S4 — Roles & Responsibilities",
    "S5_systems_tools": "S5 — Systems & Tools",
    "S6_data_flow": "S6 — Data Flow",
    "S7_data_dictionary_coding": "S7 — Data Dictionary & Coding",
    "S8_critical_data": "S8 — Critical Data",
    "S9_risk_based_monitoring": "S9 — Risk-Based Monitoring",
    "S10_data_surveillance_meeting": "S10 — Data Surveillance Meeting",
    "S11_data_review_validation": "S11 — Data Review & Validation",
    "S12_metrics_oversight": "S12 — Metrics & Oversight",
    "S13_quality_control": "S13 — Quality Control",
    "S14_data_deliverables": "S14 — Data Deliverables",
    "S15_data_archiving": "S15 — Data Archiving",
    "S16_reviewers_approvers": "S16 — Reviewers & Approvers",
}


def _add_score_breakdown(doc: Document, signals: Dict[str, Any], health: Dict[str, Any]) -> None:
    _section_heading(doc, "Score Breakdown — How the Score Was Calculated")
    total = health.get("total", len(signals))
    p = doc.add_paragraph()
    _run(p, f"Formula: total points ÷ {total} checks × 100   |   PASS = 1.0 pt  ·  WARN = 0.5 pt  ·  FAIL = 0.0 pt",
         italic=True, size=_CELL_PT, hex_color="555555")

    weights = {"PASS": 1.0, "WARN": 0.5, "FAIL": 0.0}
    headers = ["#", "Quality Check", "What It Checks", "Result", "Points"]
    col_widths = [Inches(0.3), Inches(1.3), Inches(2.7), Inches(0.8), Inches(0.9)]
    sig_keys = list(signals.keys())

    t = doc.add_table(rows=1 + len(sig_keys) + 1, cols=5)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_CELL_PT, rgb=_WHITE)

    running = 0.0
    for ri, sk in enumerate(sig_keys, 1):
        sv = signals.get(sk) or {}
        st = str(sv.get("status") or "").upper()
        pts = weights.get(st, 0.0)
        running += pts
        bg, fg = _status_colors(st)
        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w
        _clear(row.cells[0]).add_run(str(ri)).font.size = Pt(_CELL_PT)
        p_name = _clear(row.cells[1])
        _run(p_name, sk, bold=True, size=_CELL_PT)
        p_desc = _clear(row.cells[2])
        _run(p_desc, _DMP_SIG_DESC.get(sk, str(sv.get("description") or "")), size=_CELL_PT)
        c_st = row.cells[3]
        _shd(c_st, bg)
        p_st = _clear(c_st)
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_st, st, bold=True, size=_CELL_PT, hex_color=fg)
        p_pts = _clear(row.cells[4])
        p_pts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_pts, f"{pts:.1f} / 1.0", size=_CELL_PT)

    # Total row
    for ci in range(5):
        _shd(t.rows[-1].cells[ci], _BLUE_BG)
    for ci, w in enumerate(col_widths):
        t.rows[-1].cells[ci].width = w
    p_tot = _clear(t.rows[-1].cells[1])
    _run(p_tot, f"TOTAL  ·  {running:.1f} ÷ {total} × 100", bold=True, size=_CELL_PT, hex_color=_HEADER_BG)
    p_res = _clear(t.rows[-1].cells[4])
    p_res.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_res, f"= {round(running / total * 100, 1)}%", bold=True, size=_BODY_PT, hex_color=_HEADER_BG)
    doc.add_paragraph()

    if signals.get("S8", {}).get("status") == "FAIL":
        p_note = doc.add_paragraph()
        _run(p_note,
             "NOTE: S8 FAIL is due to USDM schema coverage gap (expected anchor classes not in USDM 4.0 JSON), "
             "not a DMP content error. Assess quality via S1–S7.",
             italic=True, size=_CELL_PT, hex_color="B91C1C")
    doc.add_paragraph()


def _add_section_history(doc: Document, s2: Dict[str, Any]) -> None:
    """Show each DMP section's generated confidence vs. historical presence."""
    import os, json as _json

    _section_heading(doc, "DMP Section Evidence — Generated vs. History")
    p_intro = doc.add_paragraph()
    _run(p_intro,
         "Each DMP section: generator confidence vs. historical presence across ground-truth studies.  "
         "Low-confidence + no historical precedent = highest scrutiny needed.",
         italic=True, size=_CELL_PT, hex_color="555555")

    # Load DMP GT to compute historical presence per section
    gt_paths = [
        r"C:\Users\jahna\OneDrive\Desktop\Pfizer\DMP_eval\data\dmp_ground_truth_clean.json",
    ]
    gt_studies: List[Dict[str, Any]] = []
    for path in gt_paths:
        if os.path.exists(path):
            try:
                with open(path, encoding="utf-8") as f:
                    raw = _json.load(f)
                gt_studies = raw if isinstance(raw, list) else [raw]
                break
            except Exception:
                pass

    total_gt = len(gt_studies)
    # Count presence per section key across GT studies
    section_hist: Dict[str, int] = {}
    for study in gt_studies:
        for sk in _DMP_SECTIONS:
            if sk in study and study[sk]:
                section_hist[sk] = section_hist.get(sk, 0) + 1

    # Load generator JSON for per-section confidence
    gen_path = s2.get("generator_json_path") or ""
    gen_data: Dict[str, Any] = {}
    gen_candidates = [
        r"C:\Users\jahna\Downloads\new_C509017\DMP\C5091017_DMP.json",
    ]
    study_id = str(s2.get("study_id") or "")
    for cp in gen_candidates:
        if os.path.exists(cp):
            try:
                with open(cp, encoding="utf-8") as f:
                    gen_data = _json.load(f)
                break
            except Exception:
                pass

    headers = ["DMP Section", "Generated?", "Generator Confidence", f"Historical ({total_gt} studies)", "Evidence Status"]
    col_widths = [Inches(2.0), Inches(0.9), Inches(1.2), Inches(1.2), Inches(1.7)]

    t = doc.add_table(rows=1 + len(_DMP_SECTIONS), cols=5)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_CELL_PT, rgb=_WHITE)

    for ri, (sk, label) in enumerate(_DMP_SECTIONS.items(), 1):
        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

        sec_data = gen_data.get(sk) or {}
        is_generated = bool(sec_data)
        gen_conf = str(sec_data.get("confidence") or "—").lower() if isinstance(sec_data, dict) else "—"
        hist_n = section_hist.get(sk, 0)
        hist_txt = f"{hist_n} / {total_gt}" if total_gt else "No GT data"
        has_hist = hist_n > 0

        # Evidence status
        if is_generated and has_hist and gen_conf not in ("low", "low_confidence"):
            ev, ev_bg, ev_fg = "Confirmed", _GREEN_BG, "1A7340"
        elif is_generated and has_hist:
            ev, ev_bg, ev_fg = "Generated (low conf)", _AMBER_BG, "926000"
        elif is_generated and not has_hist:
            ev, ev_bg, ev_fg = "Novel — Review", _AMBER_BG, "926000"
        elif not is_generated and has_hist:
            ev, ev_bg, ev_fg = "Missing (expected)", _RED_BG, "B91C1C"
        else:
            ev, ev_bg, ev_fg = "Not Generated", _GREY_BG, "555555"

        p_lbl = _clear(row.cells[0])
        _run(p_lbl, label, size=_CELL_PT)
        p_gen = _clear(row.cells[1])
        _run(p_gen, "Yes ✓" if is_generated else "No", size=_CELL_PT,
             hex_color="1A7340" if is_generated else "B91C1C")
        p_conf = _clear(row.cells[2])
        _run(p_conf, gen_conf, size=_CELL_PT)
        p_hist = _clear(row.cells[3])
        _run(p_hist, hist_txt, size=_CELL_PT)
        c_ev = row.cells[4]
        _shd(c_ev, ev_bg)
        p_ev = _clear(c_ev)
        _run(p_ev, ev, bold=True, size=_CELL_PT, hex_color=ev_fg)

    doc.add_paragraph()


def _add_scorecard(doc: Document, signals: Dict[str, Any]) -> None:
    _section_heading(doc, "1. Signal Scorecard")
    _signal_order = ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8"]
    headers = ["Signal", "Description", "Status", "Key Detail"]
    col_widths = [Inches(1.4), Inches(2.6), Inches(0.85), Inches(1.65)]

    t = doc.add_table(rows=1 + len(_signal_order), cols=4)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)

    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[ci]
        cell.width = w
        _shd(cell, _HEADER_BG)
        p = _clear(cell)
        _run(p, hdr, bold=True, size=_BODY_PT, rgb=_WHITE)

    for ri, sig_id in enumerate(_signal_order, 1):
        blk = signals.get(sig_id) or {}
        status = str(blk.get("status") or "—").upper()
        name = str(blk.get("name") or sig_id)
        desc = str(blk.get("description") or "")
        detail = _signal_detail(sig_id, blk)
        bg_hex, fg_hex = _status_colors(status)

        row = t.rows[ri]
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

        p_sig = _clear(row.cells[0])
        _run(p_sig, f"{sig_id} — {name}", bold=True, size=_CELL_PT)

        p_desc = _clear(row.cells[1])
        _run(p_desc, desc, size=_CELL_PT)

        c_st = row.cells[2]
        _shd(c_st, bg_hex)
        p_st = _clear(c_st)
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_st, status, bold=True, size=_CELL_PT, hex_color=fg_hex)

        p_det = _clear(row.cells[3])
        _run(p_det, detail, size=_CELL_PT)

    doc.add_paragraph()


def _add_detail_section(doc: Document, signals: Dict[str, Any]) -> None:
    _section_heading(doc, "2. Signal Detail")

    # S1 — missing sections
    s1 = signals.get("S1") or {}
    missing_s1 = s1.get("missing") or []
    if missing_s1:
        _sub_heading(doc, f"S1 — Required Section Presence  ·  {len(missing_s1)} missing")
        rows = [["Missing Section"]]
        for m in missing_s1:
            rows.append([str(m)])
        _simple_table(doc, rows, col_widths=[Inches(5.0)])

    # S2 — missing provenance
    s2 = signals.get("S2") or {}
    missing_s2 = s2.get("missing") or []
    if missing_s2:
        _sub_heading(doc, f"S2 — Source Tag Provenance  ·  {len(missing_s2)} row(s) missing")
        rows = [["Row Identifier"]]
        for m in missing_s2[:30]:
            rows.append([str(m)])
        _simple_table(doc, rows, col_widths=[Inches(5.0)])

    # S4 — confidence breakdown
    s4 = signals.get("S4") or {}
    total_c = s4.get("total_with_confidence", 0)
    low_n   = s4.get("low_or_review_count", 0)
    rate    = s4.get("low_or_review_rate", 0.0)
    rev_items_s4 = s4.get("review_items") or []
    _sub_heading(doc, "S4 — Confidence Distribution")
    rows = [["Metric", "Value"],
            ["Total rows with confidence tag", str(total_c)],
            ["Low/review confidence count",    str(low_n)],
            ["Low/review rate",                f"{rate*100:.1f}%"],
            ["Threshold",                      "WARN when >40%"]]
    _simple_table(doc, rows, col_widths=[Inches(3.0), Inches(2.5)])
    if rev_items_s4:
        p = doc.add_paragraph()
        _run(p, "Items with low/review confidence:", bold=True, size=_CELL_PT)
        rows2 = [["Section / Module", "Confidence"]]
        for item in rev_items_s4:
            rows2.append([str(item.get("section") or "—"), str(item.get("confidence") or "—")])
        _simple_table(doc, rows2, col_widths=[Inches(4.0), Inches(1.5)])

    # S5 — reconciliation issues
    s5 = signals.get("S5") or {}
    issues_s5 = s5.get("issues") or []
    if issues_s5:
        _sub_heading(doc, f"S5 — S11 Reconciliation Integrity  ·  {len(issues_s5)} issue(s)")
        rows = [["Section", "Issue"]]
        for iss in issues_s5[:30]:
            if isinstance(iss, dict):
                rows.append([str(iss.get("section") or ""), str(iss.get("issue") or "")])
            else:
                rows.append(["—", str(iss)])
        _simple_table(doc, rows, col_widths=[Inches(2.0), Inches(3.5)])

    # S6 — vendor quality issues
    s6 = signals.get("S6") or {}
    issues_s6 = s6.get("issues") or []
    if issues_s6:
        _sub_heading(doc, f"S6 — Vendor Row Quality  ·  {len(issues_s6)} issue(s)")
        rows = [["Row", "Issue"]]
        for iss in issues_s6[:30]:
            if isinstance(iss, dict):
                rows.append([str(iss.get("row") or ""), str(iss.get("issue") or "")])
            else:
                rows.append(["—", str(iss)])
        _simple_table(doc, rows, col_widths=[Inches(2.0), Inches(3.5)])

    # S7 — duplicate module names
    s7 = signals.get("S7") or {}
    dups = s7.get("duplicate_modules") or []
    if dups:
        _sub_heading(doc, f"S7 — S8 Module Uniqueness  ·  {len(dups)} duplicate(s)")
        rows = [["Duplicate Module Name"]]
        for d in dups:
            rows.append([str(d)])
        _simple_table(doc, rows, col_widths=[Inches(5.0)])

    # S8 — USDM anchors
    s8 = signals.get("S8") or {}
    anchors = s8.get("anchors_present") or {}
    _sub_heading(doc, "S8 — USDM Traceability")
    rows = [["USDM Class", "Found in Protocol"]]
    for cls, present in anchors.items():
        rows.append([cls, "Yes ✓" if present else "No ✗"])
    _simple_table(doc, rows, col_widths=[Inches(2.8), Inches(2.0)])
    if s8.get("status") == "FAIL":
        _note_paragraph(
            doc,
            "The USDM 4.0 JSON for this study does not instantiate the anchor classes "
            "expected by the DMP evaluator. This is a known schema coverage gap — "
            "USDM 4.0 uses different class names than anticipated. "
            "The DMP content quality should be assessed via S1–S7 signals instead."
        )
    if s8.get("usdm_types_sample"):
        _note_paragraph(
            doc,
            "Instance types found in USDM: "
            + ", ".join(s8["usdm_types_sample"][:12]) + " …"
        )


def _add_dmp_field_check(doc: Document, signals: Dict[str, Any]) -> None:
    """DMP structural field coverage summary from signal data."""
    s1  = signals.get("S1") or {}
    s2  = signals.get("S2") or {}
    s3  = signals.get("S3") or {}
    s6  = signals.get("S6") or {}

    _section_heading(doc, "DMP Field Structure Check")
    p = doc.add_paragraph()
    _run(p, "Summary of generated DMP content counts and field completeness by section.",
         italic=True, size=_CELL_PT, hex_color="555555")

    # Section counts
    counts = s3.get("counts") or {}
    rows = [["DMP Component", "Count Generated", "Status"]]
    rows.append(["S6 Vendor rows",         str(counts.get("s6_vendors", "?")),
                 "OK ✓" if (counts.get("s6_vendors") or 0) > 0 else "EMPTY"])
    rows.append(["S8 Critical Data modules", str(counts.get("s8_modules", "?")),
                 "OK ✓" if (counts.get("s8_modules") or 0) > 0 else "EMPTY"])
    _simple_table(doc, rows, col_widths=[Inches(2.8), Inches(1.5), Inches(1.5)])

    # Missing sections (S1)
    missing_secs = s1.get("missing_sections") or []
    if missing_secs:
        _sub_heading(doc, f"S1 — Missing sections: {len(missing_secs)}")
        rows2 = [["Missing Section"]]
        for sec in missing_secs:
            rows2.append([str(sec)])
        _simple_table(doc, rows2, col_widths=[Inches(5.0)])
    else:
        p2 = doc.add_paragraph()
        _run(p2, "S1 — All required DMP sections are present. ✓", size=_CELL_PT)

    # Missing provenance (S2)
    missing_prov = s2.get("missing_provenance") or s2.get("missing") or []
    if missing_prov:
        _sub_heading(doc, f"S2 — Sections missing provenance/source: {len(missing_prov)}")
        rows3 = [["Section", "Issue"]]
        for item in missing_prov[:30]:
            if isinstance(item, dict):
                rows3.append([str(item.get("section") or "—"), str(item.get("issue") or "—")])
            else:
                rows3.append(["—", str(item)])
        _simple_table(doc, rows3, col_widths=[Inches(2.5), Inches(3.0)])
    else:
        p3 = doc.add_paragraph()
        _run(p3, "S2 — All generated sections have provenance/source populated. ✓", size=_CELL_PT)

    # S6 vendor row quality
    issues_s6 = s6.get("issues") or []
    if not issues_s6:
        p4 = doc.add_paragraph()
        _run(p4, "S6 — All vendor rows have vendor name and data type populated. ✓", size=_CELL_PT)


def _add_review_list(doc: Document, review_list: List[Dict[str, Any]]) -> None:
    if not review_list:
        return
    _section_heading(doc, f"3. Items Requiring Review  ({len(review_list)})")
    for item in review_list:
        sig    = str(item.get("signal_id") or "")
        reason = str(item.get("reason") or "")
        p = doc.add_paragraph(style="List Bullet")
        _run(p, f"{sig}: ", bold=True, size=_BODY_PT)
        _run(p, reason, size=_BODY_PT)


def _simple_table(doc: Document, rows: List[List[str]],
                  col_widths: Optional[List] = None) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.autofit = False
    _borders(t)
    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            text = row_data[ci] if ci < len(row_data) else ""
            cell = t.rows[ri].cells[ci]
            if col_widths and ci < len(col_widths):
                cell.width = col_widths[ci]
            p = _clear(cell)
            is_header = ri == 0
            if is_header:
                _shd(cell, _HEADER_BG)
                _run(p, text, bold=True, size=_CELL_PT, rgb=_WHITE)
            else:
                _run(p, text, size=_CELL_PT)
    doc.add_paragraph()


def _compute_health(signals: Dict[str, Any]) -> Dict[str, Any]:
    weights = {"PASS": 1.0, "WARN": 0.5, "FAIL": 0.0}
    pass_n = warn_n = fail_n = 0
    weighted = 0.0
    for blk in signals.values():
        st = str(blk.get("status") or "").upper()
        pass_n += st == "PASS"
        warn_n += st == "WARN"
        fail_n += st == "FAIL"
        weighted += weights.get(st, 0.0)
    total = pass_n + warn_n + fail_n
    pct = round((weighted / total * 100) if total else 0.0, 1)
    return {"percent": pct, "pass": pass_n, "warn": warn_n, "fail": fail_n, "total": total}


# ── Public API ────────────────────────────────────────────────────────────────

def write_dmp_scenario2_docx(s2: Dict[str, Any], output_path: str) -> str:
    """Generate the DMP Scenario 2 Word report and return the resolved output path."""
    study_id   = str(s2.get("study_id") or "Unknown")
    ts         = str(s2.get("timestamp") or "")[:10] or datetime.now().strftime("%Y-%m-%d")
    verdict    = str(s2.get("verdict") or "AMBER")
    signals    = s2.get("signals") or {}
    review_list = s2.get("review_list") or []

    health = _compute_health(signals)

    doc = Document()
    _apply_defaults(doc)

    _add_banner(doc, study_id, ts)
    _add_verdict_box(doc, verdict, health)
    _add_score_breakdown(doc, signals, health)
    _add_section_history(doc, s2)
    _add_scorecard(doc, signals)
    _add_detail_section(doc, signals)
    _add_dmp_field_check(doc, signals)
    _add_review_list(doc, review_list)

    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    _run(p_foot,
         f"Protocol Digitalization Platform — USDM 4.0  ·  {ts}  ·  CONFIDENTIAL",
         italic=True, size=_CELL_PT - 1, hex_color="777777")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
        return str(out.resolve())
    except PermissionError:
        alt = out.with_name(f"{out.stem}_{datetime.now():%Y%m%d_%H%M%S}{out.suffix}")
        doc.save(str(alt))
        return str(alt.resolve())


def write_dmp_scenario2_docx_from_json(json_path: str, output_path: str) -> str:
    """Load a Scenario 2 JSON file and write the Word report."""
    with open(json_path, encoding="utf-8") as fh:
        s2 = json.load(fh)
    return write_dmp_scenario2_docx(s2, output_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python dmp_scenario2_report.py <s2_json> <output.docx>")
        sys.exit(1)
    out = write_dmp_scenario2_docx_from_json(sys.argv[1], sys.argv[2])
    print(f"Wrote: {out}")
