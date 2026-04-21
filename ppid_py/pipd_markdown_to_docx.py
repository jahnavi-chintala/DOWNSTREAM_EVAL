"""
Convert Markdown (subset) to .docx for PIPD reports.

Supports: #–#### headings, bullets (-/*), paragraphs, fenced ``` blocks,
simple GFM tables (| col | col |).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


_FONT_NAME = "Arial"


def _apply_doc_defaults(doc: Document) -> None:
    """Arial font, narrow margins — applied once per document."""
    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = Pt(10)
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)


def _set_thin_table_borders(table) -> None:
    """Apply light-gray 0.5 pt borders to the whole table via XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove existing tblBorders if any
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 0.5 pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _clean_md_for_cell(s: str) -> str:
    """Strip simple Markdown from table cells for readable Word text."""
    t = (s or "").replace("\\|", "|")
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    if len(t) >= 2 and t.startswith("_") and t.endswith("_") and not t.startswith("__"):
        t = t[1:-1]
    return t.strip()


def _add_mixed_runs(paragraph: Paragraph, text: str, base_italic: bool = False) -> None:
    """Split on **bold** segments; optional base italic for the whole paragraph."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.italic = base_italic
        else:
            r = paragraph.add_run(part)
            r.italic = base_italic


def _add_body_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    """Paragraph with optional outer _italic_ and inner **bold**."""
    t = text.replace("\\|", "|")
    base_it = False
    if len(t) >= 2 and t.startswith("_") and t.endswith("_"):
        t = t[1:-1]
        base_it = True
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    _add_mixed_runs(p, t, base_italic=base_it)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_color(run, hex_color: str) -> None:
    """Set run font color via XML (works reliably across python-docx versions)."""
    rPr = run._r.get_or_add_rPr()
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), hex_color)
    rPr.append(color_el)


def _shape_reference_eval_table(table: Table, rows: List[List[str]]) -> None:
    """Fixed column widths for PIPD eval summary/scorecard tables (reference Word layout)."""
    if not rows:
        return
    header = [_clean_md_for_cell(rows[0][i]) for i in range(len(rows[0]))]
    n = len(header)
    h0 = (header[0] or "").lower()
    widths: Optional[List] = None
    if n == 5 and h0.startswith("metric"):
        widths = [Inches(2.5), Inches(0.82), Inches(0.72), Inches(0.85), Inches(1.45)]
    elif n == 6 and h0.startswith("category"):
        # §2 scorecard: Category | Weight | Weighted | Score | Subcats | Status
        widths = [Inches(1.95), Inches(0.62), Inches(0.62), Inches(0.72), Inches(0.55), Inches(1.85)]
    elif n == 6 and "subcategory" in h0:
        # Detail table after GSOP column removal: Subcat | Score | Match | YES/NO | USDM | Conf.
        widths = [Inches(2.2), Inches(0.5), Inches(0.6), Inches(0.6), Inches(1.05), Inches(0.5)]
    elif n == 7 and "subcategory" in h0:
        # Legacy 7-col with GSOP (kept for compatibility)
        widths = [Inches(2.0), Inches(0.45), Inches(0.55), Inches(0.55), Inches(0.85), Inches(0.45), Inches(1.15)]
    if not widths:
        return
    for tr in table.rows:
        for ci, w in enumerate(widths):
            if ci < len(tr.cells):
                tr.cells[ci].width = w


def _parse_doc_score_params(encoded: str) -> dict:
    """Parse key=value|key=value params from a <!-- DOC_SCORE: ... --> comment."""
    d: dict = {}
    # Split only on | that precede a known key= pattern so label values with spaces work
    import re as _re
    parts = _re.split(r"\|(?=\w+=)", encoded)
    for part in parts:
        if "=" in part:
            k, _, v = part.partition("=")
            d[k.strip()] = v.strip()
    return d


def _add_doc_score_table(doc: Document, encoded: str) -> None:
    """Render the Document Score as a styled 3-panel table matching the reference spec."""
    p = _parse_doc_score_params(encoded)
    seg = p.get("seg", "")
    score = p.get("score", "—")
    passed = p.get("pass", "0") == "1"
    threshold = p.get("threshold", "")
    target = p.get("target", "")
    label = p.get("label", "").rstrip()

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    col_widths = [Inches(2.0), Inches(1.5), Inches(2.0)]
    for ci, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[ci].width = w

    # ── Cell 0: DOCUMENT SCORE label + segment ───────────────────────────────
    c0 = table.rows[0].cells[0]
    _set_cell_shading(c0, "D9E2F3")
    for p0 in c0.paragraphs:
        p0.clear()
    p0 = c0.paragraphs[0]
    r0 = p0.add_run("DOCUMENT SCORE")
    r0.bold = True
    r0.font.size = Pt(11)
    _set_run_color(r0, "003087")
    if seg:
        p0b = c0.add_paragraph()
        r0b = p0b.add_run(seg)
        r0b.font.size = Pt(9)
        _set_run_color(r0b, "003087")

    # ── Cell 1: score value + "/ 100" + optional label ───────────────────────
    c1 = table.rows[0].cells[1]
    _set_cell_shading(c1, "EEF7F1")
    for p1 in c1.paragraphs:
        p1.clear()
    p1a = c1.paragraphs[0]
    r1a = p1a.add_run(str(score))
    r1a.bold = True
    r1a.font.size = Pt(28)
    _set_run_color(r1a, "186A5A")
    p1b = c1.add_paragraph()
    r1b = p1b.add_run("/ 100")
    r1b.font.size = Pt(9)
    _set_run_color(r1b, "555555")
    if label:
        p1c = c1.add_paragraph()
        r1c = p1c.add_run(label)
        r1c.italic = True
        r1c.font.size = Pt(8)
        _set_run_color(r1c, "555555")

    # ── Cell 2: PASS/FAIL + threshold + target ────────────────────────────────
    c2 = table.rows[0].cells[2]
    for p2 in c2.paragraphs:
        p2.clear()
    pass_color = "186A5A" if passed else "C0392B"
    pass_txt = "PASS ✓" if passed else "FAIL ✗"
    p2a = c2.paragraphs[0]
    r2a = p2a.add_run(pass_txt)
    r2a.bold = True
    r2a.font.size = Pt(12)
    _set_run_color(r2a, pass_color)
    if threshold:
        p2b = c2.add_paragraph()
        r2b = p2b.add_run(f"Threshold: {threshold}")
        r2b.font.size = Pt(9)
        _set_run_color(r2b, "555555")
    if target:
        p2c = c2.add_paragraph()
        r2c = p2c.add_run(f"Target: {target}")
        r2c.font.size = Pt(9)
        _set_run_color(r2c, "555555")

    doc.add_paragraph()


def _add_table(
    doc: Document,
    rows: List[List[str]],
    *,
    reference_eval: bool = False,
    first_row_is_header: bool = True,
) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    _set_thin_table_borders(table)
    for ri, row_data in enumerate(rows):
        for ci in range(cols):
            text = row_data[ci] if ci < len(row_data) else ""
            cell = table.rows[ri].cells[ci]
            plain = _clean_md_for_cell(text)
            for p in cell.paragraphs:
                p.clear()
            p = cell.paragraphs[0]
            r = p.add_run(plain)
            r.font.name = _FONT_NAME
            r.font.size = Pt(9)
            if first_row_is_header and ri == 0:
                r.bold = True
                _set_cell_shading(cell, "D9E2F3")
    if reference_eval:
        _shape_reference_eval_table(table, rows)
    doc.add_paragraph()


def _split_table_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    core = s.replace("|", "")
    return bool(re.match(r"^[\s:\-]+$", core))


def _reference_section_paragraph(doc: Document, text: str) -> None:
    """Bold section label (Normal style) — matches reference Word (no Heading styles)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_banner_box(doc: Document, banner_lines: List[str]) -> None:
    """Render the PIPD report banner inside a single-cell shaded box table."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]

    # Dark blue background for header cell
    _set_cell_shading(cell, "003087")

    # Clear default paragraph, add styled content
    for p in cell.paragraphs:
        p.clear()

    first = True
    for raw in banner_lines:
        stripped = raw.strip()
        if not stripped:
            if not first:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
            continue

        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False

        # Italic line (eval date)
        is_italic = stripped.startswith("_") and stripped.endswith("_")
        text = stripped[1:-1] if is_italic else stripped
        # Bold markers
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)

        r = p.add_run(text)
        r.font.color.rgb = None
        from docx.dml.color import ColorFormat
        from docx.oxml import OxmlElement as _OE
        rPr = r._r.get_or_add_rPr()
        color_el = _OE("w:color")
        color_el.set(qn("w:val"), "FFFFFF")
        rPr.append(color_el)
        r.font.size = Pt(10 if is_italic else 11)
        r.bold = not is_italic
        r.italic = is_italic

    doc.add_paragraph()


def write_docx_from_markdown(
    markdown_text: str,
    output_path: str,
    *,
    reference_eval: bool = False,
) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    base_dir = path.parent

    doc = Document()
    _apply_doc_defaults(doc)

    lines = markdown_text.replace("\r\n", "\n").split("\n")

    # Pre-pass: collect banner lines (before ## DOCUMENT SCORE) for box rendering
    banner_end_idx = None
    if reference_eval:
        for idx, ln in enumerate(lines):
            if ln.strip().startswith("## DOCUMENT SCORE"):
                banner_end_idx = idx
                break
    if banner_end_idx and banner_end_idx > 0:
        banner_block = lines[:banner_end_idx]
        lines = lines[banner_end_idx:]
        _add_banner_box(doc, banner_block)

    i = 0
    in_code = False
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph("\n".join(code_lines) if code_lines else "")
                para.paragraph_format.left_indent = Pt(12)
                for r in para.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        doc_score_m = re.match(r"<!--\s*DOC_SCORE:\s*(.+?)\s*-->", stripped)
        if doc_score_m:
            if reference_eval:
                _add_doc_score_table(doc, doc_score_m.group(1))
            i += 1
            continue

        img_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_m:
            alt = (img_m.group(1) or "Figure").strip()
            rel = img_m.group(2).strip()
            p_img = Path(rel)
            if not p_img.is_file():
                p_img = base_dir / rel
            if p_img.is_file():
                cap = doc.add_paragraph()
                r = cap.add_run(alt)
                r.bold = True
                doc.add_picture(str(p_img), width=Inches(6.2))
            else:
                doc.add_paragraph(f"[Figure missing: {rel}]")
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: List[List[str]] = []
            header_sep_seen = False
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_separator(row_line):
                    if len(table_rows) == 1:
                        header_sep_seen = True
                    i += 1
                    continue
                table_rows.append(_split_table_row(row_line))
                i += 1
            _add_table(
                doc,
                table_rows,
                reference_eval=reference_eval,
                first_row_is_header=header_sep_seen,
            )
            continue

        if stripped.startswith("#### "):
            if reference_eval:
                _reference_section_paragraph(doc, stripped[5:].strip())
            else:
                doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            if reference_eval:
                _reference_section_paragraph(doc, stripped[4:].strip())
            else:
                doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            if reference_eval:
                _reference_section_paragraph(doc, stripped[3:].strip())
            else:
                doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            if reference_eval:
                _reference_section_paragraph(doc, stripped[2:].strip())
            else:
                doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_body_paragraph(doc, stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            _add_body_paragraph(doc, re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph()
        elif stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            _add_body_paragraph(doc, quote)
            doc.paragraphs[-1].paragraph_format.left_indent = Pt(18)
        else:
            if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
                p = doc.add_paragraph()
                r = p.add_run(stripped[2:-2])
                r.bold = True
            else:
                _add_body_paragraph(doc, stripped)
        i += 1

    doc.save(str(path))
